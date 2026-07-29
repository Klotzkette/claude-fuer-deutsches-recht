#!/usr/bin/env python3
"""Prüft die formale Dokumentqualität zentraler juristischer Aktenbestände."""

from __future__ import annotations

import hashlib
import csv
import re
import sys
from collections import defaultdict
from email import policy
from email.parser import BytesParser
from pathlib import Path

from docx import Document
from openpyxl import load_workbook
from pypdf import PdfReader

from testakte_file_filter import META_EXACT_NAMES, include_in_working_dump


REPO = Path(__file__).resolve().parent.parent
TESTAKTEN = REPO / "testakten"
PREFIXES = (
    "arbeitsrecht-",
    "arbeitszeugnis-",
    "elternunterhalt-",
    "erbrecht-",
    "erbstreit-",
    "familienrecht-",
    "kuendigungsschutzklage-",
    "longcovid-erwerbsminderung-",
    "nachehelicher-unterhalt-",
    "rentenberater-",
    "rentenrecht-",
    "schwerbehindertenrecht-",
    "selbstvertreter-sozialgericht-",
    "sozialrecht-",
    "statusfeststellung-",
    "unfallversicherung-",
    "unterhalt-",
    "versausgleich-",
)
FORMAL_EXTS = {".docx", ".eml", ".pdf"}
EXPORT_TEXT_EXTS = {
    ".csv",
    ".docx",
    ".eml",
    ".htm",
    ".html",
    ".json",
    ".pdf",
    ".rtf",
    ".tsv",
    ".txt",
    ".xlsm",
    ".xlsx",
    ".xml",
}
MIN_FORMAL_TEXT = 600
META_MARKERS = (
    "beispielakte",
    "demoakte",
    "lorem ipsum",
    "platzhalter",
    "testdaten",
    "testfall",
    "formathinweis",
    "dokumententyp",
    "für testzwecke",
    "testakte",
    "enthält bewusst",
    "bewusst keine",
    "bewusst mehrere",
)
EXPORT_META_PATTERNS = {
    "Testkennzeichnung": re.compile(
        r"\b(?:testakte|testdokument|test-dokument|testmaterial|plugin-test|plugin-testakte)\b",
        re.IGNORECASE,
    ),
    "Vorführkennzeichnung": re.compile(
        r"\b(?:demonstrationsdokument|demonstrationsauszug|demonstrations-testakte|demonstrationszweck(?:en)?)\b",
        re.IGNORECASE,
    ),
    "Übungskennzeichnung": re.compile(
        r"\b(?:übungsakte|uebungsakte|ausbildungszweck(?:en)?|lernakte|ki-kurs)\b",
        re.IGNORECASE,
    ),
    "Fiktionshinweis": re.compile(
        r"(?:alle|sämtliche)\s+[^\n.]{0,120}\b(?:fiktiv|erfunden)\b|"
        r"\bfiktives?\s+(?:beispiel|test|übungs|uebungs|lern)[\w-]*|"
        r"\bkein realer mandatsbezug\b",
        re.IGNORECASE,
    ),
    "Arbeitsmarker": re.compile(r"\bTODO\b|\[AZ\s+fiktiv\]|\[fiktiv\]", re.IGNORECASE),
    "Austauschbares Aktenvorblatt": re.compile(
        r"die akte beginnt nicht mit einem fertigen rechtsproblem|"
        r"eignet sich der vorgang gut, um nicht nur abstrakt zu prüfen|"
        r"bitte heute nicht nur eine kurznotiz schicken|"
        r"eine erste zuordnung zu:|"
        r"für die nächste bearbeitung fehlen noch die unten markierten unterlagen|"
        r"ärztliche oder wirtschaftliche kernbelege",
        re.IGNORECASE,
    ),
    "Austauschbare Kurzunterlage": re.compile(
        r"diese unterlage bündelt (?:den arbeitsstand|die lage) für die (?:nächste|erste) besprechung|"
        r"was muss entschieden werden\?\s*welche quelle fehlt\?|"
        r"für diese akte liegt (?:noch )?keine echte|"
        r"dieser aktenvermerk ersetzt keinen echten registerabruf",
        re.IGNORECASE,
    ),
}
SYNTHETIC_EMAIL_PATTERN = re.compile(
    r"(?:@|https?://)[^\s<>/]+[.](?:example|invalid|local|test)\b|"
    r"\b[a-z0-9][a-z0-9.-]*[.]example\b|"
    r"@example[.](?:com|de|org)\b|"
    r"\bbeispielkanzlei\b|"
    r"\bmandat[.]local\b",
    re.IGNORECASE,
)
BROKEN_ENCODING_MARKERS = ("\ufffd", "Ã", "Â", "â€", "ðŸ")
REMOVED_AGGREGATES = {
    "arbeitsrecht-kuendigungsdrama-koerber-werk/03_arbeitsvertrag_at_koerber_2012.docx",
    "arbeitsrecht-kuendigungsdrama-koerber-werk/04_organigramm_und_stellenbeschreibung.docx",
    "arbeitsrecht-kuendigungsdrama-koerber-werk/06_hinschg_anzeige_mai_2024.docx",
    "arbeitsrecht-kuendigungsdrama-koerber-werk/07_korrespondenz_kemnitz_freundschaftliche_hinweise.docx",
    "arbeitsrecht-kuendigungsdrama-koerber-werk/08_aufhebungsvertrag_entwurf_steinhoff_mai_2025.docx",
    "arbeitsrecht-kuendigungsdrama-koerber-werk/09_kuendigungsschreiben_22_05_2025.docx",
    "arbeitsrecht-kuendigungsdrama-koerber-werk/10_betriebsratsanhoerung_protokoll.docx",
    "arbeitsrecht-kuendigungsdrama-koerber-werk/12_bem_einladung_nichts_passiert.docx",
    "arbeitsrecht-kuendigungsdrama-koerber-werk/13_entgelttransparenz_anfrage_und_antwort.docx",
    "arbeitsrecht-kuendigungsdrama-koerber-werk/18_emails_branitz_aufhebungsdruck.docx",
    "arbeitsrecht-kuendigungsdrama-koerber-werk/pdfs/au_bescheinigungen_juli_2024_april_2025.pdf",
    "arbeitszeugnis-analyse-bluehendes-leben/90-ergaenzende-korrespondenz-und-vollvermerke.docx",
    "erbrecht-digitaler-nachlass-passwortsafe-berlin/05_testament_und_notizzettel.docx",
    "erbrecht-digitaler-nachlass-passwortsafe-berlin/16_vollmacht_und_postausgang.docx",
    "erbrecht-patchwork-stiefgrossvater-erbfolge-karlsruhe/02_testamente_und_scheidungen.docx",
    "erbrecht-patchwork-stiefgrossvater-erbfolge-karlsruhe/14_sterbeurkunden_abschriften.docx",
    "erbrecht-patchwork-stiefgrossvater-erbfolge-karlsruhe/15_nachlassgericht_anhaerung_und_verfuegung.docx",
    "erbrecht-patchwork-stiefgrossvater-erbfolge-karlsruhe/16_grundbuch_und_bankauskunft_armin.docx",
    "erbrecht-pflichtteilsergaenzung-hofuebertragung-lueneburg/04_immobilienwerte_und_depot.docx",
    "erbrecht-pflichtteilsergaenzung-hofuebertragung-lueneburg/08_bank_depot_und_schenkungsverkehr.docx",
    "erbrecht-pflichtteilsergaenzung-hofuebertragung-lueneburg/12_eroeffnungsniederschrift_nachlassgericht.docx",
    "erbrecht-pflichtteilsergaenzung-hofuebertragung-lueneburg/99_nachtrag_wohnrecht_pflege_ackerpacht.docx",
    "erbrecht-pflichtteilsergaenzung-hofuebertragung-lueneburg/02_testament_und_nachlassverzeichnis.docx",
    "erbrecht-pflichtteilsergaenzung-niesbrauch-depot-sylt/12_vollmacht_und_postausgang.docx",
    "erbrecht-volljaehrigenadoption-unternehmer-erfurt/02_adoptionsbeschluss_und_beziehung.docx",
    "erbrecht-volljaehrigenadoption-unternehmer-erfurt/03_testament_und_gesellschaft.docx",
    "erbrecht-volljaehrigenadoption-unternehmer-erfurt/04_pflichtteil_und_auskunft.docx",
    "erbstreit-krypto-multisig-edelmann-stuttgart/18_korrespondenz_marlies_anwalt_strecker.docx",
    "elternunterhalt-pflegeheim-sozialamtsregress-aachen/01_mandatsnotiz_und_rechtswahrungsanzeige_sozialamt.docx",
    "elternunterhalt-pflegeheim-sozialamtsregress-aachen/12_strategie_und_antwortschreiben_sozialamt.docx",
    "familienrecht-sorge-umgang-gewaltschutz-essen/10_jugendamt_hausbesuch_und_hilfeplan.docx",
    "familienrecht-unterhalt-zugewinn-bad-nauheim/04_einkommen_elena_arbeitgeber_und_nebentaetigkeit.docx",
    "familienrecht-unterhalt-zugewinn-bad-nauheim/06_selbststaendigkeit_martin_konten_und_rechnungen.docx",
    "familienrecht-wechselmodell-unterhalt-selbststaendige-kiel/pdfs/schulbescheinigung_und_kostenblatt.pdf",
    "longcovid-erwerbsminderung-feldermann-leipzig/04_bescheid_ukbw_bk3101_28_01_2026.docx",
    "longcovid-erwerbsminderung-feldermann-leipzig/06_bescheid_drv_em_rente.docx",
    "longcovid-erwerbsminderung-feldermann-leipzig/08_bescheid_lasov_gdb30.docx",
    "longcovid-erwerbsminderung-feldermann-leipzig/10_bescheid_jobcenter_sgbii.docx",
    "longcovid-erwerbsminderung-feldermann-leipzig/15_befund_charite_longcovid_ambulanz.docx",
    "longcovid-erwerbsminderung-feldermann-leipzig/16_befund_skh_psychiatrie.docx",
    "longcovid-erwerbsminderung-feldermann-leipzig/18_email_kette_sonnemann_traeger.docx",
    "longcovid-erwerbsminderung-feldermann-leipzig/19_mandantenkommunikation.docx",
    "longcovid-erwerbsminderung-feldermann-leipzig/pdfs/befund_charite_longcovid_redacted.pdf",
    "longcovid-erwerbsminderung-feldermann-leipzig/pdfs/bescheid_drv_em_rente_redacted.pdf",
    "selbstvertreter-sozialgericht-heizkosten-eilantrag/01_bescheid_jobcenter_auszug.docx",
    "selbstvertreter-sozialgericht-heizkosten-eilantrag/02_widerspruch_und_widerspruchsbescheid.docx",
    "selbstvertreter-sozialgericht-heizkosten-eilantrag/03_miete_heizung_und_konto.docx",
    "selbstvertreter-sozialgericht-heizkosten-eilantrag/06_attest_und_schulbescheinigung.docx",
    "selbstvertreter-sozialgericht-heizkosten-eilantrag/07_telefonnotizen_jobcenter.docx",
    "rentenrecht-riester-zulagenrueckforderung-augsburg/03_eigenbeitraege_und_zulagen.docx",
    "rentenrecht-riester-zulagenrueckforderung-augsburg/08_anbieter_jahresbescheinigung_und_kontoauszug.docx",
    "rentenrecht-waisenrente-ausbildung-abbruch-dortmund/04_gesundheit_und_abbruch.docx",
    "rentenrecht-waisenrente-ausbildung-abbruch-dortmund/11_studienzusage_und_praxisvertrag.docx",
    "rentenrecht-witwenrente-einkommensanrechnung-luebeck/04_kvdr_und_beitragsabzug.docx",
    "rentenrecht-witwenrente-einkommensanrechnung-luebeck/08_arbeitgeberbescheinigung_und_lohnabrechnung.docx",
    "sozialrecht-elektrorollstuhl-koerner-oldenburg/02_ueberweisung_und_hilfsmittelverordnung.docx",
    "sozialrecht-enzymsubstitution-seltene-erkrankung-jena/03_genetik_und_verlauf.docx",
    "sozialrecht-enzymsubstitution-seltene-erkrankung-jena/04_kassenbescheid_md.docx",
    "sozialrecht-orphan-drug-krebsmedikament-muenster/03_tumorboard_und_nutzenblatt.docx",
    "sozialrecht-orphan-drug-krebsmedikament-muenster/05_kassenbescheid.docx",
    "sozialrecht-orphan-drug-krebsmedikament-muenster/06_widerspruch_und_eilantrag.docx",
    "sozialrecht-rollstuhl-tannenberg/01-olaf-rollstuhl/Korrespondenz_mit_Nordsee-BKK.pdf",
    "sozialrecht-rollstuhl-tannenberg/01-olaf-rollstuhl/Aerztliches_Attest_Wallenstein_05-05-2026.pdf",
    "sozialrecht-rollstuhl-tannenberg/01-olaf-rollstuhl/Bescheid_Nordsee-BKK_18-04-2026.pdf",
    "sozialrecht-rollstuhl-tannenberg/01-olaf-rollstuhl/Kostenvoranschlag_Sanitaetshaus_Reha-Aktiv-Nord.pdf",
    "sozialrecht-rollstuhl-tannenberg/01-olaf-rollstuhl/MDK-Gutachten_03-04-2026.pdf",
    "sozialrecht-rollstuhl-tannenberg/01-olaf-rollstuhl/Pflegegrad_2_Bescheid_04-05-2023.pdf",
    "sozialrecht-rollstuhl-tannenberg/01-olaf-rollstuhl/Reha-Bericht_2024_Damp.pdf",
    "sozialrecht-rollstuhl-tannenberg/01-olaf-rollstuhl/Verordnung_Muster16_09-02-2026.pdf",
    "sozialrecht-schwerbehindertenverfahren-gdb-merkzeichen-muenster/12_akteneinsicht_versorgungsaerztliche_stellungnahme.docx",
    "sozialrecht-wohnraumanpassung-rampe-pflegegrad-wittenberge/02_pflegegrad_und_md_gutachten.docx",
    "sozialrecht-wohnraumanpassung-rampe-pflegegrad-wittenberge/04_pflegekasse_bescheid.docx",
    "sozialrecht-wohnraumanpassung-rampe-pflegegrad-wittenberge/05_widerspruch_und_sozialhilfe.docx",
    "sozialrecht-wohnraumanpassung-rampe-pflegegrad-wittenberge/08_eilrechtsschutz_und_ortstermin.docx",
    "unfallversicherung-arbeitsunfall-lagerleiter-sturz-trier/03_durchgangsarztbericht.docx",
    "unfallversicherung-arbeitsunfall-lagerleiter-sturz-trier/04_mrt_und_vorschaden.docx",
    "unfallversicherung-arbeitsunfall-lagerleiter-sturz-trier/06_bg_ablehnungsbescheid.docx",
    "unfallversicherung-arbeitsunfall-lagerleiter-sturz-trier/08_widerspruch_entwurf.docx",
    "unterhalt-berechnungsakte-vollstaendig-rosenheim/12_emailverkehr_zahlenfreigabe.docx",
}
PROTECTED_PROSE = re.compile(
    r"https?://[^\s<>()]+|"
    r"[\w.+-]+@[\w.-]+|"
    r"\b[^\s/\\]+\.(?:pdf|docx|odt|xlsx|csv|eml|jpg|jpeg|png|json|yaml)\b",
    re.IGNORECASE,
)
TRANSLITERATION = re.compile(
    r"(?:(?:fuer|ueber|kuendig|selbststaendig|vollstaendig|praezis|"
    r"geschaeft|beschaeft|verhaeltn|einschaetz|begruend|beruecksichtig|"
    r"erklaer|bestaet|verfueg|taetig|taeg|regelmaess|gemaess|grundsaetz|"
    r"tatsaech|persoen|moeg|faeh|vermoeg|zusaetz|unabhaeng|eingeschraenk|"
    r"beeintraechtig|zuverlaess|rechtskraeft|aerzt|sanitaet|mobilitaet|"
    r"liquiditaet|orthopaed|gespraech|stoer|schaed|gewaehr|zurueck|rueck|"
    r"schluessel|tuer|kuech|hueft|stuerz|naech|angehoer|behoerd|oeffn|"
    r"muenster|luebeck|wuerzburg|muenchen|nuernberg|fuerth|koerber|koerner|"
    r"schaefer|soell|roedelsteiner|zahn?aerzte|dienstbezueg|zukuenftig|"
    r"maerz|gruess|schliess|anschliess|beruehr|sozietaet|saetze|auskuenft|"
    r"aktenueber|erfaehr|verfahrensfuehr|gewuerdig|rechtsanwaelt|ruhegehaltfaeh|"
    r"haelft|spaet|einschlaeg|sondervermoeg|gefoerdert|eheschliess|"
    r"wirtschaftspruef|innenverhaeltn|berufsstaend|haerteeinwand|vorlaeuf|"
    r"demgegenueber|abzuaender|zuegig|buehl|bueroleit|sachstandsueber|abloes|"
    r"schriftsaetze|pflichtteilsanspruech|vermaechtn|ganztaeg|roeschel|"
    r"wuerttemberg|wamsdoerfer|brueckl|fachanwaelt|versorgungstraeg|muend|"
    r"beschwerdefuehr|bevollmaechtig|sachverstaend|ertraeg|groess|oeffent|"
    r"ursprueng|aender|traeg|saeul|faell|angefuehr|umstaend|schmaeler|"
    r"hoechst|auszueg|universitaet|buchfuehr|pietaet|schriftsaetz|gewoehn|"
    r"versoehn|ausgefuehr|ausdrueck|erloes|gueter|fuenf|verkuend|lymphoedem|"
    r"ortsueblic|dreissig|wuerdig|"
    r"minoritaet|buerg|haelt|bewaehr|stueck|regulaer|stationaer|genueg|laeuf|"
    r"gruend|fuehr|beschraenk|verguet|abzueg|verjaehr|oedem|gaeng|daenemark|"
    r"migraen|eigenhaend|abhaeng|ausueb|enthaelt|erhoeh|stuetz|schuerf|"
    r"gekuerz|gedaechtn|unnoet|maerker|baeck|unterstuetz|oesterling|ausloes|"
    r"pruef|eroerter|foerder|gegenstaend|zuzueg|buero|fuell|saeumn|raeum|"
    r"roentgen|knoech|geroet|geaeusser|fortzufuehr|haehnel|guete|ausgehaeng|"
    r"beduerf|zaehl|bloed|mitzaehl|erhoeh|aufhoer|gross|fueg|fuehl|kaempf|"
    r"laeng|hoehn|instabilitaet|maerkisch|neuropaediatr|ermued|waerm|praejudiz|"
    r"pflegestuetz|ausgehaend|plaene|eilbeduerft|abwaeg|lueck|gepruef|knoepf|"
    r"frühstueck|unterkoerper|fuetter|knoechel|kraehen|naesse|rollstuehle|"
    r"fluesse|beschluesse|erhaelt|ausschuett|koepfig|zufuehr|glaeub|fahrlaess|"
    r"abfuehr|uebrig|betrueg|verzoeg|europae|eigenmaechtig|protokollfuehr|"
    r"rückfuehr|dreiwoech|einwaend|behaelt|gehoer|sekundaer|maessig|transossaer|"
    r"anaesthes|interskalenaer|paeckchen|behelfsmaessig|woertlich|planmaessig|"
    r"praeoperativ|hoehe|staette|minuetig|auszuueb|ploetz|engpaess|turnusmaess|"
    r"kausalitaet|anlaess|erwaehn|ursaech|extremitaet|venoese|sphaere|hoeflinger|"
    r"loeb|moeller|moerser|oehlert|anspruech|schlaeg|aufzuklaer|foerdewerft|"
    r"roestwerk|schultergueftels|"
    r"abschliess|ausschliess|ausser|massgeb|aufmass|fussgaeng|reissverschluss|"
    r"verstoss|aussertarif)|strasse)",
    re.IGNORECASE,
)


def prose_without_technical_names(text: str, path: Path) -> str:
    prose = PROTECTED_PROSE.sub("", text)
    return re.sub(re.escape(path.stem), "", prose, flags=re.IGNORECASE)


def docx_text(path: Path) -> str:
    document = Document(path)
    paragraphs = [paragraph.text for paragraph in document.paragraphs]
    cells = [
        cell.text
        for table in document.tables
        for row in table.rows
        for cell in row.cells
    ]
    headers_and_footers = [
        paragraph.text
        for section in document.sections
        for area in (section.header, section.footer)
        for paragraph in area.paragraphs
    ]
    return "\n".join(paragraphs + cells + headers_and_footers)


def eml_text(path: Path) -> str:
    message = eml_message(path)
    part = message.get_body(preferencelist=("plain",))
    return part.get_content() if part is not None else str(message)


def eml_message(path: Path):
    return BytesParser(policy=policy.default).parsebytes(path.read_bytes())


def pdf_text(path: Path) -> str:
    reader = PdfReader(path)
    if not reader.pages:
        raise ValueError("PDF enthält keine Seite")
    return "\n".join(page.extract_text() or "" for page in reader.pages)


def pdf_is_a4(path: Path) -> bool:
    for page in PdfReader(path).pages:
        width = float(page.mediabox.width)
        height = float(page.mediabox.height)
        portrait = abs(width - 595.28) < 8 and abs(height - 841.89) < 8
        landscape = abs(width - 841.89) < 8 and abs(height - 595.28) < 8
        if not (portrait or landscape):
            return False
    return True


def export_text(path: Path) -> str:
    suffix = path.suffix.lower()
    if suffix == ".docx":
        return docx_text(path)
    if suffix in {".xlsx", ".xlsm"}:
        workbook = load_workbook(path, read_only=True, data_only=False)
        try:
            return "\n".join(
                str(value)
                for sheet in workbook.worksheets
                for row in sheet.iter_rows(values_only=True)
                for value in row
                if value is not None
            )
        finally:
            workbook.close()
    if suffix == ".pdf":
        return pdf_text(path)
    if suffix == ".eml":
        return eml_text(path)
    return path.read_text(encoding="utf-8", errors="ignore")


def is_a4(document: Document) -> bool:
    for section in document.sections:
        width = section.page_width.cm
        height = section.page_height.cm
        portrait = abs(width - 21.0) < 0.2 and abs(height - 29.7) < 0.2
        landscape = abs(width - 29.7) < 0.2 and abs(height - 21.0) < 0.2
        if not (portrait or landscape):
            return False
    return True


def eml_quality_errors(path: Path) -> list[str]:
    """Prüft eine gespeicherte E-Mail auf portable, unverfälschte Darstellung."""
    label = path.relative_to(REPO) if path.is_relative_to(REPO) else path
    errors: list[str] = []
    raw = path.read_bytes()
    try:
        decoded_raw = raw.decode("utf-8")
    except UnicodeDecodeError as exc:
        return [f"{label}: E-Mail ist nicht valides UTF-8: {exc}"]
    try:
        message = BytesParser(policy=policy.default).parsebytes(raw)
    except Exception as exc:
        return [f"{label}: E-Mail nicht lesbar: {exc}"]
    for header in ("From", "To", "Date", "Subject", "Message-ID"):
        if not message.get(header):
            errors.append(f"{label}: E-Mail-Header {header} fehlt")
    if SYNTHETIC_EMAIL_PATTERN.search(decoded_raw):
        errors.append(f"{label}: künstliche E-Mail-Domain vorhanden")
    if re.search(r"[^\x00-\x7f]", decoded_raw):
        if not message.get("MIME-Version"):
            errors.append(f"{label}: MIME-Version für Nicht-ASCII-Inhalt fehlt")
        text_parts = [
            part for part in message.walk() if part.get_content_maintype() == "text"
        ]
        if not text_parts or any(
            (part.get_content_charset() or "").lower().replace("_", "-") != "utf-8"
            for part in text_parts
        ):
            errors.append(
                f"{label}: UTF-8-Zeichensatz ist nicht für alle Textteile deklariert"
            )
    body = message.get_body(preferencelist=("plain", "html"))
    rendered = body.get_content() if body is not None else str(message)
    if any(marker in rendered for marker in BROKEN_ENCODING_MARKERS):
        errors.append(f"{label}: E-Mail-Inhalt wird fehlerhaft dekodiert")
    return errors


def main() -> int:
    errors: list[str] = []
    checked_files = 0
    cases = [
        path
        for path in sorted(TESTAKTEN.iterdir())
        if path.is_dir() and path.name.startswith(PREFIXES)
    ]

    for path in sorted(TESTAKTEN.rglob("*")):
        if not path.is_file() or "gesamt-pdf" in path.parts:
            continue
        if path.name.lower() in META_EXACT_NAMES:
            errors.append(
                f"{path.relative_to(REPO)}: überholtes Standard-Aktenstück vorhanden"
            )
        if path.suffix.lower() != ".eml":
            continue
        errors.extend(eml_quality_errors(path))

    export_files_checked = 0
    for case in sorted(path for path in TESTAKTEN.iterdir() if path.is_dir()):
        for path in sorted(case.rglob("*")):
            if not include_in_working_dump(path, case):
                continue
            export_files_checked += 1
            if path.suffix.lower() not in EXPORT_TEXT_EXTS:
                continue
            try:
                text = export_text(path)
            except Exception as exc:
                errors.append(f"{path.relative_to(REPO)}: Exportquelle nicht lesbar: {exc}")
                continue
            if SYNTHETIC_EMAIL_PATTERN.search(text):
                errors.append(
                    f"{path.relative_to(REPO)}: künstliche E-Mail-Domain vorhanden"
                )
            for label, pattern in EXPORT_META_PATTERNS.items():
                match = pattern.search(text)
                if match:
                    errors.append(
                        f"{path.relative_to(REPO)}: {label} {match.group(0)!r}"
                    )

    for relative in sorted(REMOVED_AGGREGATES):
        old_path = TESTAKTEN / relative
        case_readme = TESTAKTEN / Path(relative).parts[0] / "README.md"
        if old_path.exists():
            errors.append(f"{old_path.relative_to(REPO)}: zusammengezogenes Altstück vorhanden")
        if case_readme.exists() and Path(relative).name in case_readme.read_text(encoding="utf-8"):
            errors.append(
                f"{case_readme.relative_to(REPO)}: verweist noch auf {Path(relative).name}"
            )

    for case in cases:
        hashes: dict[str, list[Path]] = defaultdict(list)
        for path in sorted(case.rglob("*")):
            if not path.is_file() or "gesamt-pdf" in path.parts:
                continue
            hashes[hashlib.sha256(path.read_bytes()).hexdigest()].append(path)
            suffix = path.suffix.lower()
            if suffix == ".txt":
                text = path.read_text(encoding="utf-8", errors="strict")
                if len(text.strip()) < MIN_FORMAL_TEXT:
                    errors.append(
                        f"{path.relative_to(REPO)}: nur {len(text.strip())} Textzeichen"
                    )
                if "§" in text:
                    errors.append(f"{path.relative_to(REPO)}: Paragrafenzeichen vorhanden")
                if re.search(r"\bParagraph(?:en|e|s)?\b", text, re.IGNORECASE):
                    errors.append(f"{path.relative_to(REPO)}: 'Paragraf' nicht ausgeschrieben")
                prose = prose_without_technical_names(text, path)
                match = TRANSLITERATION.search(prose)
                if match:
                    errors.append(
                        f"{path.relative_to(REPO)}: unechter Umlaut {match.group(0)!r}"
                    )
                continue
            if suffix == ".csv":
                rows = list(csv.reader(path.open(encoding="utf-8-sig")))
                nonempty = [row for row in rows if any(cell.strip() for cell in row)]
                if len(nonempty) < 4:
                    errors.append(
                        f"{path.relative_to(REPO)}: nur {len(nonempty)} befüllte CSV-Zeilen"
                    )
                text = "\n".join(" ".join(row) for row in nonempty)
                if "§" in text:
                    errors.append(f"{path.relative_to(REPO)}: Paragrafenzeichen vorhanden")
                if re.search(r"\bParagraph(?:en|e|s)?\b", text, re.IGNORECASE):
                    errors.append(f"{path.relative_to(REPO)}: 'Paragraf' nicht ausgeschrieben")
                prose = prose_without_technical_names(text, path)
                match = TRANSLITERATION.search(prose)
                if match:
                    errors.append(
                        f"{path.relative_to(REPO)}: unechter Umlaut {match.group(0)!r}"
                    )
                continue
            if suffix == ".xlsx":
                workbook = load_workbook(path, read_only=True, data_only=False)
                values = [
                    str(cell.value)
                    for sheet in workbook.worksheets
                    for row in sheet.iter_rows()
                    for cell in row
                    if cell.value not in (None, "")
                ]
                filled_rows = sum(
                    1
                    for sheet in workbook.worksheets
                    for row in sheet.iter_rows()
                    if any(cell.value not in (None, "") for cell in row)
                )
                if filled_rows < 4:
                    errors.append(
                        f"{path.relative_to(REPO)}: nur {filled_rows} befüllte Tabellenzeilen"
                    )
                text = "\n".join(values)
                if "§" in text:
                    errors.append(f"{path.relative_to(REPO)}: Paragrafenzeichen vorhanden")
                if re.search(r"\bParagraph(?:en|e|s)?\b", text, re.IGNORECASE):
                    errors.append(f"{path.relative_to(REPO)}: 'Paragraf' nicht ausgeschrieben")
                prose = prose_without_technical_names(text, path)
                match = TRANSLITERATION.search(prose)
                if match:
                    errors.append(
                        f"{path.relative_to(REPO)}: unechter Umlaut {match.group(0)!r}"
                    )
                continue
            if path.suffix.lower() not in FORMAL_EXTS:
                continue
            checked_files += 1
            try:
                if path.suffix.lower() == ".docx":
                    document = Document(path)
                    text = docx_text(path)
                    if not is_a4(document):
                        errors.append(f"{path.relative_to(REPO)}: kein A4-Format")
                elif path.suffix.lower() == ".eml":
                    message = eml_message(path)
                    text = eml_text(path)
                else:
                    text = pdf_text(path)
                    if not pdf_is_a4(path):
                        errors.append(f"{path.relative_to(REPO)}: kein A4-Format")
            except Exception as exc:
                errors.append(f"{path.relative_to(REPO)}: nicht lesbar: {exc}")
                continue

            cleaned = text.strip()
            if len(cleaned) < MIN_FORMAL_TEXT:
                errors.append(
                    f"{path.relative_to(REPO)}: nur {len(cleaned)} Textzeichen"
                )
            if "\ufffd" in text or any(marker in text for marker in ("Ã", "Â", "â€", "ðŸ")):
                errors.append(f"{path.relative_to(REPO)}: fehlerhafte Zeichenkodierung")
            if "§" in text:
                errors.append(f"{path.relative_to(REPO)}: Paragrafenzeichen vorhanden")
            if re.search(r"\bParagraph(?:en|e|s)?\b", text, re.IGNORECASE):
                errors.append(f"{path.relative_to(REPO)}: 'Paragraf' nicht ausgeschrieben")
            if not re.search(r"[äöüÄÖÜß]", text):
                errors.append(f"{path.relative_to(REPO)}: weder echter Umlaut noch scharfes S")
            lowered = text.lower()
            for marker in META_MARKERS:
                if marker in lowered:
                    errors.append(
                        f"{path.relative_to(REPO)}: verräterischer Metahinweis {marker!r}"
                    )
            prose = prose_without_technical_names(text, path)
            match = TRANSLITERATION.search(prose)
            if match:
                errors.append(
                    f"{path.relative_to(REPO)}: unechter Umlaut {match.group(0)!r}"
                )

        for duplicates in hashes.values():
            if len(duplicates) > 1:
                joined = ", ".join(str(path.relative_to(case)) for path in duplicates)
                errors.append(f"{case.relative_to(REPO)}: identische Dubletten: {joined}")

    if errors:
        print("Dokumentqualitätsprüfung fehlgeschlagen:")
        for error in errors:
            print(f"- {error}")
        return 1

    print(
        f"Dokumentqualität OK: {len(cases)} Akten, "
        f"{checked_files} formale Dokumente, {export_files_checked} Exportdateien"
    )
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
