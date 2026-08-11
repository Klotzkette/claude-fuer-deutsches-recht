#!/usr/bin/env python3
"""Regressionstests fuer strenge und transaktionale Akten-PDF-Erzeugung."""

from __future__ import annotations

import importlib.util
import hashlib
import io
import sys
import tempfile
from pathlib import Path
import zipfile

from odf.opendocument import OpenDocumentText
from odf.text import H, P
from docx import Document

from testakte_disclaimer import (
    NOTICE_BYTES,
    NOTICE_DE,
    NOTICE_EN,
    NOTICE_FILENAME,
    pdf_notice_errors,
)


SCRIPTS = Path(__file__).resolve().parent


def load_module(filename: str, module_name: str):
    spec = importlib.util.spec_from_file_location(module_name, SCRIPTS / filename)
    assert spec and spec.loader
    module = importlib.util.module_from_spec(spec)
    sys.modules[module_name] = module
    spec.loader.exec_module(module)
    return module


G = load_module("build-testakte-gesamt-pdf.py", "test_build_testakte_gesamt_pdf")
E = load_module("build-testakten-einzelpdf-zips.py", "test_build_testakten_einzelpdf_zips")
W = load_module("build-testakten-release-zips.py", "test_build_testakten_release_zips")


def require(condition: bool, message: str) -> None:
    if not condition:
        raise AssertionError(message)


def main() -> int:
    letter_page = G.PageObject.create_blank_page(width=612, height=792)
    normalized_letter = G.a4_normalized_page(letter_page)
    require(
        abs(float(normalized_letter.mediabox.width) - float(G.A4[0])) < 0.5
        and abs(float(normalized_letter.mediabox.height) - float(G.A4[1])) < 0.5,
        "fremde Hochformat-Seiten muessen im Gesamt-PDF auf A4 normalisiert werden",
    )
    landscape_letter = G.PageObject.create_blank_page(width=792, height=612)
    normalized_landscape = G.a4_normalized_page(landscape_letter)
    require(
        abs(float(normalized_landscape.mediabox.width) - float(G.A4[1])) < 0.5
        and abs(float(normalized_landscape.mediabox.height) - float(G.A4[0])) < 0.5,
        "fremde Querformat-Seiten muessen im Gesamt-PDF auf A4 quer normalisiert werden",
    )

    letter_writer = G.PdfWriter()
    letter_writer.add_blank_page(width=612, height=792)
    letter_bytes = io.BytesIO()
    letter_writer.write(letter_bytes)
    normalized_bytes = E.normalize_pdf_to_a4(letter_bytes.getvalue(), "letter.pdf")
    normalized_page = G.PdfReader(io.BytesIO(normalized_bytes)).pages[0]
    require(
        abs(float(normalized_page.mediabox.width) - float(G.A4[0])) < 0.5
        and abs(float(normalized_page.mediabox.height) - float(G.A4[1])) < 0.5,
        "Einzel-PDFs muessen fremde Hochformat-Seiten auf A4 normalisieren",
    )

    with tempfile.TemporaryDirectory(prefix="testakte-pdf-") as tmp:
        root = Path(tmp)
        odt = root / "akte.odt"
        doc = OpenDocumentText()
        doc.text.addElement(H(outlinelevel=1, text="Sachstand"))
        doc.text.addElement(P(text="Der Antrag liegt vollstaendig vor."))
        doc.save(str(odt))

        flowables = E.odt_to_flowables(odt)
        require(len(flowables) == 2, "ODT-Ueberschrift und Absatz muessen gerendert werden")

        broken = root / "defekt.docx"
        broken.write_bytes(b"kein Office-Dokument")
        try:
            E.render_document_pdf(broken, root)
        except E.G.DocumentRenderError:
            pass
        else:
            raise AssertionError("Defekte DOCX-Datei darf kein Platzhalter-PDF erzeugen")
        broken.unlink()

        styled = root / "brief.docx"
        styled_doc = Document()
        styled_doc.sections[0].header.paragraphs[0].text = "Kanzlei am Markt - Aktenzeichen 26/184"
        styled_doc.add_paragraph("Vollständiger Briefkopf mit Straße und Rückrufnummer.")
        styled_doc.add_page_break()
        styled_doc.add_paragraph("Zweite Seite mit gesonderter Anlage.")
        styled_doc.sections[0].footer.paragraphs[0].text = "Seite und Vertraulichkeitsvermerk"
        styled_doc.save(styled)
        styled_pdf = E.render_document_pdf(styled, root)
        require(styled_pdf is not None, "DOCX muss ein PDF ergeben")
        styled_reader = G.PdfReader(io.BytesIO(styled_pdf))
        if E.render_office(styled) is not None:
            require(
                len(styled_reader.pages) == 3,
                "native DOCX-PDF muss Hinweis und Seitenumbrüche erhalten",
            )
            styled_text = "\n".join(page.extract_text() or "" for page in styled_reader.pages)
            require("Kanzlei am Markt" in styled_text, "native DOCX-PDF muss Kopfzeile erhalten")
            require("Vertraulichkeitsvermerk" in styled_text, "native DOCX-PDF muss Fußzeile erhalten")
        require(
            not pdf_notice_errors(styled_pdf, exactly_once=True),
            "jedes Einzel-PDF braucht den zweisprachigen Hinweis genau einmal",
        )

        pdf_data = E.render_document_pdf(odt, root)
        require(pdf_data is not None and pdf_data.startswith(b"%PDF-"), "ODT muss ein PDF ergeben")
        require(len(list(G.PdfReader(io.BytesIO(pdf_data)).pages)) >= 1, "PDF braucht mindestens eine Seite")
        require(
            not pdf_notice_errors(pdf_data, exactly_once=True),
            "gerenderte Einzel-PDFs muessen den Hinweis tragen",
        )

        raw_json = root / "messwerte.json"
        raw_json.write_text(
            '{"geraet": "Prüfstand Süd", "werte": [12, 17], "freigabe": false}\n',
            encoding="utf-8",
        )
        calendar = root / "frist.ics"
        calendar.write_text(
            "BEGIN:VCALENDAR\nVERSION:2.0\nBEGIN:VEVENT\n"
            "DTSTART:20260721T083000\nSUMMARY:Fristablauf Widerspruch\n"
            "END:VEVENT\nEND:VCALENDAR\n",
            encoding="utf-8",
        )
        html_mail = root / "eingang-html.eml"
        html_mail.write_text(
            "From: poststelle@example.invalid\nTo: kanzlei@example.invalid\n"
            "Date: Wed, 15 Jul 2026 11:24:00 +0200\nSubject: Bescheid und Straße\n"
            "MIME-Version: 1.0\nContent-Type: text/html; charset=utf-8\n\n"
            "<html><head><style>p{color:red}</style></head><body>"
            "<h1>Bescheid</h1><p>Die Behörde übersendet den Bescheid für die Südstraße.</p>"
            "</body></html>\n",
            encoding="utf-8",
        )
        json_pdf = E.render_document_pdf(raw_json, root)
        ics_pdf = E.render_document_pdf(calendar, root)
        html_mail_pdf = E.render_document_pdf(html_mail, root)
        require(json_pdf is not None and json_pdf.startswith(b"%PDF-"), "JSON muss ein PDF ergeben")
        require(ics_pdf is not None and ics_pdf.startswith(b"%PDF-"), "ICS muss ein PDF ergeben")
        require(html_mail_pdf is not None, "HTML-E-Mail muss ein PDF ergeben")
        html_mail_text = "\n".join(
            page.extract_text() or "" for page in G.PdfReader(io.BytesIO(html_mail_pdf)).pages
        )
        require("Südstraße" in html_mail_text, "HTML-E-Mail muss sichtbaren Text statt Quellcode rendern")
        require("<style>" not in html_mail_text, "HTML-Styles dürfen nicht im E-Mail-PDF erscheinen")

        dist = root / "dist"
        dist.mkdir()
        archive, count = E.build_single(root, dist)
        require(count == 5, "Office-, E-Mail- und strukturierte Dateien müssen jeweils als Einzel-PDF im ZIP landen")
        first_hash = hashlib.sha256(archive.read_bytes()).hexdigest()
        archive, _ = E.build_single(root, dist)
        second_hash = hashlib.sha256(archive.read_bytes()).hexdigest()
        require(first_hash == second_hash, "wiederholte ZIP-Bauten muessen byteidentisch sein")
        with zipfile.ZipFile(archive) as built:
            require(
                all(info.date_time == E.ZIP_TIMESTAMP for info in built.infolist()),
                "ZIP-Eintraege brauchen stabile Zeitstempel",
            )
            require(
                all("/" not in info.filename for info in built.infolist()),
                "Einzel-PDF-ZIPs duerfen keine Unterordner enthalten",
            )
            require("messwerte.pdf" in built.namelist(), "JSON braucht eine eigene flache PDF")
            require("frist.pdf" in built.namelist(), "ICS braucht eine eigene flache PDF")
            for info in built.infolist():
                require(
                    not pdf_notice_errors(built.read(info), exactly_once=True),
                    f"{info.filename} braucht den Hinweis genau einmal",
                )

        working_case = root / "arbeitsakte"
        working_case.mkdir()
        (working_case / "01_sachstand.txt").write_text(
            "Antrag und Bescheid liegen vor.\n", encoding="utf-8"
        )
        nested = working_case / "eingang"
        nested.mkdir()
        (nested / "02_nachricht.eml").write_text(
            "From: mandant@example.invalid\nTo: kanzlei@example.invalid\n"
            "Date: Tue, 14 Jul 2026 10:20:00 +0200\n"
            "Subject: Unterlagen\n\nAnbei die Unterlagen.\n",
            encoding="utf-8",
        )
        (nested / "nicht_exportieren.md").write_text(
            "# Interne Vorschau\n",
            encoding="utf-8",
        )
        placeholder_docx = working_case / "03_unfertiger_entwurf.docx"
        placeholder = Document()
        placeholder.add_paragraph("Schreiben vom [Datum einsetzen]")
        placeholder.save(placeholder_docx)
        solution_docx = working_case / "04_interne_auswertung.docx"
        solution = Document()
        solution.add_paragraph("Musterlösung und Erwartungshorizont")
        solution.save(solution_docx)
        for filename in (
            "05_rechtsprechungsanalyse.docx",
            "06_klageraster_sozialgericht.docx",
            "07_plan_quality_gate.docx",
            "08_phase_i_pruefraster.docx",
            "09_rechtsgutachten.docx",
            "10_chronologie_arbeitsstand.docx",
            "11_prozessstrategie.docx",
            "12_kanzleinotizen_intern.docx",
            "13_gerichtliche_route.docx",
            "14_zieloutput_checkliste.docx",
            "15_schutzschirmantrag_vorbereitung.docx",
            "16_sachwalter_eigenverwaltung_notiz.docx",
        ):
            meta = Document()
            meta.add_paragraph("Interne fachliche Auswertung mit abschließendem Ergebnis.")
            meta.save(working_case / filename)
        working_dist = root / "working-dist"
        working_dist.mkdir()
        working_archive, working_count = W.build_single(working_case, working_dist)
        require(
            working_count == 3,
            "Arbeitsakten-ZIP muss Hinweis und beide nativen Unterlagen enthalten",
        )
        first_working_hash = hashlib.sha256(working_archive.read_bytes()).hexdigest()
        working_archive, _ = W.build_single(working_case, working_dist)
        second_working_hash = hashlib.sha256(working_archive.read_bytes()).hexdigest()
        require(
            first_working_hash == second_working_hash,
            "Arbeitsakten-ZIPs müssen bei gleichem Bestand byteidentisch sein",
        )
        with zipfile.ZipFile(working_archive) as built:
            require(
                all(info.date_time == W.ZIP_TIMESTAMP for info in built.infolist()),
                "Arbeitsakten-ZIPs brauchen stabile Zeitstempel",
            )
            names = built.namelist()
            require(all("/" not in name for name in names), "Arbeitsakten-ZIP muss flach sein")
            require(not any(name.endswith(".md") for name in names), "Markdown darf nicht ins Akten-ZIP")
            require(NOTICE_FILENAME in names, "Arbeitsakten-ZIP braucht README.txt")
            require(built.read(NOTICE_FILENAME) == NOTICE_BYTES, "README.txt braucht verbindlichen Wortlaut")
            notice_text = built.read(NOTICE_FILENAME).decode("utf-8")
            require(NOTICE_DE in notice_text, "deutscher Hinweis fehlt in README.txt")
            require(NOTICE_EN in notice_text, "englischer Hinweis fehlt in README.txt")
            require("eingang__02_nachricht.eml" in names, "Pfadbestandteile muessen im flachen Namen erhalten bleiben")
            require(
                not any("unfertiger_entwurf" in name for name in names),
                "Dokumente mit Datumsplatzhaltern duerfen nicht ins Akten-ZIP",
            )
            require(
                not any("interne_auswertung" in name for name in names),
                "Musterloesungen duerfen nicht ins Akten-ZIP",
            )
            require(
                not any(
                    marker in name
                    for name in names
                    for marker in (
                        "rechtsprechungsanalyse",
                        "klageraster",
                        "quality_gate",
                        "pruefraster",
                        "rechtsgutachten",
                        "chronologie_arbeitsstand",
                        "prozessstrategie",
                        "kanzleinotizen_intern",
                        "gerichtliche_route",
                        "zieloutput_checkliste",
                        "schutzschirmantrag_vorbereitung",
                        "sachwalter_eigenverwaltung_notiz",
                    )
                ),
                "interne Rechtsanalysen und Pruefraster duerfen nicht ins Akten-ZIP",
            )

    print("test-testakte-pdf-build OK")
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
