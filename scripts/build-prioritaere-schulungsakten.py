#!/usr/bin/env python3
"""Erzeugt fünf fachlich priorisierte Schulungsakten aus Primärdokumenten."""

from __future__ import annotations

import csv
import shutil
import sys
from email.message import EmailMessage
from email.policy import SMTP
from pathlib import Path

from PIL import Image, ImageDraw, ImageFont
from docx import Document
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor


REPO = Path(__file__).resolve().parent.parent
ROOT = REPO / "testakten"

NAVY = RGBColor(28, 55, 76)
TEAL = RGBColor(1, 105, 111)
GRAY = RGBColor(92, 98, 102)
LIGHT = "E8EEF1"


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_repeat_table_header(row) -> None:
    tr_pr = row._tr.get_or_add_trPr()
    tbl_header = OxmlElement("w:tblHeader")
    tbl_header.set(qn("w:val"), "true")
    tr_pr.append(tbl_header)


def add_page_number(paragraph) -> None:
    run = paragraph.add_run()
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    instruction = OxmlElement("w:instrText")
    instruction.set(qn("xml:space"), "preserve")
    instruction.text = " PAGE "
    separate = OxmlElement("w:fldChar")
    separate.set(qn("w:fldCharType"), "separate")
    text = OxmlElement("w:t")
    text.text = "1"
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    run._r.extend([begin, instruction, separate, text, end])


def base_document(org: str, strapline: str, contact: str, accent: RGBColor = NAVY) -> Document:
    doc = Document()
    section = doc.sections[0]
    section.page_width = Cm(21.0)
    section.page_height = Cm(29.7)
    section.top_margin = Cm(2.35)
    section.bottom_margin = Cm(2.2)
    section.left_margin = Cm(2.45)
    section.right_margin = Cm(2.25)
    section.header_distance = Cm(0.8)
    section.footer_distance = Cm(0.9)

    normal = doc.styles["Normal"]
    normal.font.name = "Times New Roman"
    normal.font.size = Pt(11)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.08
    for style_name, size in (("Heading 1", 14), ("Heading 2", 12), ("Heading 3", 11)):
        style = doc.styles[style_name]
        style.font.name = "Times New Roman"
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = accent
        style.paragraph_format.space_before = Pt(10)
        style.paragraph_format.space_after = Pt(4)
        style.paragraph_format.keep_with_next = True

    header = section.header
    table = header.add_table(rows=1, cols=2, width=Cm(16.3))
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    table.autofit = False
    table.columns[0].width = Cm(9.7)
    table.columns[1].width = Cm(6.6)
    left = table.cell(0, 0)
    right = table.cell(0, 1)
    left.width = Cm(9.7)
    right.width = Cm(6.6)
    p = left.paragraphs[0]
    p.paragraph_format.space_after = Pt(0)
    r = p.add_run(org)
    r.bold = True
    r.font.name = "Arial"
    r.font.size = Pt(12.5)
    r.font.color.rgb = accent
    r2 = p.add_run("\n" + strapline)
    r2.font.name = "Arial"
    r2.font.size = Pt(8.5)
    r2.font.color.rgb = GRAY
    p = right.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    p.paragraph_format.space_after = Pt(0)
    r = p.add_run(contact)
    r.font.name = "Arial"
    r.font.size = Pt(7.5)
    r.font.color.rgb = GRAY
    bottom = OxmlElement("w:pBdr")
    border = OxmlElement("w:bottom")
    border.set(qn("w:val"), "single")
    border.set(qn("w:sz"), "8")
    border.set(qn("w:space"), "4")
    border.set(qn("w:color"), f"{accent[0]:02X}{accent[1]:02X}{accent[2]:02X}")
    bottom.append(border)
    header.paragraphs[-1]._p.get_or_add_pPr().append(bottom)

    footer = section.footer
    p = footer.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(f"{org}  |  {strapline}  |  Seite ")
    r.font.name = "Arial"
    r.font.size = Pt(7.5)
    r.font.color.rgb = GRAY
    add_page_number(p)
    return doc


def add_metadata(doc: Document, rows: list[tuple[str, str]]) -> None:
    table = doc.add_table(rows=len(rows), cols=2)
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    table.autofit = False
    for i, (label, value) in enumerate(rows):
        c0, c1 = table.rows[i].cells
        c0.width = Cm(3.4)
        c1.width = Cm(12.9)
        c0.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.TOP
        c1.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.TOP
        p0 = c0.paragraphs[0]
        p0.paragraph_format.space_after = Pt(1)
        r0 = p0.add_run(label)
        r0.bold = True
        r0.font.name = "Arial"
        r0.font.size = Pt(8.5)
        r0.font.color.rgb = GRAY
        p1 = c1.paragraphs[0]
        p1.paragraph_format.space_after = Pt(1)
        r1 = p1.add_run(value)
        r1.font.name = "Arial"
        r1.font.size = Pt(9)
    doc.add_paragraph().paragraph_format.space_after = Pt(2)


def add_table(doc: Document, headers: list[str], rows: list[list[str]], widths: list[float] | None = None) -> None:
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    table.autofit = False
    set_repeat_table_header(table.rows[0])
    for j, header in enumerate(headers):
        cell = table.rows[0].cells[j]
        set_cell_shading(cell, LIGHT)
        p = cell.paragraphs[0]
        p.paragraph_format.space_after = Pt(1)
        run = p.add_run(header)
        run.bold = True
        run.font.name = "Arial"
        run.font.size = Pt(8.5)
    for row in rows:
        cells = table.add_row().cells
        for j, value in enumerate(row):
            cells[j].vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.TOP
            p = cells[j].paragraphs[0]
            p.paragraph_format.space_after = Pt(1)
            run = p.add_run(value)
            run.font.name = "Times New Roman"
            run.font.size = Pt(9)
    if widths:
        for row in table.rows:
            for j, width in enumerate(widths):
                row.cells[j].width = Cm(width)
    doc.add_paragraph().paragraph_format.space_after = Pt(1)


def add_sections(doc: Document, sections: list[tuple[str, list[str]]]) -> None:
    for heading, paragraphs in sections:
        if heading:
            doc.add_heading(heading, level=1)
        for text in paragraphs:
            p = doc.add_paragraph(text)
            p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            p.paragraph_format.widow_control = True


def write_letter(
    path: Path,
    *,
    org: str,
    strapline: str,
    contact: str,
    recipient: list[str],
    date: str,
    reference: str,
    subject: str,
    salutation: str,
    sections: list[tuple[str, list[str]]],
    closing: str,
    signer: str,
    signer_role: str,
    attachments: list[str] | None = None,
    accent: RGBColor = NAVY,
) -> None:
    doc = base_document(org, strapline, contact, accent)
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(12)
    r = p.add_run("\n".join(recipient))
    r.font.name = "Times New Roman"
    r.font.size = Pt(10.5)
    add_metadata(doc, [("Datum", date), ("Zeichen", reference)])
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(12)
    r = p.add_run(subject)
    r.bold = True
    r.font.size = Pt(11.5)
    doc.add_paragraph(salutation)
    add_sections(doc, sections)
    doc.add_paragraph(closing)
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(0)
    r = p.add_run(signer)
    r.bold = True
    p2 = doc.add_paragraph(signer_role)
    p2.paragraph_format.space_after = Pt(8)
    if attachments:
        p = doc.add_paragraph()
        r = p.add_run("Anlagen")
        r.bold = True
        for index, item in enumerate(attachments, start=1):
            doc.add_paragraph(f"{index}. {item}")
    path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(path)


def write_record(
    path: Path,
    *,
    org: str,
    strapline: str,
    contact: str,
    title: str,
    metadata: list[tuple[str, str]],
    sections: list[tuple[str, list[str]]],
    tables: list[tuple[str, list[str], list[list[str]], list[float] | None]] | None = None,
    signer: str | None = None,
    signer_role: str | None = None,
    accent: RGBColor = TEAL,
) -> None:
    doc = base_document(org, strapline, contact, accent)
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(7)
    r = p.add_run(title)
    r.bold = True
    r.font.name = "Arial"
    r.font.size = Pt(17)
    r.font.color.rgb = accent
    add_metadata(doc, metadata)
    add_sections(doc, sections)
    if tables:
        for heading, headers, rows, widths in tables:
            doc.add_heading(heading, level=1)
            add_table(doc, headers, rows, widths)
    if signer:
        doc.add_paragraph()
        p = doc.add_paragraph(signer)
        p.runs[0].bold = True
        if signer_role:
            doc.add_paragraph(signer_role)
    path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(path)


def write_email(
    path: Path,
    *,
    sender: str,
    recipients: list[str],
    cc: list[str],
    date: str,
    subject: str,
    message_id: str,
    body: str,
    attachments: list[str] | None = None,
) -> None:
    message = EmailMessage()
    message["From"] = sender
    message["To"] = ", ".join(recipients)
    if cc:
        message["Cc"] = ", ".join(cc)
    message["Date"] = date
    message["Subject"] = subject
    message["Message-ID"] = message_id
    if attachments:
        message["X-Attachments"] = ", ".join(attachments)
    message.set_content(body.rstrip() + "\n", charset="utf-8")
    path.parent.mkdir(parents=True, exist_ok=True)
    path.write_bytes(message.as_bytes(policy=SMTP))


def write_csv(path: Path, headers: list[str], rows: list[list[object]]) -> None:
    path.parent.mkdir(parents=True, exist_ok=True)
    with path.open("w", encoding="utf-8", newline="") as stream:
        writer = csv.writer(stream, delimiter=";", quoting=csv.QUOTE_MINIMAL)
        writer.writerow(headers)
        writer.writerows(rows)


def write_txt(path: Path, text: str) -> None:
    path.parent.mkdir(parents=True, exist_ok=True)
    path.write_text(text.rstrip() + "\n", encoding="utf-8")


def font(size: int, bold: bool = False):
    candidates = [
        "/System/Library/Fonts/Supplemental/Arial Bold.ttf" if bold else "/System/Library/Fonts/Supplemental/Arial.ttf",
        "/System/Library/Fonts/Helvetica.ttc",
    ]
    for candidate in candidates:
        try:
            return ImageFont.truetype(candidate, size)
        except OSError:
            continue
    return ImageFont.load_default()


def write_device_photo(path: Path, title: str, lines: list[str], accent: tuple[int, int, int]) -> None:
    image = Image.new("RGB", (1600, 1050), (226, 222, 210))
    draw = ImageDraw.Draw(image)
    draw.rounded_rectangle((170, 115, 1430, 930), radius=30, fill=(245, 245, 242), outline=(88, 88, 84), width=5)
    draw.rectangle((170, 115, 1430, 205), fill=accent)
    draw.text((220, 140), title, font=font(34, True), fill="white")
    y = 265
    for line in lines:
        draw.text((250, y), line, font=font(30, False), fill=(28, 32, 34))
        y += 78
    draw.text((220, 875), "Foto aus dem Vorgangsordner · aufgenommen mit Mobiltelefon", font=font(22), fill=(90, 90, 86))
    path.parent.mkdir(parents=True, exist_ok=True)
    image.save(path, quality=92, optimize=True)


def prepare(slug: str) -> Path:
    target = ROOT / slug
    if target.exists():
        if "--force" not in sys.argv:
            raise SystemExit(f"Ordner besteht bereits: {target}")
        shutil.rmtree(target)
    target.mkdir(parents=True)
    return target


def build_apothekenrecht() -> None:
    slug = "apothekenrecht-retaxation-kuehlkette-bremen"
    d = prepare(slug)
    pharmacy_contact = "Am Wall 147 · 28195 Bremen\nTelefon 0421 438 71 20 · Fax 0421 438 71 29\nservice@adler-apotheke-wall.de · IK 302 041 889"
    insurer_contact = "Konsul-Smidt-Straße 54 · 28217 Bremen\nTelefon 0421 590 88 0 · Fax 0421 590 88 199\narzneimittel@hansemerk-bkk.de"

    write_record(
        d / "01_e_rezept_abgabeprotokoll.docx",
        org="Adler-Apotheke am Wall",
        strapline="Dr. Helena Kröger e. K. · Öffentliche Apotheke",
        contact=pharmacy_contact,
        title="Abgabeprotokoll Hochpreistherapie",
        metadata=[("Vorgang", "HM-26-0218-7741"), ("Versicherte", "Yasemin Elmas, geb. 12.09.1978"), ("Abgabe", "21.02.2026, 10:42 Uhr"), ("Bearbeitung", "PTA Nora Ahlers / Freigabe Dr. Kröger")],
        sections=[
            ("1. Verordnung und Warenbezug", ["Das elektronische Rezept der Gemeinschaftspraxis Dr. Wiemer und Dr. Barten wurde am 18.02.2026 um 16:08 Uhr aus der Telematikinfrastruktur abgerufen. Verordnet waren zwei Packungen Selvarimab 120 mg, jeweils eine Fertigspritze, Anwendung nach schriftlichem Schema der rheumatologischen Praxis. Das Aut-idem-Feld war gesperrt. Die Warenwirtschaft zeigte keinen Rabattpartner und meldete eine Beschaffung ausschließlich über den pharmazeutischen Großhandel Nordpharm Logistik GmbH.", "Die Bestellung wurde am selben Tag um 16:24 Uhr als Kühlware ausgelöst. Nordpharm bestätigte die Auslieferung für den 20.02.2026 zwischen 07:00 und 08:00 Uhr. Der Einkaufspreis je Packung betrug 28740.50 EUR. Die Ware war auf den Namen der Versicherten reserviert und konnte nach Öffnung der Transportbox nicht zurückgegeben werden."]),
            ("2. Eingang und Freigabe", ["Fahrer Malte Rehberg übergab die Transportbox am 20.02.2026 um 07:31 Uhr an Apothekerin Jelena Markowitsch. Der Einwegindikator war grün. Der elektronische Logger zeigte beim Auslesen einen Höchstwert von 9.1 Grad Celsius für die Zeit von 07:18 bis 07:32 Uhr. Der zweite, in der Box liegende Referenzlogger wies maximal 7.8 Grad Celsius aus. Frau Markowitsch vermerkte telefonisch um 07:44 Uhr bei Nordpharm die Abweichung. Um 09:16 Uhr ging die schriftliche Freigabe des Großhandels ein.", "Dr. Kröger prüfte Chargen, Verfalldaten, Verordnung, Patientenstammdaten und das Dosierschema. Die Versicherte erhielt beide Packungen in einer Kühltasche mit zwei Akkus. Die Beratung umfasste Lagerung zwischen 2 und 8 Grad Celsius, Transport, Injektionszeitpunkte am 22.02. und 07.03.2026 sowie das Vorgehen bei Fieber. Frau Elmas bestätigte die Übernahme um 10:47 Uhr auf dem Abholbeleg."]),
            ("3. Abrechnung", ["Die Abrechnung wurde mit dem Datensatz 260221-1042 an das Rechenzentrum Nordwest übermittelt. Bruttobetrag einschließlich Zuschlägen und gesetzlicher Abschläge: 59842.18 EUR. Das Rezeptbild, die elektronische Dispensierinformation und die Chargendokumentation wurden im Vorgang HM-26-0218-7741 abgelegt. Eine zusätzliche ärztliche Papierbescheinigung lag bei der Abgabe nicht vor."]),
        ],
        signer="Dr. Helena Kröger",
        signer_role="Apothekerin und Inhaberin",
    )
    write_record(
        d / "02_grosshandel_lieferschein_freigabe.docx",
        org="Nordpharm Logistik GmbH",
        strapline="Pharmagroßhandel · Kühlzentrum Bremen-Hemelingen",
        contact="Europaallee 22 · 28309 Bremen\nTelefon 0421 699 42 0 · kühlservice@nordpharm-logistik.de\nHRB 29184 HB · USt-IdNr. DE291884302",
        title="Lieferschein und Qualitätsfreigabe",
        metadata=[("Lieferschein", "NP-260220-88173"), ("Kunde", "Adler-Apotheke am Wall, IK 302 041 889"), ("Tour", "HB-K2 / Fahrzeug HB-NP 418"), ("Übergabe", "20.02.2026, 07:31 Uhr")],
        sections=[
            ("1. Sendung", ["Transportbox K-99218 enthielt zwei Packungen Selvarimab 120 mg Fertigspritze, Artikel 17820411, Charge SVB26A17, Verfall 31.01.2028. Packmittel: validierte Mehrwegbox ThermoSafe 48, zwei vorkonditionierte Kühlelemente, Referenzlogger NL-44018 und Einwegindikator CoolMark 8. Die Box verließ das Kühlzentrum um 05:54 Uhr."]),
            ("2. Abweichungsmeldung", ["Die Apotheke meldete um 07:44 Uhr eine kurzzeitige Anzeige von 9.1 Grad Celsius am Fahrzeuglogger. Der Boxlogger NL-44018 blieb während des gesamten Transportes zwischen 4.3 und 7.8 Grad Celsius. Nach Prüfung durch die Qualitätssicherung wurde die Ware um 09:12 Uhr zur Abgabe freigegeben. Die Freigabe bezieht sich auf die Transportphase bis zur Übergabe und ersetzt nicht die Lagerungsprüfung der Apotheke."]),
            ("3. Abrechnung", ["Nettowarenwert 57481.00 EUR; Umsatzsteuer 10921.39 EUR; Rechnungsbetrag 68402.39 EUR. Zahlungsziel 14 Tage ohne Abzug. Eigentumsvorbehalt bis zum vollständigen Zahlungseingang. Die Retourensperre für patientenbezogene Kühlware wurde im Auftrag vermerkt."]),
        ],
        signer="Maren Voß",
        signer_role="Leitung Qualitätssicherung / Sachkundige Person",
    )
    write_letter(
        d / "03_retaxationsankuendigung_hansemerk_bkk.docx",
        org="HanseMerk BKK",
        strapline="Team Arzneimittelabrechnung Nord",
        contact=insurer_contact,
        recipient=["Adler-Apotheke am Wall", "z. Hd. Frau Dr. Helena Kröger", "Am Wall 147", "28195 Bremen"],
        date="18.06.2026",
        reference="AM-RX 7741/26 · IK 302 041 889",
        subject="Beanstandung der Abrechnung vom 21.02.2026 – Selvarimab 120 mg – 59842.18 EUR",
        salutation="Sehr geehrte Frau Dr. Kröger,",
        sections=[
            ("1. Beanstandete Abgabe", ["Für unsere Versicherte Yasemin Elmas wurden am 21.02.2026 zwei Packungen Selvarimab 120 mg abgerechnet. Der übermittelte Datensatz enthält als Dosierangabe lediglich den Hinweis auf ein schriftliches Schema. Eine eigenständige Dosierung ist aus den übermittelten Verordnungsdaten nicht ersichtlich. Außerdem liegt uns ein Hinweis des Rechenzentrums auf eine Temperaturabweichung während der Anlieferung vor."]),
            ("2. Vorgesehene Korrektur", ["Wir beabsichtigen, den Abrechnungsbetrag von 59842.18 EUR vollständig zu beanstanden und mit der nächsten Sammelabrechnung zu verrechnen. Vor einer endgültigen Buchung geben wir Ihnen Gelegenheit, bis zum 09.07.2026 die Verordnungsdaten, die dokumentierte pharmazeutische Prüfung, die Temperaturaufzeichnung beider Logger und eine Bestätigung der verordnenden Praxis vorzulegen."]),
            ("3. Unterlagenweg", ["Bitte übermitteln Sie keine Originale. Ordnen Sie jede Datei dem Zeichen AM-RX 7741/26 zu. Falls die Unterlagen nicht fristgerecht eingehen, entscheiden wir nach dem vorhandenen Datenbestand. Diese Nachricht trifft noch keine Aussage zur Erstattungsfähigkeit der Therapie im Verhältnis zur Versicherten."]),
        ],
        closing="Mit freundlichen Grüßen",
        signer="Imke Halberstadt",
        signer_role="Sachbearbeitung Arzneimittelabrechnung · Durchwahl 0421 590 88 231",
        attachments=["Abrechnungsauszug vom 10.06.2026", "Prüfdatensatz RX-7741"],
    )
    write_letter(
        d / "04_stellungnahme_rheumatologische_praxis.docx",
        org="Rheumatologie am Bürgerpark",
        strapline="Dr. med. Claas Wiemer · Dr. med. Antonia Barten",
        contact="Parkallee 92 · 28209 Bremen\nTelefon 0421 347 90 11 · Fax 0421 347 90 19\npraxis@rheuma-buergerpark.de · BSNR 038812700",
        recipient=["HanseMerk BKK", "Team Arzneimittelabrechnung Nord", "Konsul-Smidt-Straße 54", "28217 Bremen"],
        date="24.06.2026",
        reference="WE/ELM-120978",
        subject="Yasemin Elmas – Verordnung Selvarimab 120 mg vom 18.02.2026",
        salutation="Sehr geehrte Damen und Herren,",
        sections=[
            ("1. Verordnung", ["Wir bestätigen, dass wir am 18.02.2026 für Frau Yasemin Elmas zwei Packungen Selvarimab 120 mg als elektronische Verordnung ausgestellt haben. Die Verordnung war medizinisch beabsichtigt und sollte nicht ausgetauscht werden. Das Dosierschema wurde der Patientin als Ausdruck mitgegeben und am selben Tag in der Patientenakte hinterlegt."]),
            ("2. Dosierung", ["Vorgesehen war eine subkutane Injektion von 120 mg am 22.02.2026 und eine zweite Injektion von 120 mg am 07.03.2026. Die Behandlung schloss an die zuletzt am 08.02.2026 verabreichte Dosis an. Frau Elmas war seit Oktober 2025 auf diese Therapie eingestellt. Ein Abbruch oder eine Dosisänderung wurde nicht angeordnet."]),
            ("3. Rückfrage der Apotheke", ["Apothekerin Jelena Markowitsch rief am 20.02.2026 um 09:28 Uhr in unserer Praxis an und ließ sich Wirkstärke, Packungsanzahl und Anwendung bestätigen. Medizinische Fachangestellte Maren Segel vermerkte das Gespräch. Eine neue Verordnung wurde nicht erstellt, weil die bereits ausgestellte Verordnung inhaltlich zutraf."]),
        ],
        closing="Mit freundlichen Grüßen",
        signer="Dr. med. Claas Wiemer",
        signer_role="Facharzt für Innere Medizin und Rheumatologie",
    )
    write_letter(
        d / "05_einspruch_apotheke_retaxation.docx",
        org="Adler-Apotheke am Wall",
        strapline="Dr. Helena Kröger e. K. · Öffentliche Apotheke",
        contact=pharmacy_contact,
        recipient=["HanseMerk BKK", "Team Arzneimittelabrechnung Nord", "Konsul-Smidt-Straße 54", "28217 Bremen"],
        date="02.07.2026",
        reference="HK/NM 26-114 · Ihr Zeichen AM-RX 7741/26",
        subject="Unterlagen zur Abgabe Selvarimab 120 mg am 21.02.2026",
        salutation="Sehr geehrte Frau Halberstadt,",
        sections=[
            ("1. Verordnungsdaten", ["Zu Ihrem Schreiben vom 18.06.2026 übersenden wir das Abgabeprotokoll, die Bestätigung der verordnenden Praxis und den Ausdruck des Dosierschemas. Die Packungsanzahl entsprach den zwei vorgesehenen Injektionen. Die Versicherte war mit dem Arzneimittel bereits eingestellt. Eine Änderung der verordneten Therapie fand nicht statt."]),
            ("2. Kühlkette", ["Der von Ihnen erwähnte Wert von 9.1 Grad Celsius stammt aus dem Fahrzeuglogger, nicht aus dem validierten Logger der ungeöffneten Transportbox. Der Boxlogger blieb bis zur Übergabe innerhalb von 4.3 bis 7.8 Grad Celsius. Die Qualitätssicherung des Großhandels erteilte vor der Abgabe eine schriftliche Freigabe. In unserem Kühlschrank A-2 lagen die Werte zwischen 4.8 und 5.6 Grad Celsius."]),
            ("3. Wirtschaftliche Bedeutung", ["Die vollständige Verrechnung würde den monatlichen Rohertrag der Apotheke erheblich übersteigen, während die Ware ordnungsgemäß an die Versicherte gelangt ist. Wir bitten um Prüfung anhand der beigefügten Primärunterlagen und um Mitteilung, welche konkrete Angabe aus Ihrer Sicht trotz der ärztlichen Bestätigung noch fehlt."]),
        ],
        closing="Mit freundlichen Grüßen",
        signer="Dr. Helena Kröger",
        signer_role="Apothekerin und Inhaberin",
        attachments=["Abgabeprotokoll", "Lieferschein und Qualitätsfreigabe", "Temperaturdaten", "Ärztliche Bestätigung", "Abholbeleg"],
        accent=TEAL,
    )
    write_record(
        d / "06_abholbeleg_beratung.docx",
        org="Adler-Apotheke am Wall",
        strapline="Dr. Helena Kröger e. K. · Öffentliche Apotheke",
        contact=pharmacy_contact,
        title="Abhol- und Beratungsbeleg",
        metadata=[("Versicherte", "Yasemin Elmas, Versichertennummer HMBK-44721091"), ("Datum", "21.02.2026"), ("Uhrzeit", "10:42 bis 10:47 Uhr"), ("Kassenbon", "260221-0187")],
        sections=[
            ("1. Übergebene Arzneimittel", ["Zwei Packungen Selvarimab 120 mg Fertigspritze, Charge SVB26A17, Verfall 31.01.2028. Übergabe in isolierter Kühltasche mit zwei Kühlelementen. Die Packungen waren unbeschädigt und mit dem Namen der Versicherten gekennzeichnet."]),
            ("2. Besprochene Punkte", ["Lagerung im mittleren Kühlschrankfach zwischen 2 und 8 Grad Celsius; keine Lagerung an der Rückwand oder in der Tür; erste Injektion am 22.02.2026, zweite Injektion am 07.03.2026; Entnahme 20 Minuten vor Anwendung; Kontakt zur Praxis bei Fieber, Infektzeichen oder unterbrochener Kühlung. Frau Elmas erklärte, dass der Heimweg ungefähr zwölf Minuten dauere und sie unmittelbar nach Hause fahre."]),
            ("3. Bestätigung", ["Ich habe die genannten Packungen und das Dosierschema erhalten. Die Hinweise zur Lagerung und Anwendung wurden mir erklärt. Die Unterschrift befindet sich auf dem in der Apotheke verwahrten Papieroriginal; diese Datei ist die am 23.02.2026 erstellte Abschrift aus dem Warenwirtschaftsvorgang."]),
        ],
        signer="Nora Ahlers / Yasemin Elmas",
        signer_role="Abgebende PTA / Versicherte",
    )
    write_email(
        d / "07_email_nordpharm_qualitaetsfreigabe.eml",
        sender="Maren Voß <qualitaet@nordpharm-logistik.de>",
        recipients=["Jelena Markowitsch <j.markowitsch@adler-apotheke-wall.de>"],
        cc=["Dr. Helena Kröger <h.kroeger@adler-apotheke-wall.de>"],
        date="Fri, 20 Feb 2026 09:16:00 +0100",
        subject="Freigabe Box K-99218 / Lieferschein NP-260220-88173",
        message_id="<20260220.091600.K99218@nordpharm-logistik.de>",
        body="""Guten Morgen Frau Markowitsch,

wir haben die Rohdaten beider Geräte geprüft. Der Wert 9.1 Grad stammt vom Fahrzeuglogger an der Ladetür. Der Logger NL-44018 in der verschlossenen Transportbox lag vom Packen um 05:41 Uhr bis zur Übergabe um 07:31 Uhr zwischen 4.3 und 7.8 Grad Celsius. Der Einwegindikator hat nicht ausgelöst.

Die Qualitätssicherung gibt die beiden Packungen Selvarimab 120 mg, Charge SVB26A17, zur bestimmungsgemäßen Verwendung frei. Bitte legen Sie diese Nachricht zusammen mit dem CSV-Auszug in Ihrem Wareneingangsvorgang ab. Der Fahrer hat die handschriftliche Abweichungsnotiz auf seinem Tourenblatt ergänzt.

Freundliche Grüße
Maren Voß
Leitung Qualitätssicherung
Nordpharm Logistik GmbH
Telefon 0421 699 42 318""",
        attachments=["logger_NL-44018_20260220.csv"],
    )
    write_email(
        d / "08_email_rechenzentrum_verrechnung.eml",
        sender="Retax-Team <retax@nordwest-rezept.de>",
        recipients=["Buchhaltung <buchhaltung@adler-apotheke-wall.de>"],
        cc=["Dr. Helena Kröger <h.kroeger@adler-apotheke-wall.de>"],
        date="Tue, 23 Jun 2026 14:38:00 +0200",
        subject="Vormerkung Retax 59842.18 EUR / HanseMerk BKK / Juli-Abrechnung",
        message_id="<rx-7741-2606231438@nordwest-rezept.de>",
        body="""Guten Tag,

die HanseMerk BKK hat für die Verordnung vom 18.02.2026 eine Vollabsetzung angekündigt. Der Betrag von 59842.18 EUR ist im Abrechnungskonto zunächst nur vorgemerkt. Wenn bis zum 14.07.2026 keine Rücknahme oder Aussetzung eingeht, wird die Position mit dem Abschlag Juli verrechnet.

Bitte beachten Sie, dass unser Portal nur die Retaxdaten der Kasse wiedergibt. Medizinische oder pharmazeutische Unterlagen prüfen wir nicht. Im Portal liegt unter Vorgang RW-882901 ein Scan des Kassenschreibens. Für Rückfragen zur Verrechnung erreichen Sie Frau Bettina Reese unter 0511 407 81 226.

Mit freundlichen Grüßen
Nordwest Rezeptabrechnung eG
Team Retaxationen""",
    )
    write_txt(
        d / "09_chat_wareneingang_20_februar.txt",
        """Export aus dem internen Apothekenchat · Kanal #wareneingang-kühlware
20.02.2026

07:35 Jelena Markowitsch: Box für Frau Elmas ist da. Außenlogger zeigt kurz 9.1, Boxlogger 7.8. CoolMark grün. Ich stelle alles in Quarantänefach A-2.
07:37 Nora Ahlers: Soll ich die Patientin für 10 Uhr erst einmal nicht bestätigen?
07:39 Jelena Markowitsch: Bitte warten. Ich rufe Nordpharm an und dokumentiere beide Gerätenummern.
07:52 Dr. Helena Kröger: Bin ab 08:30 da. Nichts abgeben, bis die Freigabe schriftlich da ist. Bitte auch Praxis wegen Schema anrufen.
09:18 Jelena Markowitsch: Nordpharm-Mail ist eingegangen, Boxlogger im Bereich, Charge freigegeben. Mail liegt im Vorgang HM-26-0218-7741.
09:31 Nora Ahlers: Praxis bestätigt zwei Termine 22.02. und 07.03. Frau Segel schickt noch ein Schreiben.
10:06 Dr. Helena Kröger: Freigabe erteilt. Kühltasche mit zwei Akkus, Lagerung und Fieberhinweis ausführlich erklären. Frau Elmas soll direkt nach Hause.
10:49 Nora Ahlers: Abholung erledigt, unterschriebener Beleg im roten Tagesordner. Sie fährt mit ihrer Schwester, ungefähr zwölf Minuten.""",
    )
    write_csv(
        d / "10_temperaturprotokoll_kuehlkette.csv",
        ["Zeitpunkt", "Gerät", "Ort", "Temperatur_C", "Status", "Bemerkung"],
        [
            ["2026-02-20 05:41", "NL-44018", "Transportbox K-99218", "4.3", "ok", "Packvorgang beendet"],
            ["2026-02-20 06:15", "NL-44018", "Transportbox K-99218", "5.2", "ok", "Tour HB-K2"],
            ["2026-02-20 07:18", "FZ-1182", "Laderaum Türzone", "9.1", "Warnung", "Türöffnung bei erster Zustellung"],
            ["2026-02-20 07:18", "NL-44018", "Transportbox K-99218", "7.8", "ok", "Höchstwert Boxlogger"],
            ["2026-02-20 07:31", "NL-44018", "Wareneingang Apotheke", "7.1", "ok", "Übergabe"],
            ["2026-02-20 08:00", "A2-APO", "Quarantänefach A-2", "5.6", "ok", "Zwischenlagerung"],
            ["2026-02-20 12:00", "A2-APO", "Kühlschrank A-2", "4.9", "ok", "Routinewert"],
            ["2026-02-21 10:35", "A2-APO", "Kühlschrank A-2", "5.1", "ok", "Entnahme zur Abgabe"],
        ],
    )
    write_csv(
        d / "11_abrechnung_und_zahlungsfluss.csv",
        ["Datum", "Beleg", "Beteiligter", "Betrag_EUR", "Buchungsart", "Status", "Bezug"],
        [
            ["2026-02-20", "NP-260220-88173", "Nordpharm Logistik GmbH", "68402.39", "Wareneingang", "bezahlt 2026-03-05", "zwei Packungen SVB26A17"],
            ["2026-02-21", "260221-0187", "Yasemin Elmas", "20.00", "Zuzahlung", "bezahlt bar", "Abgabe"],
            ["2026-03-10", "260221-1042", "HanseMerk BKK", "59842.18", "Kassenforderung", "ausgezahlt 2026-03-31", "Monatsabrechnung Februar"],
            ["2026-06-18", "AM-RX 7741/26", "HanseMerk BKK", "-59842.18", "Retaxankündigung", "offen", "Stellungnahme bis 2026-07-09"],
            ["2026-06-23", "RW-882901", "Nordwest Rezeptabrechnung eG", "-59842.18", "Verrechnungsvormerkung", "geplant 2026-07-14", "Juli-Abschlag"],
            ["2026-07-02", "HK/NM 26-114", "Adler-Apotheke am Wall", "0.00", "Unterlageneingang", "übermittelt", "fünf Anlagen"],
        ],
    )
    write_device_photo(
        d / "12_foto_loggeranzeige.jpg",
        "ThermoSafe NL-44018",
        ["Sendung: K-99218", "Minimum: 4.3 °C", "Maximum: 7.8 °C", "Übergabe: 20.02.2026 · 07:31", "Status: Grenzwert nicht überschritten"],
        (15, 110, 117),
    )

    write_readme(
        d,
        "Apothekenrecht: Retaxation und Kühlkette in Bremen",
        ["`apothekenrecht`"],
        "Die Adler-Apotheke am Wall gibt zwei Packungen eines hochpreisigen, kühlpflichtigen Biologikums auf elektronisches Rezept ab. Monate später kündigt die HanseMerk BKK eine Vollabsetzung von 59842.18 EUR an: Die Dosierangabe sei unzureichend und ein Temperaturwert auffällig. Verordnung, zwei Logger, ärztliche Bestätigung, Warenfluss und Abrechnung erzählen nicht durchgehend dieselbe Geschichte.",
        [
            ("01_e_rezept_abgabeprotokoll.docx", "Vollständiges Abgabeprotokoll mit Verordnung, Wareneingang, Beratung und Abrechnung"),
            ("02_grosshandel_lieferschein_freigabe.docx", "Lieferschein des Großhandels mit Chargen-, Preis- und Loggerdaten"),
            ("03_retaxationsankuendigung_hansemerk_bkk.docx", "Kassenschreiben mit angekündigter Vollabsetzung und Unterlagenfrist"),
            ("04_stellungnahme_rheumatologische_praxis.docx", "Ärztliche Bestätigung zu Verordnung, Dosierung und telefonischer Rückfrage"),
            ("05_einspruch_apotheke_retaxation.docx", "Stellungnahme der Apotheke mit Anlagenbezug und wirtschaftlicher Einordnung"),
            ("06_abholbeleg_beratung.docx", "Abhol- und Beratungsbeleg der Versicherten"),
            ("07_email_nordpharm_qualitaetsfreigabe.eml", "E-Mail der Qualitätssicherung zur Transportfreigabe"),
            ("08_email_rechenzentrum_verrechnung.eml", "E-Mail des Rechenzentrums zur drohenden Verrechnung"),
            ("09_chat_wareneingang_20_februar.txt", "Interner Chat aus der Stunde zwischen Wareneingang und Freigabe"),
            ("10_temperaturprotokoll_kuehlkette.csv", "Rohwerte aus Fahrzeug-, Box- und Apothekenlogger"),
            ("11_abrechnung_und_zahlungsfluss.csv", "Einkauf, Auszahlung, Retaxankündigung und Verrechnungsvormerkung"),
            ("12_foto_loggeranzeige.jpg", "Mobiltelefonfoto der Loggeranzeige"),
        ],
        "Die Unterlagen erlauben die getrennte Prüfung von Abgabe, Dokumentation, Kühlkette, Formmangel, Retaxationshöhe, Frist und Zahlungsfluss. Sie legen keine rechtliche Bewertung fest.",
    )
    write_rubric(
        d,
        slug,
        "apothekenrecht",
        "Retaxationsstreit um ein hochpreisiges Kühlprodukt mit elektronischer Verordnung, zwei widersprüchlich wirkenden Loggern und vollständigem Zahlungsfluss.",
        [
            ("retaxationsgrund", "Beanstandete Dosierangabe und behauptete Temperaturabweichung getrennt prüfen und jeweils an Primärbelegen festmachen."),
            ("formfehler", "Paragraf 129 Abs. 4 SGB V und die einschlägigen Rahmenvertragsregeln zur vollständigen oder teilweisen Retaxation wegen Formfehlern in die Prüfung einordnen."),
            ("kuehlkette", "Fahrzeuglogger, Boxlogger, Einwegindikator, Freigabemail und Apothekenkühlschrank zeitlich sauber auseinanderhalten."),
            ("zahlung", "Einkauf, ursprüngliche Auszahlung, angekündigte Retaxation und geplante Verrechnung betragsgenau abstimmen."),
        ],
    )


def build_krankenhausrecht() -> None:
    slug = "krankenhausrecht-md-intensivabrechnung-halle"
    d = prepare(slug)
    clinic_contact = "Merseburger Straße 118 · 06110 Halle (Saale)\nTelefon 0345 728 0 · Fax 0345 728 2299\npost@klinikum-saalebogen.de · IK 261 500 447"
    insurer_contact = "Willy-Brandt-Platz 7 · 04109 Leipzig\nTelefon 0341 219 77 0 · krankenhaus@gesundhanse-bkk.de"
    write_record(
        d / "01_entlassungsbericht_intensivstation.docx",
        org="Klinikum Saalebogen gGmbH",
        strapline="Klinik für Anästhesiologie und operative Intensivmedizin",
        contact=clinic_contact,
        title="Entlassungsbericht Intensivstation ITS 2",
        metadata=[("Patient", "Jochen Wendt, geb. 04.06.1959"), ("Fallnummer", "260112-88341"), ("Aufenthalt ITS", "12.01.2026, 19:42 Uhr bis 28.01.2026, 11:18 Uhr"), ("Weiterbehandlung", "Allgemeinstation C3 bis 03.02.2026")],
        sections=[
            ("1. Aufnahmegrund", ["Notfallmäßige Übernahme aus dem Operationssaal nach Ösophagusresektion wegen einer Anastomoseninsuffizienz mit Mediastinitis und septischem Kreislaufversagen. Bei Aufnahme war der Patient invasiv beatmet, mit Noradrenalin kreislaufunterstützt und über einen zentralvenösen Katheter versorgt. Ehefrau Irmgard Wendt wurde am 12.01.2026 um 21:06 Uhr telefonisch informiert."]),
            ("2. Verlauf", ["Am 13.01. erfolgten Revision und Spülung des Mediastinums. Vom 12.01. bis 20.01. bestand invasive Beatmung, anschließend High-Flow-Sauerstofftherapie mit wiederholten nichtinvasiven Unterstützungsphasen. Nierenersatztherapie wurde vom 14.01. bis 18.01. durchgeführt. Eine erneute Kreislaufinstabilität trat in der Nacht vom 21. auf den 22.01. auf. Oberärztin Dr. Leonie Tesch übernahm die Behandlungsleitung am 19.01. während der Fortbildung von Chefarzt Prof. Dr. Birkner.", "Die täglichen SAPS-II- und TISS-Erhebungen wurden im Intensivsystem dokumentiert. Für den 17.01. fehlt im exportierten Datensatz die elektronische Signatur des Frühdienstes; der Papierbogen trägt die Handzeichen LZ und MB. Für den 22.01. wurde ein nachträglicher Korrektureintrag angelegt, weil die Dialyseleistung zunächst dem Vortag zugeordnet war."]),
            ("3. Entlassungszustand", ["Bei Verlegung auf Station C3 war Herr Wendt wach, kontaktfähig und mit zwei Litern Sauerstoff über Nasenbrille stabil. Die Ernährung erfolgte über Jejunalsonde. Physiotherapeutische Mobilisation bis an die Bettkante war möglich. Die weitere Wundversorgung und antibiotische Therapie wurden an die Station übergeben."]),
        ],
        signer="Dr. med. Leonie Tesch",
        signer_role="Oberärztin · Zusatzweiterbildung Intensivmedizin",
    )
    write_record(
        d / "02_intensivmedizinischer_leistungsnachweis.docx",
        org="Klinikum Saalebogen gGmbH",
        strapline="Medizincontrolling · Fallprüfung",
        contact=clinic_contact,
        title="Leistungsnachweis Fall 260112-88341",
        metadata=[("Erstellt", "10.02.2026, 15:20 Uhr"), ("Kodierfachkraft", "Sabine Lodemann"), ("Abrechnungsstand", "Version 3 nach Freigabe Medizincontrolling"), ("OPS-Gruppe", "8-98f, Aufwandspunkte aus SAPS II und TISS")],
        sections=[
            ("1. Behandlungsorganisation", ["Die Station ITS 2 verfügt über zwölf Betten. Die ärztliche Behandlungsleitung lag bei Prof. Dr. Marius Birkner und vom 19.01. bis 23.01. bei Dr. Leonie Tesch. Die Dienstpläne weisen in diesem Zeitraum rund um die Uhr einen auf der Station eingesetzten Arzt aus. Am 24.01. war Dr. Tesch von 09:05 bis 10:14 Uhr bei einer Reanimation auf Station B2; Assistenzarzt Dr. Pohl blieb auf ITS 2."]),
            ("2. Dokumentationsabweichungen", ["Für den 17.01. liegt im Datenexport kein finaler Freigabestatus vor, obwohl der Papierbogen Werte und Handzeichen enthält. Für den 22.01. ist die Dialyseleistung durch einen Korrektureintrag vom 23.01. ergänzt. Die ursprüngliche und die korrigierte Version sind im Archiv erhalten. In der Pflegeplanung des Nachtdienstes 16./17.01. fehlt die zweite Gegenzeichnung für die Lagerungsmaßnahme."]),
            ("3. Abrechnung", ["Die Summe des Klinikexports beträgt 1246 Aufwandspunkte. Ohne den 17.01. und ohne den korrigierten Dialyseeintrag ergäben sich 1038 Punkte. Die Gruppierung führte zusammen mit den Beatmungs- und Operationsdaten zu einem Rechnungsbetrag von 132488.17 EUR einschließlich Pflegeentgelt. Die Krankenkasse zahlte zunächst unter Vorbehalt."]),
        ],
        tables=[("4. Tageswerte", ["Datum", "SAPS II", "TISS", "Summe", "Freigabe"], [["12.01.", "41", "32", "73", "signiert"], ["13.01.", "44", "38", "82", "signiert"], ["17.01.", "36", "29", "65", "Papierbogen"], ["22.01.", "39", "37", "76", "korrigiert"], ["28.01.", "18", "14", "32", "signiert"]], [2.7, 2.7, 2.7, 2.7, 5.0])],
        signer="Sabine Lodemann",
        signer_role="Leitende Kodierfachkraft",
    )
    write_letter(
        d / "03_pruefanzeige_gesundhanse_bkk.docx",
        org="GesundHanse BKK",
        strapline="Krankenhausfallmanagement",
        contact=insurer_contact,
        recipient=["Klinikum Saalebogen gGmbH", "Medizincontrolling", "Merseburger Straße 118", "06110 Halle (Saale)"],
        date="19.02.2026",
        reference="KH-260112-88341 / GH-922174",
        subject="Prüfanzeige zum Behandlungsfall Jochen Wendt – Aufnahme 12.01.2026",
        salutation="Sehr geehrte Damen und Herren,",
        sections=[
            ("1. Prüfgegenstand", ["Wir leiten eine Prüfung der sachlich-rechnerischen Richtigkeit und der Voraussetzungen der abgerechneten intensivmedizinischen Komplexbehandlung ein. Auffällig sind die nachträgliche Änderung von Aufwandspunkten, ein fehlender Freigabestatus am 17.01.2026 und die dokumentierte Abwesenheit der Behandlungsleitung am 24.01.2026."]),
            ("2. Beauftragter Medizinischer Dienst", ["Der Medizinische Dienst Sachsen-Anhalt erhält den Prüfauftrag elektronisch. Bitte übermitteln Sie die angeforderten Unterlagen ausschließlich über den vereinbarten Übermittlungsweg. Die Anforderung des Medizinischen Dienstes bestimmt Umfang und Frist. Nicht angeforderte Unterlagen werden von uns nicht vorab bewertet."]),
            ("3. Zahlungsstatus", ["Die Rechnung vom 10.02.2026 über 132488.17 EUR wurde am 17.02.2026 unter Vorbehalt ausgeglichen. Eine spätere Korrektur oder Aufrechnung bleibt vorbehalten. Unser Schreiben enthält noch keine abschließende Kürzungsentscheidung."]),
        ],
        closing="Mit freundlichen Grüßen",
        signer="Marvin Kleeberg",
        signer_role="Fallmanager stationäre Versorgung · Durchwahl 0341 219 77 416",
    )
    write_letter(
        d / "04_unterlagenanforderung_md_sachsen_anhalt.docx",
        org="Medizinischer Dienst Sachsen-Anhalt",
        strapline="Fachbereich Krankenhaus · Prüfteam Süd",
        contact="Breiter Weg 16 · 39104 Magdeburg\nTelefon 0391 5661 0 · Fax 0391 5661 188\nkh-pruefung@md-san.de",
        recipient=["Klinikum Saalebogen gGmbH", "Medizincontrolling", "Merseburger Straße 118", "06110 Halle (Saale)"],
        date="23.02.2026",
        reference="MD-KH 26/18841",
        subject="Unterlagenanforderung – Fall Jochen Wendt – Fallnummer 260112-88341",
        salutation="Sehr geehrte Damen und Herren,",
        sections=[
            ("1. Benötigte Unterlagen", ["Bitte übermitteln Sie bis zum 23.03.2026 den vollständigen Intensivkurvenexport, die SAPS-II- und TISS-Tagesbögen, Beatmungsdaten, Dialyseprotokolle, ärztlichen und pflegerischen Dienstpläne, Nachweise über die Behandlungsleitung sowie die Versionshistorie der Korrektureinträge vom 22. und 23.01.2026."]),
            ("2. Fragestellung", ["Zu beurteilen ist, ob die dokumentierten Leistungen und Strukturvoraussetzungen für die abgerechnete OPS-Gruppe im gesamten geltend gemachten Zeitraum vorlagen und ob die Aufwandspunkte nachvollziehbar ermittelt wurden. Die medizinische Notwendigkeit der stationären Behandlung als solche ist nicht Gegenstand des Auftrags."]),
            ("3. Übermittlung", ["Jede Datei ist mit Fallnummer und Dokumentart zu benennen. Bei technischen Schwierigkeiten ist vor Fristablauf Frau Nele Thormann unter der Durchwahl 0391 5661 442 zu kontaktieren. Ein Uploadprotokoll wird nach Abschluss der Übertragung bereitgestellt."]),
        ],
        closing="Mit freundlichen Grüßen",
        signer="Dr. med. Jasper Rohwedder",
        signer_role="Facharzt für Anästhesiologie · Sozialmedizin",
    )
    write_record(
        d / "05_md_gutachten_nach_aktenlage.docx",
        org="Medizinischer Dienst Sachsen-Anhalt",
        strapline="Fachbereich Krankenhaus · Prüfteam Süd",
        contact="Breiter Weg 16 · 39104 Magdeburg\nTelefon 0391 5661 0 · kh-pruefung@md-san.de",
        title="Gutachtliche Stellungnahme nach Aktenlage",
        metadata=[("Aktenzeichen", "MD-KH 26/18841"), ("Versicherter", "Jochen Wendt, geb. 04.06.1959"), ("Begutachtung", "28.04.2026"), ("Gutachter", "Dr. med. Jasper Rohwedder")],
        sections=[
            ("1. Auftrag und Material", ["Geprüft wurden Abrechnung, Entlassungsbericht, Intensivkurvenexport, 17 Tagesbögen, Beatmungs- und Dialyseprotokolle, Dienstpläne sowie die vom Krankenhaus am 22.03.2026 übermittelte Versionshistorie. Der Papierbogen für den 17.01.2026 lag als Scan vor. Die Rückseite mit Erläuterungen war im Upload nicht enthalten."]),
            ("2. Medizinischer Verlauf", ["Die intensivmedizinische Behandlungsbedürftigkeit während des streitigen Zeitraums ist anhand der Organunterstützung und der wiederholten Kreislaufinstabilität nachvollziehbar. Die Beatmungszeiten lassen sich aus Geräteexport und Kurve im Wesentlichen abgleichen. Für zwei High-Flow-Intervalle unterscheiden sich die Zeitangaben um jeweils ungefähr 35 Minuten."]),
            ("3. Aufwandspunkte", ["Die elektronisch signierten Tageswerte sind rechnerisch nachvollziehbar. Für den 17.01. fehlt die Rückseite des Papierbogens; die dort angesetzten 65 Punkte können aus dem übrigen Export nicht vollständig rekonstruiert werden. Beim 22.01. ist der Korrektureintrag technisch nachvollziehbar, die medizinische Begründung nennt jedoch keine Uhrzeit der Dialyseaufnahme."]),
            ("4. Behandlungsleitung", ["Die Dienstpläne benennen Prof. Dr. Birkner und Dr. Tesch. Für den 24.01. ist eine Abwesenheit von 69 Minuten wegen eines Reanimationseinsatzes dokumentiert. Ob und wie die Verantwortung während dieses Einsatzes organisatorisch fortgeführt wurde, ergibt sich aus der vorgelegten Organisationsanweisung nur teilweise."]),
            ("5. Mitteilung an die Krankenkasse", ["Die abschließende sozialmedizinische Beurteilung wird der beauftragenden Krankenkasse im vorgesehenen elektronischen Datensatz übermittelt. Dieses Exemplar dokumentiert die tatsächlichen Feststellungen und die verbliebenen Dokumentationslücken."]),
        ],
        signer="Dr. med. Jasper Rohwedder",
        signer_role="Facharzt für Anästhesiologie · Sozialmedizin",
        accent=NAVY,
    )
    write_letter(
        d / "06_aufrechnungsmitteilung_krankenkasse.docx",
        org="GesundHanse BKK",
        strapline="Krankenhausfallmanagement",
        contact=insurer_contact,
        recipient=["Klinikum Saalebogen gGmbH", "Debitorenbuchhaltung", "Merseburger Straße 118", "06110 Halle (Saale)"],
        date="12.05.2026",
        reference="KH-260112-88341 / GH-922174",
        subject="Rechnungskorrektur und Aufrechnung – 74612.80 EUR",
        salutation="Sehr geehrte Damen und Herren,",
        sections=[
            ("1. Korrekturbetrag", ["Nach Abschluss der Prüfung erkennen wir die abgerechnete intensivmedizinische Komplexbehandlung nicht im geltend gemachten Umfang an. Aus der Neugruppierung ergibt sich gegenüber der Zahlung vom 17.02.2026 eine Differenz von 74612.80 EUR. Die medizinische Notwendigkeit des stationären Aufenthalts bleibt unberührt."]),
            ("2. Verrechnung", ["Wir rechnen den Betrag mit den in der Anlage bezeichneten unstreitigen Forderungen aus den Fällen 260401-11402 bis 260418-15577 auf. Die Buchung ist für den Zahlungslauf 20.05.2026 vorgesehen. Der Prüfaufschlag wird nach gesonderter Prüfung abgerechnet."]),
            ("3. Erörterung", ["Falls Sie die Korrektur nicht akzeptieren, teilen Sie dies bis zum 02.06.2026 unter Angabe unseres Zeichens mit. Benennen Sie bitte konkret, welche Tatsachenfeststellung oder Berechnung angegriffen wird. Die bereits ausgetauschten Unterlagen werden nicht erneut angefordert."]),
        ],
        closing="Mit freundlichen Grüßen",
        signer="Marvin Kleeberg",
        signer_role="Fallmanager stationäre Versorgung",
        attachments=["Aufstellung der Aufrechnungsforderungen", "Korrekturdatensatz vom 11.05.2026"],
    )
    write_email(
        d / "07_email_medizincontrolling_an_intensivstation.eml",
        sender="Sabine Lodemann <s.lodemann@klinikum-saalebogen.de>",
        recipients=["Dr. Leonie Tesch <l.tesch@klinikum-saalebogen.de>"],
        cc=["Prof. Dr. Marius Birkner <m.birkner@klinikum-saalebogen.de>", "Recht <recht@klinikum-saalebogen.de>"],
        date="Thu, 14 May 2026 08:26:00 +0200",
        subject="Fall Wendt / Aufrechnung 74612.80 EUR / Rückseite 17.01.",
        message_id="<20260514.082600.26011288341@klinikum-saalebogen.de>",
        body="""Guten Morgen Frau Dr. Tesch,

die Kasse hat gestern die Kürzung eingestellt. Der MD schreibt, beim 17.01. habe die Rückseite des Papierbogens gefehlt. In unserem Scanordner liegt tatsächlich nur die Vorderseite. Frau Kühne aus dem Archiv meint, der Originalordner sei nach dem Umzug der ITS 2 noch in Kiste 14 im Kellerraum K-03.

Können Sie außerdem kurz erklären, wer am 24.01. zwischen 09:05 und 10:14 Uhr die unmittelbare Stationsverantwortung hatte, als Sie bei der Reanimation auf B2 waren? Im Dienstplan steht Dr. Pohl, in der Organisationsanweisung aber nur die allgemeine Vertretungsregel. Ich brauche keine juristische Einschätzung, sondern Namen, Uhrzeiten und vorhandene Belege.

Viele Grüße
Sabine Lodemann
Leitende Kodierfachkraft
Telefon 0345 728 3841""",
    )
    write_email(
        d / "08_email_archiv_fund_papierbogen.eml",
        sender="Kathrin Kühne <k.kuehne@klinikum-saalebogen.de>",
        recipients=["Sabine Lodemann <s.lodemann@klinikum-saalebogen.de>"],
        cc=["Dr. Leonie Tesch <l.tesch@klinikum-saalebogen.de>"],
        date="Mon, 18 May 2026 16:12:00 +0200",
        subject="Kiste 14 gefunden – Tagesbogen Wendt 17.01. vollständig",
        message_id="<archiv.20260518.161200.k14@klinikum-saalebogen.de>",
        body="""Hallo Frau Lodemann,

wir haben Kiste 14 heute um 15:35 Uhr mit Herrn Böhme aus dem Gebäudeservice geöffnet. Der Tagesbogen Wendt vom 17.01. liegt vollständig darin, Vorder- und Rückseite sind zusammengetackert. Auf der Rückseite stehen die Einzelwerte, die Handzeichen LZ und MB sowie ein Vermerk zur Noradrenalinsteigerung um 03:18 Uhr.

Ich habe beide Seiten mit 300 dpi als eine Datei gescannt und im revisionsgeschützten Archiv unter 260112-88341 abgelegt. Das Papieroriginal bleibt in Fach K-03-14-7. Bitte sagt mir, ob ein dokumentierter Auszug an die Kasse oder nur an die Rechtsabteilung gehen soll.

Viele Grüße
Kathrin Kühne
Zentralarchiv · Durchwahl 2274""",
        attachments=["260112-88341_tagesbogen_2026-01-17_vollstaendig.pdf"],
    )
    write_txt(
        d / "09_chat_dienstplan_und_reanimation.txt",
        """Export aus dem Stationschat ITS 2 · 24.01.2026

08:57 Dr. Leonie Tesch: Morgenrunde abgeschlossen. Jochen Wendt heute weiter High Flow 35 Liter, Kreislauf stabil.
09:04 Zentrale: Reanimation B2 Zimmer 214, Intensivteam sofort.
09:05 Dr. Leonie Tesch: Ich gehe mit Pflegekraft Möller nach B2. Daniel, du bleibst auf ITS 2 und rufst mich bei jeder Änderung direkt an.
09:06 Dr. Daniel Pohl: Verstanden. Ich bin am Arztstützpunkt, Telefon 419.
09:42 Dr. Daniel Pohl: Bei Wendt MAP 68, Noradrenalin unverändert. Blutgas ist da, ich stelle es in die Kurve.
10:03 Dr. Leonie Tesch: Reanimation läuft noch. Gibt es auf ITS etwas Dringendes?
10:04 Dr. Daniel Pohl: Nein. Wendt stabil, Zimmer 8 Bronchoskopie beendet, Oberarzt Riemer ist jetzt auch hier.
10:14 Dr. Leonie Tesch: Zurück auf ITS 2.
10:18 Pflegekraft Eva Lenz: Wendt gelagert, Sättigung 95 Prozent. Eintrag ist fertig.""",
    )
    write_csv(
        d / "10_saps_tiss_tageswerte.csv",
        ["Datum", "SAPS_II", "TISS", "Gesamt", "Quelle", "Freigabe", "Hinweis"],
        [["2026-01-12", 41, 32, 73, "Intensivsystem", "signiert", "Aufnahme 19:42"], ["2026-01-13", 44, 38, 82, "Intensivsystem", "signiert", "Revision"], ["2026-01-14", 47, 41, 88, "Intensivsystem", "signiert", "Dialysebeginn"], ["2026-01-15", 42, 35, 77, "Intensivsystem", "signiert", ""], ["2026-01-16", 38, 31, 69, "Intensivsystem", "signiert", ""], ["2026-01-17", 36, 29, 65, "Papierbogen", "Handzeichen LZ MB", "Rückseite im Erstupload fehlte"], ["2026-01-18", 40, 34, 74, "Intensivsystem", "signiert", ""], ["2026-01-19", 35, 27, 62, "Intensivsystem", "signiert", "Leitungswechsel"], ["2026-01-20", 30, 24, 54, "Intensivsystem", "signiert", "Extubation"], ["2026-01-21", 37, 31, 68, "Intensivsystem", "signiert", ""], ["2026-01-22", 39, 37, 76, "Intensivsystem", "korrigiert 2026-01-23", "Dialysezeit ergänzt"], ["2026-01-23", 33, 26, 59, "Intensivsystem", "signiert", ""], ["2026-01-24", 29, 23, 52, "Intensivsystem", "signiert", "Reanimation B2"], ["2026-01-25", 28, 21, 49, "Intensivsystem", "signiert", ""], ["2026-01-26", 24, 19, 43, "Intensivsystem", "signiert", ""], ["2026-01-27", 21, 17, 38, "Intensivsystem", "signiert", ""], ["2026-01-28", 18, 14, 32, "Intensivsystem", "signiert", "Verlegung 11:18"]],
    )
    write_csv(
        d / "11_rechnung_aufrechnung_fallliste.csv",
        ["Fallnummer", "Rechnungsdatum", "Ursprungsbetrag_EUR", "Zahlung_EUR", "Aufrechnung_EUR", "Status", "Bezug"],
        [["260112-88341", "2026-02-10", "132488.17", "132488.17", "-74612.80", "streitige Kürzung", "Jochen Wendt"], ["260401-11402", "2026-04-12", "18442.50", "0.00", "18442.50", "für Aufrechnung benannt", "unstrittig"], ["260405-12881", "2026-04-16", "27611.20", "0.00", "27611.20", "für Aufrechnung benannt", "unstrittig"], ["260411-14309", "2026-04-22", "19874.10", "0.00", "19874.10", "für Aufrechnung benannt", "unstrittig"], ["260418-15577", "2026-04-29", "8685.00", "0.00", "8685.00", "für Aufrechnung benannt", "unstrittig"]],
    )
    write_device_photo(d / "12_foto_archivkiste_14.jpg", "Zentralarchiv · Keller K-03", ["Kiste 14 / ITS 2", "Fach 7: Wendt, Jochen", "Tagesbogen 17.01.2026", "Vorder- und Rückseite vorhanden", "Fund: 18.05.2026 · 15:35 Uhr"], (43, 73, 101))
    write_readme(d, "Krankenhausrecht: MD-Prüfung einer Intensivabrechnung in Halle", ["`krankenhausrecht`"], "Das Klinikum Saalebogen rechnet einen langen intensivmedizinischen Verlauf mit 1246 Aufwandspunkten ab. Die Krankenkasse zahlt zunächst, veranlasst dann eine Prüfung und rechnet 74612.80 EUR auf. Ein fehlender Bogenscan, eine nachgetragene Dialyseleistung und ein 69-minütiger Reanimationseinsatz außerhalb der Station prägen den Streit.", [("01_entlassungsbericht_intensivstation.docx", "Ärztlicher Entlassungsbericht mit Verlauf und Dokumentationsabweichungen"), ("02_intensivmedizinischer_leistungsnachweis.docx", "Leistungsnachweis des Medizincontrollings"), ("03_pruefanzeige_gesundhanse_bkk.docx", "Prüfanzeige der Krankenkasse"), ("04_unterlagenanforderung_md_sachsen_anhalt.docx", "Konkrete Unterlagenanforderung des Medizinischen Dienstes"), ("05_md_gutachten_nach_aktenlage.docx", "Gutachtliche Tatsachenfeststellungen nach Aktenlage"), ("06_aufrechnungsmitteilung_krankenkasse.docx", "Kürzungs- und Aufrechnungsmitteilung"), ("07_email_medizincontrolling_an_intensivstation.eml", "Interne Nachfrage zu Scanlücke und Stationsverantwortung"), ("08_email_archiv_fund_papierbogen.eml", "Nachricht über den späteren Fund des vollständigen Papierbogens"), ("09_chat_dienstplan_und_reanimation.txt", "Zeitnaher Stationschat zum Reanimationseinsatz"), ("10_saps_tiss_tageswerte.csv", "Sämtliche Tageswerte mit Quellen- und Freigabestatus"), ("11_rechnung_aufrechnung_fallliste.csv", "Abgleich von Rechnung, Zahlung und Aufrechnungsfällen"), ("12_foto_archivkiste_14.jpg", "Mobiltelefonfoto der aufgefundenen Archivkiste")], "Die Akte lässt Prüfauftrag, Unterlagenfrist, OPS-Dokumentation, Behandlungsleitung, nachgereichte Belege und Aufrechnung getrennt bearbeiten. Sie enthält keine abschließende rechtliche oder medizinische Lösung.")
    write_rubric(d, slug, "krankenhausrecht", "Krankenhausabrechnungsprüfung zu OPS 8-98f, Aufwandspunkten, Behandlungsleitung, Nachreichung und Aufrechnung.", [("pruefverfahren", "Prüfanzeige, MD-Unterlagenanforderung, Gutachten, Kassenentscheidung und Erörterungsweg nach Paragraf 275c SGB V und Paragraf 17c KHG zeitlich trennen."), ("ops", "Die 2026 maßgeblichen Mindestmerkmale und Aufwandspunkte der OPS-Gruppe 8-98f am amtlichen Katalog prüfen."), ("dokumente", "Erstupload, später gefundene Bogenrückseite und Korrektureintrag quellenkritisch würdigen, ohne bloß auf das Gutachten zu vertrauen."), ("aufrechnung", "Die Summe der vier benannten Aufrechnungsforderungen exakt mit 74612.80 EUR abstimmen und prozessuale Folgen gesondert prüfen.")])


def build_eu_prozessrecht() -> None:
    slug = "eu-prozessrecht-beihilfe-nichtigkeitsklage-leipzig"
    d = prepare(slug)
    company = "Lipsia Speichertechnik AG"
    company_contact = "Torgauer Straße 231 · 04347 Leipzig\nTelefon 0341 604 72 0 · legal@lipsia-speicher.de\nHRB 41882 Amtsgericht Leipzig"
    law_contact = "Taunusanlage 18 · 60325 Frankfurt am Main\nTelefon 069 348 19 0 · eu-litigation@westphal-counsel.de\nPartnerschaftsregister PR 2881 AG Frankfurt am Main"
    write_letter(d / "01_zustellung_kommissionsentscheidung.docx", org="Europäische Kommission", strapline="Generaldirektion Wettbewerb · Staatliche Beihilfen", contact="1049 Brüssel · Belgien\nElektronische Zustellung über EU Secure Exchange\nRegistratur COMP-GREFFE-STATE-AID", recipient=["Ständige Vertretung der Bundesrepublik Deutschland", "Rue Jacques de Lalaing 8-14", "1040 Brüssel", "Kopie: Lipsia Speichertechnik AG"], date="17.04.2026", reference="SA.118742 (2025/C) · C(2026) 2417 final", subject="Staatliche Beihilfe – kommunale Garantie zugunsten der Lipsia Speichertechnik AG", salutation="Sehr geehrte Damen und Herren,", sections=[("1. Entscheidung", ["Die Kommission hat am 15.04.2026 die Entscheidung C(2026) 2417 final erlassen. Nach ihrer Auffassung stellt die von der Stadt Leipzig am 30.06.2023 übernommene Garantie für ein Darlehen der Mitteldeutschen Aufbaubank eine mit dem Binnenmarkt unvereinbare staatliche Beihilfe dar. Deutschland wird aufgefordert, den Vorteil nebst Zinsen vom Begünstigten zurückzufordern."]), ("2. Zustellung", ["Die authentische deutsche Fassung und die Anlagen wurden am 17.04.2026 um 10:06 Uhr im sicheren Austauschraum bereitgestellt. Der Empfang durch die Ständige Vertretung wurde um 10:21 Uhr bestätigt. Eine Informationskopie ging gleichzeitig an die im Verwaltungsverfahren benannte Rechtsabteilung der Lipsia Speichertechnik AG."]), ("3. Weiteres Verfahren", ["Die Entscheidung nennt die Rechtsbehelfsbelehrung in Abschnitt XII. Veröffentlichung und Behandlung vertraulicher Angaben erfolgen gesondert. Die Kommission bittet Deutschland binnen zwei Monaten um eine Übersicht der bereits ergriffenen Rückforderungsmaßnahmen."])], closing="Mit freundlichen Grüßen", signer="Elise Van der Meer", signer_role="Referatsleiterin COMP.H.3", attachments=["Entscheidung C(2026) 2417 final, deutsche Fassung", "Vertrauliche Anlage 1", "Zustellprotokoll EU Secure Exchange"])
    write_record(d / "02_vorstandsprotokoll_sondersitzung.docx", org=company, strapline="Vorstandsbüro · Vertrauliche Niederschrift", contact=company_contact, title="Niederschrift der Vorstandssondersitzung", metadata=[("Datum", "20.04.2026, 08:30 bis 10:05 Uhr"), ("Ort", "Leipzig, Besprechungsraum Elster"), ("Teilnehmer", "Nora Linden, Viktor Hagedorn, Dr. Edda Noll, RA Dr. Felix Westphal zugeschaltet"), ("Protokoll", "Miriam Kroll")], sections=[("1. Eingang der Entscheidung", ["Syndikusrechtsanwältin Dr. Edda Noll berichtet, dass die Informationskopie am 17.04.2026 um 10:08 Uhr im Funktionspostfach legal@lipsia-speicher.de eingegangen sei. Die Stadt Leipzig habe am 17.04. um 16:42 Uhr telefonisch angekündigt, den Rückforderungsbetrag vorläufig mit rund 2.74 Mio. EUR einschließlich Zinsen zu berechnen. Ein förmlicher Leistungsbescheid liege noch nicht vor."]), ("2. Liquidität und Außenwirkung", ["Finanzvorstand Viktor Hagedorn weist auf zwei Kreditklauseln hin, die bei einer sofortigen Rückforderung eine Zustimmung der Konsortialbanken verlangen. Die nächste Ziehung für das Werk Delitzsch ist am 04.05.2026 vorgesehen. Vorstandsvorsitzende Nora Linden bittet, externe Kommunikation bis zur Abstimmung mit der Stadt und den Banken auf die Bestätigung des Eingangs zu beschränken."]), ("3. Verfahrensschritte", ["Dr. Westphal nennt als mögliche Schritte eine Klage beim Gericht der Europäischen Union und einen Antrag auf vorläufigen Rechtsschutz. Er bittet um die vollständige Verwaltungsakte, die ursprüngliche Garantiezusage, alle Finanzierungsangebote ohne Garantie und die Berechnungen der Kommission. Im Kalender des Vorstandsbüros wird zunächst der 26.06.2026 als interne Sicherheitsfrist eingetragen. Frau Kroll soll prüfen, ob der 27.06.2026 auf ein Wochenende fällt."]), ("4. Beschluss", ["Der Vorstand beauftragt die Kanzlei mit der Vorbereitung der gerichtlichen Schritte und ermächtigt Dr. Noll, e-Curia-Unterlagen, Vollmacht und Identitätsnachweise zusammenzustellen. Über die Einreichung soll nach Vorlage eines vollständigen Entwurfs entschieden werden."])], signer="Miriam Kroll", signer_role="Leiterin Vorstandsbüro")
    write_letter(d / "03_aufforderung_rueckforderung_stadt_leipzig.docx", org="Stadt Leipzig", strapline="Dezernat Wirtschaft, Arbeit und Digitales · Referat Beteiligungen", contact="Martin-Luther-Ring 4-6 · 04109 Leipzig\nTelefon 0341 123 0 · beteiligungen@leipzig.de", recipient=[company, "Vorstand", "Torgauer Straße 231", "04347 Leipzig"], date="04.05.2026", reference="VII-20.31/LSA-2026-17", subject="Anhörung zur Rückforderung im Verfahren SA.118742 (2025/C)", salutation="Sehr geehrte Frau Linden, sehr geehrter Herr Hagedorn,", sections=[("1. Gegenstand", ["Die Bundesrepublik Deutschland ist nach der Entscheidung C(2026) 2417 final gehalten, den nach Auffassung der Kommission gewährten Vorteil zurückzufordern. Die Stadt bereitet als Garantiegeberin die Festsetzung gegenüber Ihrer Gesellschaft vor. Vor Erlass des Bescheids erhalten Sie Gelegenheit, sich zu Berechnungsgrundlagen und Zahlungsmodalitäten zu äußern."]), ("2. Vorläufige Berechnung", ["Ausgehend von der Garantiesumme, dem von der Kommission angesetzten marktüblichen Garantieentgelt und den bislang gemeldeten Zinsperioden ergibt sich vorläufig ein Betrag von 2743513.06 EUR zum 30.04.2026. Die beigefügte Tabelle enthält die quartalsweise Berechnung. Änderungen bleiben nach Abstimmung mit Bund und Kommission möglich."]), ("3. Frist", ["Ihre Stellungnahme erwarten wir bis zum 25.05.2026. Teilen Sie bitte zugleich mit, ob Sie eine Ratenzahlung beantragen und welche Sicherheiten angeboten werden. Ein gegen die Kommissionsentscheidung eingeleitetes gerichtliches Verfahren hindert die Vorbereitung der Rückforderung nicht ohne Weiteres."])], closing="Mit freundlichen Grüßen", signer="Dr. Amelie Hartung", signer_role="Referatsleiterin Beteiligungen", attachments=["Vorläufige Zinsberechnung", "Auszug aus C(2026) 2417 final"])
    write_letter(d / "04_klageschrift_fassung_24_juni.docx", org="Westphal · Dierkes · Mohn Partnerschaft mbB", strapline="Rechtsanwälte · European Litigation", contact=law_contact, recipient=["Gericht der Europäischen Union", "Kanzlei", "Rue du Fort Niedergrünewald", "L-2925 Luxemburg"], date="24.06.2026", reference="FW/EN 61-2026", subject="Klage der Lipsia Speichertechnik AG gegen die Europäische Kommission", salutation="Namens und in Vollmacht der Klägerin erheben wir Klage gegen die Entscheidung C(2026) 2417 final vom 15.04.2026.", sections=[("1. Parteien und angegriffener Rechtsakt", ["Klägerin ist die Lipsia Speichertechnik AG, Torgauer Straße 231, 04347 Leipzig, Deutschland, vertreten durch ihren Vorstand Nora Linden und Viktor Hagedorn. Beklagte ist die Europäische Kommission. Angegriffen wird die der Bundesrepublik Deutschland am 17.04.2026 zugestellte Entscheidung im Verfahren SA.118742 (2025/C), soweit sie die kommunale Garantie als unvereinbare Beihilfe einordnet und deren Rückforderung verlangt."]), ("2. Anträge", ["Die Klägerin beantragt, die Entscheidung C(2026) 2417 final für nichtig zu erklären, soweit sie die Klägerin betrifft, der Kommission die Kosten aufzuerlegen und die in Anlage A.12 bezeichneten Geschäftsgeheimnisse gegenüber möglichen Streithelfern vertraulich zu behandeln."]), ("3. Erster Klagegrund", ["Die Kommission habe den wirtschaftlichen Vorteil fehlerhaft bestimmt. Das im Verwaltungsverfahren vorgelegte Angebot der Elbe Privatbank vom 14.06.2023 sei nicht mit der garantierten Finanzierung vergleichbar gewesen. Laufzeit, Rang, Sicherheiten und Auszahlungsbedingungen wichen ab. Die Entscheidung setze gleichwohl die Differenz der Entgelte ohne weitere Anpassung an."]), ("4. Zweiter Klagegrund", ["Die Klägerin rügt, dass die Kommission die Stellungnahme vom 19.11.2025 und das Gutachten der Wirtschaftsprüfungsgesellschaft Harenberg Consult nur auszugsweise würdige. Aus der Entscheidung sei nicht erkennbar, weshalb die dort berechnete marktübliche Prämie verworfen werde."]), ("5. Dritter Klagegrund", ["Hilfsweise wird geltend gemacht, dass die Rückforderungsberechnung die bereits gezahlten Garantieentgelte und die Verringerung der ausstehenden Darlehenssumme nicht periodengerecht berücksichtige. Die von der Stadt übersandte vorläufige Berechnung weiche zudem von Anlage 3 der Entscheidung ab."]), ("6. Anlagen und Vertraulichkeit", ["Die Anlagen A.1 bis A.14 werden über e-Curia eingereicht. Für A.7, A.9 und A.12 liegt jeweils eine nicht vertrauliche Fassung vor. Das Vorstandsprotokoll vom 20.04.2026 wird nur zur Begründung des Antrags auf vorläufigen Rechtsschutz verwendet und ist nicht Bestandteil dieser Klagefassung."])], closing="Dr. Felix Westphal · Rechtsanwalt", signer="Dr. Elin Nørgaard", signer_role="Advokatin und Rechtsanwältin", attachments=["A.1 Vollmacht", "A.2 Handelsregisterauszug", "A.3 angegriffene Entscheidung", "A.4 Zustellnachweis", "A.5 bis A.14 Finanzierungs- und Verfahrensunterlagen"])
    write_record(d / "05_e_curia_eingangsbestaetigung.docx", org="Gericht der Europäischen Union", strapline="Kanzlei · e-Curia", contact="Rue du Fort Niedergrünewald · L-2925 Luxemburg\ncuria.europa.eu · Technischer Dienst e-Curia", title="Bestätigung der Einreichung", metadata=[("Einreichung", "ECU-GC-2026-0626-184431"), ("Zeitpunkt", "26.06.2026, 18:44:31 Uhr Luxemburger Zeit"), ("Einreicher", "Dr. Elin Nørgaard für Lipsia Speichertechnik AG"), ("Dokumentart", "Klage gemäß Artikel 263 AEUV")], sections=[("1. Übertragene Dateien", ["Die Übertragung umfasste die Klageschrift mit 48 Seiten, das Anlagenverzeichnis und vierzehn Anlagen. Die technische Prüfung hat sämtliche Dateien angenommen. Der für die Klageschrift berechnete SHA-256-Wert beginnt mit 8F2C7A91D4E6. Die vollständige Prüfsumme ist im elektronischen Konto abrufbar."]), ("2. Vorläufige Registrierung", ["Die Einreichung wurde unter dem vorläufigen Zeichen T-PRE-18842/26 erfasst. Diese Bestätigung besagt, dass Dateien technisch eingegangen sind. Sie enthält keine Entscheidung über Zulässigkeit, Fristwahrung, Vollmacht, Verfahrenssprache oder die Ordnungsmäßigkeit der Anlagen."]), ("3. Zustellungen", ["Weitere Nachrichten der Kanzlei werden im e-Curia-Konto der Vertreter bereitgestellt. Der Zeitpunkt der tatsächlichen oder vermuteten Zustellung wird dort für jedes Dokument gesondert ausgewiesen."])], signer="Automatisch erzeugte Eingangsbestätigung", signer_role="Kanzlei des Gerichts")
    write_letter(d / "06_kanzleischreiben_nachbesserung.docx", org="Gericht der Europäischen Union", strapline="Kanzlei · Direkte Klagen", contact="Rue du Fort Niedergrünewald · L-2925 Luxemburg\nTelefon +352 4303 1 · RegistryGC@curia.europa.eu", recipient=["Dr. Elin Nørgaard", "Westphal · Dierkes · Mohn Partnerschaft mbB", "Taunusanlage 18", "60325 Frankfurt am Main"], date="03.07.2026", reference="T-412/26 · Lipsia Speichertechnik/Kommission", subject="Klageeingang – zu behebende formale Mängel", salutation="Sehr geehrte Frau Kollegin,", sections=[("1. Vollmacht und Registerunterlage", ["Die hochgeladene Vollmacht trägt zwei eingescannte Unterschriften, lässt jedoch die Vertretungsbefugnis des zweiten Vorstandsmitglieds nicht aus sich heraus erkennen. Bitte reichen Sie einen aktuellen Handelsregisterauszug und eine lesbare Fassung des Vorstandsbeschlusses ein."]), ("2. Anlagen", ["Anlage A.9 wurde in der Klageschrift auf Seite 31 zitiert, ist im Anlagenverzeichnis jedoch als A.8 bezeichnet. Für Anlage A.12 liegt eine vertrauliche Fassung vor; die hochgeladene nicht vertrauliche Fassung enthält auf Seite 6 weiterhin Namen und Einzelpreise. Bitte berichtigen Sie Nummerierung und Fassung."]), ("3. Frist", ["Die genannten Punkte sind bis zum 13.07.2026 über e-Curia zu beheben. Dieses Schreiben greift einer Entscheidung über die Zulässigkeit der Klage nicht vor."])], closing="Mit vorzüglicher Hochachtung", signer="L. Ferreira", signer_role="Kanzler der Dritten Kammer")
    write_email(d / "07_email_fristsorge_kanzlei.eml", sender="Dr. Felix Westphal <f.westphal@westphal-counsel.de>", recipients=["Dr. Elin Nørgaard <e.norgaard@westphal-counsel.de>"], cc=["Dr. Edda Noll <edda.noll@lipsia-speicher.de>"], date="Fri, 19 Jun 2026 21:17:00 +0200", subject="Lipsia / SA.118742 – bitte Fristrechnung morgen doppelt prüfen", message_id="<wdm.61-2026.20260619.211700@westphal-counsel.de>", body="""Elin,

im Vorstandskalender steht 26.06. als Sicherheitsfrist. Meine handschriftliche Rechnung kommt bei Zustellung am 17.04. auf den 27.06.; das ist ein Samstag. Bitte prüfe anhand der aktuellen Verfahrensordnung, wie Monatsfrist, Entfernungsfrist und Wochenende zusammenspielen. Wir reichen auf keinen Fall erst am letzten rechnerischen Tag ein.

Noch offen: Handelsregisterauszug nicht älter als drei Monate, saubere Vollmacht, nicht vertrauliche A.12 und die Frage, ob die deutsche Fassung der Entscheidung vollständig ist. Der Mandant möchte am Montag um 08:00 Uhr einen belastbaren Stand.

Felix""", attachments=["fristnotiz_fw_scan.jpg"])
    write_email(d / "08_email_stadt_zinsdaten.eml", sender="Dr. Amelie Hartung <amelie.hartung@leipzig.de>", recipients=["Dr. Edda Noll <edda.noll@lipsia-speicher.de>"], cc=["Referat Europa <europa@leipzig.de>"], date="Wed, 13 May 2026 11:24:00 +0200", subject="VII-20.31/LSA-2026-17 – aktualisierte Zinsdatei", message_id="<lsa-2026-17.20260513.112400@leipzig.de>", body="""Sehr geehrte Frau Dr. Noll,

in der am 04.05. übersandten Tabelle war die Tilgung vom 30.09.2024 in der Spalte Restvaluta versehentlich erst zum Folgequartal berücksichtigt. Anbei erhalten Sie Version 2. Der vorläufige Gesamtbetrag zum 30.04.2026 reduziert sich dadurch von 2743513.06 EUR auf 2582004.34 EUR.

Die Stellungnahmefrist 25.05.2026 bleibt bestehen. Bitte verwenden Sie für Ihre Rückmeldung nur die neue Datei und geben Sie an, falls Ihre Bankunterlagen von den in Spalte D genannten Valuten abweichen.

Mit freundlichen Grüßen
Dr. Amelie Hartung
Stadt Leipzig · Referat Beteiligungen
Telefon 0341 123 4192""", attachments=["rueckforderung_berechnung_v2_2026-05-13.csv"])
    write_txt(d / "09_chat_e_curia_upload.txt", """Kanzleichat · Matter 61-2026 · 26.06.2026

16:08 Elin Nørgaard: Klageschrift final 48 Seiten. Bitte keine weiteren Änderungen ohne Rückruf.
16:11 Miriam Vogelsang: Anlagen A.1 bis A.14 sind im Uploadordner. A.12 hat zwei Fassungen, vertraulich und geschwärzt.
16:44 Felix Westphal: Ist A.9 im Verzeichnis richtig? Im Schriftsatz Seite 31 steht Gutachten Harenberg.
16:47 Miriam Vogelsang: Im Verzeichnis heißt das Gutachten gerade A.8. Ich prüfe die Nummern gegen die Dateinamen.
17:22 Elin Nørgaard: Nummerierung im Schriftsatz bleibt. Bitte Verzeichnis korrigieren und beide A.12-Fassungen erneut exportieren.
18:31 Miriam Vogelsang: Upload vollständig, 16 Dateien. e-Curia zeigt 187 MB und keine Fehlermeldung.
18:44 Elin Nørgaard: Eingereicht. Bestätigung ECU-GC-2026-0626-184431 ist da.
18:51 Felix Westphal: Bitte Quittung, Hash und vollständigen Uploadordner schreibgeschützt ablegen. Am Montag Registerauszug aktualisieren.""")
    write_csv(d / "10_rueckforderungsberechnung_stadt.csv", ["Periode", "Restvaluta_EUR", "Marktpreis_Prozent", "Gezahlt_Prozent", "Differenz_EUR", "Zinsen_EUR", "Version"], [["2023-Q3", "38000000.00", "3.10", "0.65", "232750.00", "18941.22", "v2 gemäß Stadt"], ["2023-Q4", "38000000.00", "3.10", "0.65", "232750.00", "21880.04", "v2 gemäß Stadt"], ["2024-Q1", "36500000.00", "3.25", "0.65", "237250.00", "20612.18", "v2 gemäß Stadt"], ["2024-Q2", "35000000.00", "3.25", "0.65", "227500.00", "18440.77", "v2 gemäß Stadt"], ["2024-Q3", "33500000.00", "3.40", "0.65", "230312.50", "16382.02", "v2 gemäß Stadt"], ["2024-Q4", "31500000.00", "3.40", "0.65", "216562.50", "13409.91", "v2 gemäß Stadt"], ["2025", "27000000.00", "3.60", "0.65", "796500.00", "48711.63", "v2 gemäß Stadt"], ["2026-bis-04-30", "23500000.00", "3.75", "0.65", "242708.33", "7293.24", "v2 gemäß Stadt"]])
    write_csv(d / "11_anlagen_und_vertraulichkeit.csv", ["Anlage", "Dateiname", "Seiten", "Vertraulich", "Nicht_vertrauliche_Fassung", "Offener_Punkt"], [["A.1", "vollmacht_2026-06-22.pdf", 3, "nein", "entfällt", "Registervertretung prüfen"], ["A.2", "handelsregister_2026-03-11.pdf", 5, "nein", "entfällt", "älter als drei Monate"], ["A.3", "C_2026_2417_de.pdf", 84, "teilweise", "ja", "authentische Fassung"], ["A.8", "bankangebote_2023.pdf", 22, "ja", "ja", "Nummerierung kollidiert"], ["A.9", "gutachten_harenberg_2025.pdf", 41, "ja", "ja", "im Verzeichnis als A.8"], ["A.12", "kundenpreise_delitzsch.xlsx", 9, "ja", "ja", "Seite 6 noch unvollständig geschwärzt"], ["A.14", "zinsberechnung_stadt_v2.csv", 8, "nein", "entfällt", "Betrag weicht von Entscheidung ab"]])
    write_device_photo(d / "12_foto_fristskizze.jpg", "Handschriftliche Fristnotiz", ["Zustellung: Freitag, 17.04.2026", "+ 2 Monate: 17.06.2026?", "+ 10 Tage: 27.06.2026", "27.06. ist Samstag", "Interne Abgabe: Freitag, 26.06. · 18 Uhr"], (36, 64, 95))
    write_readme(d, "Europäisches Prozessrecht: Beihilfe-Nichtigkeitsklage aus Leipzig", ["`europaeisches-prozessrecht`"], "Die Lipsia Speichertechnik AG erhält eine Kommissionsentscheidung zur Rückforderung einer kommunalen Garantie. Zwischen Zustellung, städtischer Anhörung, drohender Liquiditätsbelastung und Klage beim Gericht der Europäischen Union entstehen Frist-, Vollmachts-, Vertraulichkeits- und Anlagenprobleme. Eine technisch erfolgreiche e-Curia-Einreichung wird später formal beanstandet.", [("01_zustellung_kommissionsentscheidung.docx", "Zustellschreiben der Kommission mit Entscheidungs- und Veröffentlichungsbezug"), ("02_vorstandsprotokoll_sondersitzung.docx", "Vorstandsprotokoll zu Liquidität, Kommunikation und Verfahrensauftrag"), ("03_aufforderung_rueckforderung_stadt_leipzig.docx", "Anhörung der Stadt mit vorläufiger Rückforderungsberechnung"), ("04_klageschrift_fassung_24_juni.docx", "Ausformulierte Klagefassung mit Anträgen und Anlagenbezug"), ("05_e_curia_eingangsbestaetigung.docx", "Technische Eingangsbestätigung der Einreichung"), ("06_kanzleischreiben_nachbesserung.docx", "Gerichtliches Schreiben zu Vollmacht, Register und Anlagen"), ("07_email_fristsorge_kanzlei.eml", "Partner-E-Mail zur Fristberechnung"), ("08_email_stadt_zinsdaten.eml", "Berichtigte Zinsdatei und geänderter Rückforderungsbetrag"), ("09_chat_e_curia_upload.txt", "Zeitnaher Upload-Chat mit Nummerierungsproblemen"), ("10_rueckforderungsberechnung_stadt.csv", "Quartalsbezogene Berechnungsdaten"), ("11_anlagen_und_vertraulichkeit.csv", "Anlagenstatus mit vertraulichen und offenen Fassungen"), ("12_foto_fristskizze.jpg", "Mobiltelefonfoto der handschriftlichen Fristskizze")], "Der Fall trainiert Klageart, Betroffenheit, Fristberechnung, einstweiligen Rechtsschutz, e-Curia, Verfahrenssprache, Vollmacht, Anlagenordnung und vertrauliche Fassungen. Er enthält Parteivortrag, aber keine gerichtliche Lösung.")
    write_rubric(d, slug, "europaeisches-prozessrecht", "Direkte Nichtigkeitsklage eines Beihilfebegünstigten mit e-Curia, Frist, Vollmacht, Vertraulichkeit und möglichem Eilrechtsschutz.", [("frist", "Zustellung am 17.04.2026, Zweimonatsfrist nach Artikel 263 Abs. 6 AEUV, Berechnung nach Artikel 58 und Entfernungsfrist nach Artikel 60 der Verfahrensordnung des Gerichts sowie das Wochenende sauber zusammenführen."), ("klagebefugnis", "Adressat der Rückforderungswirkung, unmittelbare und individuelle Betroffenheit sowie Rechtsschutzbedürfnis getrennt prüfen."), ("e-curia", "Technische Eingangsbestätigung nicht mit gerichtlicher Zulässigkeitsprüfung verwechseln und Nachbesserung fristgerecht organisieren."), ("vertraulichkeit", "Anlagenverzeichnis, vertrauliche Fassung und nicht vertrauliche Fassung datei- und seitenbezogen abstimmen."), ("eilrechtsschutz", "Klage und möglichen Antrag nach Artikeln 278 und 279 AEUV einschließlich Dringlichkeit und Interessenabwägung getrennt strukturieren.")])


def build_treuepflicht() -> None:
    slug = "gesellschaftsrecht-treuepflicht-kapitalerhoehung-ulm"
    d = prepare(slug)
    company = "Brenner Präzisionsguss GmbH"
    contact = "Magirusstraße 28 · 89077 Ulm\nTelefon 0731 509 60 · gesellschaft@brenner-guss.de\nHRB 728441 Amtsgericht Ulm · Stammkapital 500000 EUR"
    write_record(d / "01_gesellschaftsvertrag_lesefassung.docx", org=company, strapline="Gesellschaftsunterlagen · Lesefassung vom 12.01.2022", contact=contact, title="Gesellschaftsvertrag", metadata=[("Notarielle Urschrift", "UR-Nr. 41/2022 Notarin Dr. Cora Mertens, Ulm"), ("Register", "HRB 728441 Amtsgericht Ulm"), ("Stammkapital", "500000 EUR"), ("Geschäftsjahr", "Kalenderjahr")], sections=[("1. Firma und Sitz", ["Die Gesellschaft führt die Firma Brenner Präzisionsguss GmbH und hat ihren Sitz in Ulm. Gegenstand ist die Entwicklung, Herstellung und der Vertrieb von Präzisionsgussteilen sowie die Erbringung damit verbundener technischer Dienstleistungen."]), ("2. Geschäftsanteile", ["Vom Stammkapital halten Viktor Brenner einen Geschäftsanteil von 200000 EUR, Dr. Anna Brenner einen Geschäftsanteil von 175000 EUR und die Familienkontor Brenner KG einen Geschäftsanteil von 125000 EUR. Jeder Euro eines Geschäftsanteils gewährt eine Stimme."]), ("3. Gesellschafterbeschlüsse", ["Beschlüsse werden grundsätzlich mit einfacher Mehrheit der abgegebenen Stimmen gefasst. Änderungen des Gesellschaftsvertrags, Kapitalmaßnahmen, Veräußerung wesentlicher Betriebsteile und Rechtsgeschäfte mit Gesellschaftern oder ihnen nahestehenden Unternehmen bedürfen einer Mehrheit von 75 Prozent der abgegebenen Stimmen. Gesetzliche Stimmverbote bleiben unberührt."]), ("4. Informationsrechte", ["Jeder Gesellschafter kann während der Geschäftszeiten Einsicht in Bücher und Schriften verlangen. Bei laufenden Verhandlungen mit Banken darf die Geschäftsführung besonders sensible Unterlagen zunächst in einem gesicherten Datenraum bereitstellen."]), ("5. Wettbewerbs- und Nachfolgeregelung", ["Gesellschafter dürfen ohne Zustimmung keine unmittelbar konkurrierende Tätigkeit aufnehmen. Übertragungen an Abkömmlinge sind nach schriftlicher Anzeige zulässig; im Übrigen besteht ein Vorerwerbsrecht der Mitgesellschafter."])], signer="Dr. Cora Mertens", signer_role="Notarin · beglaubigte Lesefassung")
    write_record(d / "02_gesellschafterliste_registerstand.docx", org="Notariat Mertens & Reuß", strapline="Notarinnen in Ulm · Registervollzug", contact="Neue Straße 42 · 89073 Ulm\nTelefon 0731 880 44 0 · notarinnen@mertens-reuss.de", title="Gesellschafterliste zum 15.03.2026", metadata=[("Gesellschaft", company), ("Register", "HRB 728441 Amtsgericht Ulm"), ("Liste", "eingereicht am 15.03.2026"), ("Gesamt", "500000 EUR")], sections=[("1. Beteiligungen", ["Die nachfolgende Liste gibt den bei Erstellung mitgeteilten Beteiligungsstand wieder. Veränderungen seit der zuletzt aufgenommenen Liste: keine."])], tables=[("2. Gesellschafter", ["Nr.", "Gesellschafter", "Nennbetrag", "Anteil"], [["1", "Viktor Brenner, Ulm", "200000 EUR", "40 Prozent"], ["2", "Dr. Anna Brenner, Neu-Ulm", "175000 EUR", "35 Prozent"], ["3", "Familienkontor Brenner KG, Blaubeuren", "125000 EUR", "25 Prozent"]], [1.4, 7.2, 3.4, 3.5])], signer="Dr. Cora Mertens", signer_role="Notarin")
    write_letter(d / "03_einladung_gesellschafterversammlung.docx", org=company, strapline="Geschäftsführung", contact=contact, recipient=["Dr. Anna Brenner", "Silcherstraße 19", "89231 Neu-Ulm"], date="06.05.2026", reference="VB/CS 05-26", subject="Einladung zur Gesellschafterversammlung am 21.05.2026 um 14:00 Uhr", salutation="Sehr geehrte Frau Dr. Brenner,", sections=[("1. Ort und Teilnahme", ["Die Versammlung findet im Verwaltungsgebäude, Magirusstraße 28, Besprechungsraum Gießerei, statt. Eine Teilnahme per Videokonferenz wird zusätzlich angeboten. Bitte teilen Sie bis zum 18.05. mit, ob Sie persönlich oder durch Bevollmächtigten teilnehmen."]), ("2. Tagesordnung", ["Tagesordnungspunkt 1: Bericht zur Liquiditäts- und Auftragslage. Tagesordnungspunkt 2: Erhöhung des Stammkapitals von 500000 EUR auf 2000000 EUR durch Ausgabe neuer Geschäftsanteile. Tagesordnungspunkt 3: Zulassung der Familienkontor Brenner KG zur Übernahme der neuen Geschäftsanteile; Bezugsfrist für übrige Gesellschafter bis 27.05.2026. Tagesordnungspunkt 4: Abschluss eines Lager- und Logistikvertrags mit der VB Industriepark GmbH. Tagesordnungspunkt 5: Erweiterung der Zustimmungsvorbehalte für die Geschäftsführung."]), ("3. Unterlagen", ["Im Datenraum stehen der Liquiditätsplan vom 04.05.2026, das Schreiben der Alb-Donau Bank, der Entwurf des Übernahmebeschlusses und der Entwurf des Lagervertrags. Das Verkehrswertgutachten für Halle 3 wird nach Eingang ergänzt."])], closing="Mit freundlichen Grüßen", signer="Viktor Brenner", signer_role="Geschäftsführer und Gesellschafter", attachments=["Beschlussentwürfe", "Datenraumindex", "Vollmachtsformular"])
    write_letter(d / "04_bankschreiben_eigenkapitalauflage.docx", org="Alb-Donau Bank AG", strapline="Firmenkunden Süd · Kreditmanagement", contact="Olgastraße 83 · 89073 Ulm\nTelefon 0731 141 0 · firmenkunden@alb-donau-bank.de\nBIC ALBDDE6UXXX", recipient=[company, "Herrn Viktor Brenner", "Magirusstraße 28", "89077 Ulm"], date="30.04.2026", reference="K 4188-77 / CS", subject="Konsortiallinie 2024/188 – Liquiditätsauflage zum 31.05.2026", salutation="Sehr geehrter Herr Brenner,", sections=[("1. Auslastung", ["Die Betriebsmittellinie von 6000000 EUR war am 28.04.2026 mit 5748810.44 EUR in Anspruch genommen. Die im Vertrag vereinbarte Mindestliquidität wurde in den Kalenderwochen 14 bis 17 dreimal unterschritten. Zugleich ist der Zahlungseingang des Kunden NordRail Systems GmbH über 2180000 EUR weiterhin offen."]), ("2. Auflage", ["Das Konsortium erwartet bis zum 31.05.2026 einen nachgewiesenen Eigenkapitalzufluss von mindestens 1500000 EUR oder eine wirtschaftlich gleichwertige, nachrangige Finanzierung. Ohne fristgerechten Nachweis wird die weitere Ziehung unter dem Investitionsrahmen für die Formanlage G-7 ausgesetzt."]), ("3. Unterlagen", ["Bitte reichen Sie Gesellschafterbeschluss, Übernahmeerklärungen, Zahlungsnachweise und aktualisierte Zwölf-Wochen-Planung ein. Eine Finanzierung durch ein nahestehendes Unternehmen wird nur berücksichtigt, wenn Konditionen, Rang und Rückzahlung transparent dokumentiert sind."])], closing="Mit freundlichen Grüßen", signer="Carolin Sauter", signer_role="Direktorin Firmenkunden", attachments=["Covenant-Auswertung April 2026"])
    write_record(d / "05_protokoll_gesellschafterversammlung.docx", org=company, strapline="Gesellschafterversammlung · Niederschrift", contact=contact, title="Niederschrift vom 21.05.2026", metadata=[("Beginn", "14:03 Uhr"), ("Ende", "17:18 Uhr"), ("Ort", "Magirusstraße 28, Ulm"), ("Protokoll", "Rechtsanwältin Lea Konrad")], sections=[("1. Anwesenheit", ["Viktor Brenner ist persönlich anwesend. Für die Familienkontor Brenner KG nimmt Komplementärin Helene Brenner teil. Dr. Anna Brenner ist mit Rechtsanwalt Dr. Malte Rupp anwesend. Sämtliche Geschäftsanteile sind vertreten. Notarin Dr. Cora Mertens nimmt ab 15:35 Uhr für die Beurkundung der Kapitalmaßnahme teil."]), ("2. Bericht zur Lage", ["Viktor Brenner erläutert den offenen NordRail-Zahlungseingang und die Kreditauflage. Dr. Anna Brenner verlangt die Debitorenliste, den Vertrag mit NordRail, die Korrespondenz mit den Banken und eine Erläuterung, warum die Kapitalerhöhung innerhalb von sechs Tagen vollständig eingezahlt werden müsse. Ihr wird Einsicht in den Datenraum angeboten; die Debitorenliste sei dort nur ohne Einzelkundenpreise vorhanden."]), ("3. Kapitalerhöhung", ["Viktor Brenner und Helene Brenner stimmen für die Erhöhung des Stammkapitals auf 2000000 EUR. Dr. Anna Brenner stimmt dagegen und erklärt Widerspruch zur Niederschrift. Sie beantragt eine Bezugsfrist bis 30.06.2026 und die Zulassung einer anteiligen Übernahme. Der Versammlungsleiter stellt fest, die Familienkontor Brenner KG sei zur Übernahme der gesamten Erhöhung zugelassen. Über den Verlängerungsantrag wird nach kurzer Unterbrechung nicht gesondert abgestimmt."]), ("4. Lagervertrag", ["Viktor Brenner erklärt vor der Abstimmung, alleiniger Gesellschafter der VB Industriepark GmbH zu sein. Dr. Rupp beantragt, sein Stimmrecht bei Tagesordnungspunkt 4 nicht zu berücksichtigen. Der Versammlungsleiter lässt Viktor Brenner abstimmen und stellt 65 Prozent Zustimmung sowie 35 Prozent Gegenstimmen fest. Das Verkehrswertgutachten liegt noch nicht vor; der Vertragsentwurf nennt eine Jahresmiete von 486000 EUR und zehn Jahre feste Laufzeit."]), ("5. Nachbemerkungen", ["Dr. Anna Brenner verlangt eine Abschrift der Niederschrift, sämtliche Abstimmungsergebnisse und die Beurkundungsunterlagen. Um 17:11 Uhr erklärt Notarin Dr. Mertens, die Anmeldung werde erst nach Eingang der Übernahmeerklärung und der Geschäftsführerbestätigung vorbereitet."])], signer="Lea Konrad", signer_role="Protokollführerin")
    write_letter(d / "06_notariat_vollzugshinweis.docx", org="Notariat Mertens & Reuß", strapline="Notarinnen in Ulm · Registervollzug", contact="Neue Straße 42 · 89073 Ulm\nTelefon 0731 880 44 0 · notarinnen@mertens-reuss.de", recipient=[company, "Geschäftsführung", "Magirusstraße 28", "89077 Ulm"], date="27.05.2026", reference="CM/kr UR 388/2026", subject="Kapitalerhöhung vom 21.05.2026 – fehlende Vollzugsunterlagen", salutation="Sehr geehrter Herr Brenner,", sections=[("1. Beurkundung", ["Die Niederschrift der Gesellschafterversammlung und der Kapitalerhöhungsbeschluss wurden unter UR-Nr. 388/2026 beurkundet. Die Familienkontor Brenner KG hat die Übernahme eines neuen Geschäftsanteils von 1500000 EUR erklärt."]), ("2. Fehlende Unterlagen", ["Für die Anmeldung fehlen der Banknachweis über die Einzahlung, die Versicherung der Geschäftsführung zur freien Verfügung und eine abschließende Gesellschafterliste. Zudem liegt uns das Schreiben von Rechtsanwalt Dr. Rupp vom 26.05.2026 vor, mit dem Einwendungen gegen Beschlussfassung und Bezugsfrist erhoben werden."]), ("3. Registeranmeldung", ["Wir haben die Anmeldung noch nicht an das Handelsregister übermittelt. Bitte teilen Sie bis zum 01.06.2026 mit, ob die Gesellschaft trotz des angekündigten gerichtlichen Antrags den Vollzug wünscht. Wir beraten keine Seite zu den materiellen Erfolgsaussichten des Gesellschafterstreits."])], closing="Mit freundlichen Grüßen", signer="Dr. Cora Mertens", signer_role="Notarin")
    write_letter(d / "07_antrag_einstweilige_verfuegung.docx", org="Rupp · Eberlein Rechtsanwälte", strapline="Gesellschaftsrecht · Prozessführung", contact="Bahnhofstraße 17 · 89073 Ulm\nTelefon 0731 602 18 0 · kanzlei@rupp-eberlein.de\nAnderkonto IBAN DE72 6305 0000 0088 4112 19", recipient=["Landgericht Ulm", "Kammer für Handelssachen", "Olgastraße 106", "89073 Ulm"], date="29.05.2026", reference="MR/AB 118/26", subject="Dr. Anna Brenner gegen Brenner Präzisionsguss GmbH – Antrag auf Erlass einer einstweiligen Verfügung", salutation="Namens und in Vollmacht der Antragstellerin beantragen wir,", sections=[("1. Anträge", ["der Antragsgegnerin bis zu einer Entscheidung in der Hauptsache zu untersagen, die am 21.05.2026 beschlossene Kapitalerhöhung zum Handelsregister anzumelden oder weiterzuverfolgen; hilfsweise, der Antragsgegnerin aufzugeben, eine Anmeldung nur unter Mitteilung des anhängigen Streits vorzunehmen."]), ("2. Beteiligungen und Beschluss", ["Die Antragstellerin hält 35 Prozent des Stammkapitals. Die übrigen Stimmen entfallen auf Viktor Brenner und die von seiner Mutter vertretene Familienkontor Brenner KG. Der Beschluss erhöht das Stammkapital um 1500000 EUR und weist die gesamte Übernahme der Familienkontor Brenner KG zu. Die Antragstellerin erhielt sechs Tage Zeit, ohne dass ihr ein konkretes anteiliges Bezugsangebot unterbreitet wurde."]), ("3. Informationslage", ["Trotz wiederholter Nachfrage waren im Datenraum weder die vollständige Debitorenliste noch der NordRail-Vertrag oder ein Unternehmenswert verfügbar. Der Bankbrief verlangt einen Eigenkapitalzufluss, schreibt aber weder Übernehmer noch gesellschaftsrechtliche Ausgestaltung vor. Die Antragstellerin hat am 22.05.2026 schriftlich erklärt, bis zu 300000 EUR kurzfristig und weitere Mittel nach Prüfung aufzubringen."]), ("4. Dringlichkeit", ["Das Notariat hat die Anmeldung zwar noch nicht eingereicht, fordert aber eine Entscheidung der Gesellschaft bis zum 01.06.2026. Nach Eintragung und Einzahlung würden sich Beteiligungsquoten und Einflussmöglichkeiten der Antragstellerin erheblich verändern. Eine Schutzschrift der Gesellschaft ist der Antragstellerin nicht bekannt."])], closing="Dr. Malte Rupp", signer="Rechtsanwalt", signer_role="Fachanwalt für Handels- und Gesellschaftsrecht", attachments=["Gesellschaftsvertrag", "Einladung", "Bankbrief", "Protokoll", "E-Mail-Verkehr", "eidesstattliche Versicherung Dr. Anna Brenner"])
    write_email(d / "08_email_anna_brenner_finanzierungsangebot.eml", sender="Dr. Anna Brenner <anna.brenner@ab-materialtechnik.de>", recipients=["Viktor Brenner <v.brenner@brenner-guss.de>"], cc=["Helene Brenner <h.brenner@familienkontor-brenner.de>", "Dr. Malte Rupp <m.rupp@rupp-eberlein.de>"], date="Fri, 22 May 2026 09:14:00 +0200", subject="Kapitalerhöhung – mein anteiliges Angebot und fehlende Unterlagen", message_id="<ab.20260522.091400.kapital@ab-materialtechnik.de>", body="""Viktor,

ich halte meine Einwände aus der gestrigen Versammlung aufrecht. Unabhängig davon bin ich bereit, kurzfristig 300000 EUR einzuzahlen. Für einen darüber hinausgehenden Betrag brauche ich bis spätestens Dienstag die vollständige Debitorenliste, den NordRail-Vertrag, die Bankkorrespondenz und die Herleitung des Ausgabebetrags.

Bitte bestätige, auf welchen neuen Geschäftsanteil ich zeichnen kann und ob dieselben Bedingungen gelten wie für das Familienkontor. Die Frist bis 27.05. ist ohne diese Angaben nicht ernsthaft nutzbar. Ich bin am Montag ab 11 Uhr telefonisch erreichbar.

Anna""")
    write_email(d / "09_email_geschaeftsfuehrer_an_notariat.eml", sender="Viktor Brenner <v.brenner@brenner-guss.de>", recipients=["Dr. Cora Mertens <c.mertens@mertens-reuss.de>"], cc=["Helene Brenner <h.brenner@familienkontor-brenner.de>"], date="Thu, 28 May 2026 18:37:00 +0200", subject="UR 388/2026 – bitte Anmeldung vorbereiten", message_id="<vb.20260528.183700.ur388@brenner-guss.de>", body="""Sehr geehrte Frau Dr. Mertens,

die Einzahlung des Familienkontors ist heute um 16:21 Uhr auf dem Kapitalerhöhungskonto eingegangen. Der Kontoauszug folgt morgen früh. Bitte bereiten Sie die Anmeldung vor. Nach unserer Auffassung ist das Anwaltsschreiben meiner Schwester kein Grund, den Vollzug aufzuhalten.

Die Bank erwartet den Registervollzug nicht ausdrücklich bis 31.05., aber sie verlangt einen belastbaren Eigenkapitalnachweis. Ohne Freigabe der nächsten Kredittranche können wir die Formanlage G-7 nicht abnehmen. Bitte rufen Sie mich morgen vor 09:00 Uhr an, falls Sie noch eine Erklärung benötigen.

Mit freundlichen Grüßen
Viktor Brenner
Geschäftsführer""", attachments=["zahlungsavis_familienkontor_1500000.pdf"])
    write_txt(d / "10_chat_geschaeftsleitung_lagerhalle.txt", """Auszug aus dem Leitungschat · Kanal #halle3 · 06.05. bis 20.05.2026

06.05. 18:02 Viktor Brenner: Gutachten für Halle 3 kommt laut Sachverständigem erst Ende nächster Woche.
06.05. 18:11 CFO Peter Renz: Wir brauchen für die Bank vorher eine belastbare Mietzahl. Im Modell stehen 40500 EUR pro Monat.
07.05. 07:46 Viktor Brenner: Bitte beim Entwurf bleiben. VB Industriepark hat den Umbau vorfinanziert.
12.05. 16:28 Peter Renz: Vergleichsmiete aus Blaubeuren liegt eher bei 31800 bis 35000 EUR. Ich habe die Mail in den Datenraum gelegt.
12.05. 16:41 Viktor Brenner: Nicht in den allgemeinen Ordner. Erst mit dem Gutachten zusammen einstellen, sonst gibt es wieder eine Grundsatzdiskussion.
20.05. 19:13 Peter Renz: Gutachten ist noch nicht da. Soll TOP 4 trotzdem abgestimmt werden?
20.05. 19:19 Viktor Brenner: Ja. Wir können im Vertrag eine Anpassung nach Gutachten vorsehen. Der aktuelle Entwurf hat die Klausel aber noch nicht.""")
    write_csv(d / "11_liquiditaetsplan_12_wochen.csv", ["Woche", "Anfang_EUR", "Einzahlungen_EUR", "Auszahlungen_EUR", "Kredit_EUR", "Ende_EUR", "Annahme"], [["2026-05-04", "612400", "988000", "1215000", "0", "385400", "NordRail offen"], ["2026-05-11", "385400", "712000", "1068000", "0", "29400", "Kurzarbeit nicht beantragt"], ["2026-05-18", "29400", "634000", "912000", "250000", "1400", "Banktranche teilweise"], ["2026-05-25", "1400", "901000", "1186000", "0", "-283600", "ohne Kapitalzufluss"], ["2026-06-01", "-283600", "2780000", "1325000", "0", "1171400", "NordRail Zahlung vollständig"], ["2026-06-08", "1171400", "822000", "1044000", "0", "949400", "Normalbetrieb"], ["2026-06-15", "949400", "785000", "1299000", "0", "435400", "Anzahlung Formanlage"], ["2026-06-22", "435400", "866000", "945000", "0", "356400", ""], ["2026-06-29", "356400", "914000", "880000", "0", "390400", ""], ["2026-07-06", "390400", "801000", "967000", "0", "224400", ""], ["2026-07-13", "224400", "854000", "923000", "0", "155400", ""], ["2026-07-20", "155400", "889000", "901000", "0", "143400", ""]])
    write_csv(d / "12_kapitalangebote_und_stimmrechte.csv", ["Beteiligter", "Altanteil_EUR", "Altquote_Prozent", "Angebot_neu_EUR", "Zahlungsdatum", "Neue_Quote_wenn_angenommen", "Bemerkung"], [["Viktor Brenner", "200000", "40", "0", "", "10", "keine eigene Übernahme"], ["Dr. Anna Brenner", "175000", "35", "300000", "offen", "23.75", "weitere Mittel nach Unterlagenprüfung"], ["Familienkontor Brenner KG", "125000", "25", "1500000", "2026-05-28", "81.25", "gesamte Erhöhung übernommen"], ["Gesamt", "500000", "100", "1800000", "", "", "Angebote übersteigen beschlossene Erhöhung um 300000 EUR"]])
    write_device_photo(d / "13_foto_whiteboard_kapitalrunde.jpg", "Besprechungsraum Gießerei", ["Bankauflage: 1.5 Mio. bis 31.05.", "Familienkontor: 1.5 Mio. sofort", "Anna: 0.3 Mio. + Prüfung", "NordRail: 2.18 Mio. offen", "Halle 3: Gutachten fehlt"], (111, 64, 45))
    write_readme(d, "Gesellschaftsrechtliche Treuepflicht: Kapitalerhöhung in Ulm", ["`gesellschaftsrechtliche-treuepflicht`"], "Eine Familien-GmbH gerät wegen eines großen offenen Kundenzahlungseingangs unter Bankdruck. Mehrheit und Familienkontor beschließen eine Kapitalerhöhung, die die Minderheitsgesellschafterin stark verwässern würde. Gleichzeitig stimmt der Geschäftsführer über einen langfristigen Lagervertrag mit seiner eigenen Immobiliengesellschaft ab.", [("01_gesellschaftsvertrag_lesefassung.docx", "Gesellschaftsvertrag mit Mehrheiten, Informationsrechten und Related-Party-Klausel"), ("02_gesellschafterliste_registerstand.docx", "Notarielle Gesellschafterliste"), ("03_einladung_gesellschafterversammlung.docx", "Einladung mit Kapitalerhöhung und Lagervertrag"), ("04_bankschreiben_eigenkapitalauflage.docx", "Bankauflage und Kreditstatus"), ("05_protokoll_gesellschafterversammlung.docx", "Ausführliche Niederschrift mit Anträgen, Widerspruch und Abstimmungen"), ("06_notariat_vollzugshinweis.docx", "Notarieller Zwischenstand vor Registeranmeldung"), ("07_antrag_einstweilige_verfuegung.docx", "Ausformulierter Eilantrag der Minderheitsgesellschafterin"), ("08_email_anna_brenner_finanzierungsangebot.eml", "Anteiliges Finanzierungsangebot und Unterlagenverlangen"), ("09_email_geschaeftsfuehrer_an_notariat.eml", "Vollzugswunsch nach Zahlungseingang"), ("10_chat_geschaeftsleitung_lagerhalle.txt", "Interner Chat zum noch fehlenden Verkehrswertgutachten"), ("11_liquiditaetsplan_12_wochen.csv", "Zwölf-Wochen-Plan mit und ohne erwarteten Kundenzahlungseingang"), ("12_kapitalangebote_und_stimmrechte.csv", "Beteiligungen, Angebote und mögliche Verwässerung"), ("13_foto_whiteboard_kapitalrunde.jpg", "Mobiltelefonfoto der internen Eckdaten")], "Der Vorgang verbindet Informationsrechte, Mehrheitsmacht, Bezugsrecht, Stimmverbot, Related-Party-Geschäft, Registervollzug und einstweiligen Rechtsschutz. Die Primärstücke lassen unterschiedliche Deutungen zu.")
    write_rubric(d, slug, "gesellschaftsrechtliche-treuepflicht", "Familien-GmbH zwischen Bankauflage, selektiver Kapitalerhöhung, Verwässerung, Related-Party-Vertrag und Registervollzug.", [("satzung", "Satzungsmehrheiten, gesetzliche Anforderungen der Kapitalerhöhung und tatsächliche Abstimmung getrennt erfassen."), ("treuepflicht", "Sanierungsbedarf, Gleichbehandlung, Bezugschance, Information und weniger belastende Alternativen in einer fallbezogenen Treuepflichtprüfung abwägen."), ("stimmverbot", "Die Abstimmung über den Vertrag mit der vom Geschäftsführer beherrschten VB Industriepark GmbH eigenständig nach Paragraf 47 Abs. 4 GmbHG prüfen."), ("eilschutz", "Registerstand, Notariatsschreiben, Zahlungsnachweis und drohende Eintragung für Antrag, Glaubhaftmachung und Dringlichkeit zusammenführen."), ("zahlen", "Liquiditätsplan und Kapitalangebote rechnerisch aufeinander beziehen, ohne den NordRail-Eingang als sicher zu behandeln.")])


def build_wirtschaftsprueferrecht() -> None:
    slug = "wirtschaftsprueferrecht-unabhaengigkeit-hamburg"
    d = prepare(slug)
    firm = "Hanseatische Revision Partnerschaft mbB"
    firm_contact = "Ballindamm 39 · 20095 Hamburg\nTelefon 040 380 91 0 · Fax 040 380 91 199\npost@hanseatische-revision.de · PR 1188 AG Hamburg"
    client = "Nordwerft Energieanlagen AG"
    client_contact = "Australiastraße 51 · 20457 Hamburg\nTelefon 040 771 82 0 · governance@nordwerft-energie.de\nHRB 188241 AG Hamburg · ISIN DE000NW00018"
    write_letter(d / "01_pruefungsauftrag_2025.docx", org=firm, strapline="Wirtschaftsprüfungsgesellschaft · Niederlassung Hamburg", contact=firm_contact, recipient=[client, "z. Hd. Vorsitzende des Prüfungsausschusses", "Australiastraße 51", "20457 Hamburg"], date="18.07.2025", reference="HRP/NWE-AP-2025", subject="Auftragsbestätigung zur Prüfung des Jahres- und Konzernabschlusses 2025", salutation="Sehr geehrte Frau Dr. Mahler,", sections=[("1. Gegenstand", ["Nach Wahl durch die Hauptversammlung und Auftrag des Aufsichtsrats prüfen wir Jahresabschluss, Lagebericht, Konzernabschluss und Konzernlagebericht zum 31.12.2025. Verantwortlicher Prüfungspartner ist Wirtschaftsprüfer Dr. Henrik Clausen. Das Kernteam umfasst 18 Mitarbeiter an den Standorten Hamburg und Bremen."]), ("2. Honorar", ["Das voraussichtliche Prüfungshonorar beträgt 420000 EUR zuzüglich gesetzlicher Umsatzsteuer und Auslagen. Wesentliche Erweiterungen werden vor Ausführung mit dem Prüfungsausschuss abgestimmt. Leistungen anderer Netzwerkgesellschaften werden gesondert beauftragt und abgerechnet."]), ("3. Unabhängigkeit", ["Wir bestätigen vorbehaltlich der fortlaufenden Überwachung, dass uns bei Auftragsannahme keine Umstände bekannt sind, die unsere Unabhängigkeit ausschließen. Der Prüfungsausschuss erhält vor Prüfungsbeginn eine Aufstellung sämtlicher im Netzwerk erbrachter Leistungen. Neue Nichtprüfungsleistungen bedürfen der vorgesehenen Freigabe."]), ("4. Kommunikation", ["Feststellungen von besonderer Bedeutung werden unverzüglich an die Vorsitzende des Prüfungsausschusses berichtet. Der abschließende Prüfungsbericht und die Unabhängigkeitserklärung werden dem Aufsichtsrat vorgelegt."])], closing="Mit freundlichen Grüßen", signer="Dr. Henrik Clausen · WP StB", signer_role="Partner · Verantwortlicher Abschlussprüfer", attachments=["Allgemeine Auftragsbedingungen", "Team- und Zeitplan", "Datenschutz- und Portalhinweise"])
    write_record(d / "02_unabhaengigkeitserklaerung_team.docx", org=firm, strapline="Qualitätsmanagement · Independence Office", contact=firm_contact, title="Unabhängigkeitserklärung und Konfliktabfrage", metadata=[("Mandant", client), ("Prüfungszeitraum", "Geschäftsjahr 2025"), ("Abfrage geschlossen", "29.08.2025"), ("Fall-ID", "IND-25-4471")], sections=[("1. Netzwerkabfrage", ["Die elektronische Konfliktabfrage wurde an 184 Partner und Direktoren des deutschen Netzwerks versandt. 181 Antworten gingen bis zum Stichtag ein; drei Antworten wurden am 02.09.2025 nachgeholt. Die Netzwerkgesellschaft Hanseatic Advisory GmbH meldete ein Projekt zur Einführung des Konsolidierungssystems OrbisCon bei der Nordwerft-Tochter Baltic Rotor Services GmbH."]), ("2. Finanzielle Interessen", ["Für die Mitglieder des Prüfungsteams ergab die Depotabfrage keine unmittelbar gehaltenen Aktien der Nordwerft Energieanlagen AG. Prüfungsmanagerin Lena Brandt meldete am 28.08.2025, dass ihr Ehepartner seit 01.06.2025 Leiter Treasury der Baltic Rotor Services GmbH ist. Sie wurde zunächst aus dem Teilteam Finanzierung genommen; die Systemberechtigung blieb bis 12.09.2025 aktiv."]), ("3. Leistungen", ["Neben der Abschlussprüfung waren Steuer-Compliance-Leistungen, eine prüfungsnahe Bestätigung zu Fördermitteln und das OrbisCon-Projekt gemeldet. Für die Steuerleistungen lag eine Freigabe des Prüfungsausschusses vom 22.05.2025 vor. Für zwei Erweiterungsmodule des OrbisCon-Projekts enthielt der Datenraum lediglich eine E-Mail des Finanzvorstands an den Ausschussvorsitzenden."]), ("4. Schlussvermerk", ["Das Independence Office bat am 03.09.2025 um Nachweise zu Leistungsumfang, Entscheidungsverantwortung des Mandanten und Freigabe der Erweiterungsmodule. Der Auftragsstatus blieb bis zur Antwort auf gelb gesetzt. Am 15.09.2025 wurde er nach Vorlage eines Management-Responsibility-Memos auf grün geändert."])], signer="Mareike Oltmann", signer_role="Director Independence & Ethics")
    write_record(d / "03_pruefungsausschuss_protokoll.docx", org=client, strapline="Aufsichtsrat · Prüfungsausschuss", contact=client_contact, title="Niederschrift der 18. Sitzung des Prüfungsausschusses", metadata=[("Datum", "22.05.2025, 09:00 bis 11:46 Uhr"), ("Ort", "Hamburg, Verwaltungsgebäude Speicherstadt"), ("Vorsitz", "Dr. Hilde Mahler"), ("Protokoll", "Svenja Karthaus")], sections=[("1. Abschlussprüfung 2025", ["Der Ausschuss empfiehlt dem Aufsichtsrat, die Hanseatische Revision Partnerschaft mbB nach der Wahl durch die Hauptversammlung zu beauftragen. Dr. Clausen stellt Team, Zeitplan und Schwerpunkte vor. Der Ausschuss verlangt einen gesonderten Bericht zu Projektkalkulationen, Covenants der Anleihe und dem Konsolidierungsprozess der Auslandstöchter."]), ("2. Steuer-Compliance", ["Der Ausschuss genehmigt für 2025 wiederkehrende Steuer-Compliance-Leistungen bis zu einem Honorar von 140000 EUR. Die Leistungen umfassen Steuererklärungen und die Abstimmung latenter Steuern. Gestaltungsberatung und Vertretung in wesentlichen Streitverfahren sind nicht umfasst."]), ("3. Konsolidierungssystem", ["Finanzvorstand Dr. Martin Veer berichtet über die Einführung von OrbisCon durch Hanseatic Advisory GmbH bei Baltic Rotor Services. Dr. Mahler fragt, ob das Projekt die spätere Prüfung eigener Arbeiten berührt. Dr. Clausen erklärt, die Tochtergesellschaft treffe sämtliche Konfigurations- und Freigabeentscheidungen; das Prüfungsteam werde die Implementierung nicht übernehmen. Der Ausschuss bittet um eine schriftliche Abgrenzung vor Produktivsetzung."]), ("4. Weitere Beratung", ["Eine gesonderte Beschlussfassung zu Erweiterungsmodulen findet in der Sitzung nicht statt. Dr. Veer soll bei wesentlichen Änderungen erneut berichten."])], signer="Dr. Hilde Mahler", signer_role="Vorsitzende des Prüfungsausschusses")
    write_letter(d / "04_ankuenigung_apas_inspektion.docx", org="Abschlussprüferaufsichtsstelle beim Bundesamt für Wirtschaft und Ausfuhrkontrolle", strapline="Inspektionen bei Abschlussprüfern von Unternehmen von öffentlichem Interesse", contact="Uhlandstraße 88-90 · 10717 Berlin\nTelefon 030 590 099 0 · poststelle@apasbafa.bund.de", recipient=[firm, "Geschäftsführung", "Ballindamm 39", "20095 Hamburg"], date="08.04.2026", reference="APAS-I 4-26/0318", subject="Inspektion 2026 – Anforderung von Unterlagen und Auswahl einer Prüfungsakte", salutation="Sehr geehrte Damen und Herren,", sections=[("1. Inspektion", ["Die Abschlussprüferaufsichtsstelle führt bei Ihrer Praxis eine Inspektion durch. Der Vor-Ort-Termin ist für den 11. bis 15.05.2026 in Hamburg vorgesehen. Inspektionsleiterin ist Regierungsdirektorin Dr. Paula Kern."]), ("2. Ausgewählte Abschlussprüfung", ["Für die auftragsbezogene Prüfung wurde unter anderem die gesetzliche Abschlussprüfung der Nordwerft Energieanlagen AG zum 31.12.2025 ausgewählt. Bitte stellen Sie bis zum 24.04.2026 den Prüfungsbericht, die Prüfungsakte, Unabhängigkeitsabfragen, Honorarübersichten, Freigaben von Nichtprüfungsleistungen und die Kommunikation mit dem Prüfungsausschuss bereit."]), ("3. Zugang", ["Der Datenraum ist so einzurichten, dass Version, Einstellzeitpunkt und nachträgliche Änderungen erkennbar bleiben. Benennen Sie bis zum 17.04.2026 einen organisatorischen Ansprechpartner. Die Auswahl sagt nichts über das Ergebnis der Inspektion aus."])], closing="Mit freundlichen Grüßen", signer="Dr. Paula Kern", signer_role="Inspektionsleiterin")
    write_letter(d / "05_zwischenmitteilung_apas.docx", org="Abschlussprüferaufsichtsstelle beim Bundesamt für Wirtschaft und Ausfuhrkontrolle", strapline="Inspektionen bei Abschlussprüfern von Unternehmen von öffentlichem Interesse", contact="Uhlandstraße 88-90 · 10717 Berlin\nTelefon 030 590 099 0 · poststelle@apasbafa.bund.de", recipient=[firm, "z. Hd. Dr. Henrik Clausen", "Ballindamm 39", "20095 Hamburg"], date="28.05.2026", reference="APAS-I 4-26/0318", subject="Vorläufige Feststellungen – Nordwerft Energieanlagen AG", salutation="Sehr geehrter Herr Dr. Clausen,", sections=[("1. Unabhängigkeitsüberwachung", ["Die Inspektion hat festgestellt, dass die Meldung zur Beschäftigung des Ehepartners der Prüfungsmanagerin am 28.08.2025 einging, ihr Portalzugang zum Prüfungsbereich Finanzierung jedoch erst am 12.09.2025 entzogen wurde. Bitte erläutern Sie, welche Arbeitsschritte sie in diesem Zeitraum tatsächlich ausführte und wie deren Ergebnisse überprüft wurden."]), ("2. Nichtprüfungsleistungen", ["Für die Erweiterungsmodule des OrbisCon-Projekts liegen Rechnungen über insgesamt 310000 EUR und eine E-Mail des Finanzvorstands vor. Ein gesonderter Ausschussbeschluss war im bereitgestellten Datenraum nicht auffindbar. Bitte legen Sie dar, auf welche Freigabe und welche Abgrenzung der Entscheidungsverantwortung sich die Praxis stützte."]), ("3. Prüfungsdokumentation", ["Mehrere Arbeitspapiere wurden nach dem Datum des Bestätigungsvermerks geändert. Das Systemprotokoll weist Änderungen am 19.03., 02.04. und 07.04.2026 aus. Benennen Sie Anlass, Inhalt, Autor und ursprüngliche Fassung jeder Änderung. Ihre Stellungnahme erwarten wir bis zum 18.06.2026."])], closing="Mit freundlichen Grüßen", signer="Dr. Paula Kern", signer_role="Inspektionsleiterin", attachments=["Liste der erbetenen Einzelnachweise"])
    write_record(d / "06_interne_nachtragsdokumentation.docx", org=firm, strapline="Qualitätsmanagement · Prüfungsakte Nordwerft 2025", contact=firm_contact, title="Dokumentation nach Abschluss der Prüfungsakte", metadata=[("Prüfungsbericht", "erteilt am 11.03.2026"), ("Archivschluss", "10.05.2026"), ("Erstellt", "09.04.2026"), ("Verantwortlich", "Dr. Henrik Clausen")], sections=[("1. Änderung vom 19.03.2026", ["Senior Manager Torben Lask ergänzte im Arbeitspapier FIN-4.7 den Querverweis auf die Bankbestätigung vom 28.02.2026. Zahlen und Prüfungsschlussfolgerung blieben unverändert. Die frühere PDF-Fassung liegt im Versionsarchiv."]), ("2. Änderung vom 02.04.2026", ["Nach einer Rückfrage des Prüfungsausschusses wurde dem Arbeitspapier IT-2.3 die schriftliche Abgrenzung der Managementverantwortung beim OrbisCon-Projekt beigefügt. Das Memo trägt das Datum 12.09.2025, wurde nach Angaben von Hanseatic Advisory jedoch erst am 31.03.2026 aus deren Projektarchiv exportiert."]), ("3. Änderung vom 07.04.2026", ["Im Arbeitspapier IND-3.2 wurde der Zeitraum der aktiven Systemberechtigung von Prüfungsmanagerin Lena Brandt von 28.08. bis 12.09.2025 ergänzt. Die Zugriffsprotokolle zeigen in diesem Zeitraum 17 Dokumentaufrufe und zwei Kommentierungen. Der Engagement Quality Reviewer erhielt am 08.04.2026 eine Kopie."]), ("4. Sicherung", ["Alle Änderungen sind in der Prüfungssoftware mit Nutzer, Uhrzeit und Begründung gespeichert. Die Arbeitspapiere wurden nicht gelöscht. Die ursprünglichen Fassungen sind im schreibgeschützten Export NWE-2025-FINAL-110326 enthalten."])], signer="Dr. Henrik Clausen", signer_role="Verantwortlicher Abschlussprüfer")
    write_email(d / "07_email_hinweis_pruefungsmanagerin.eml", sender="Lena Brandt <l.brandt@hanseatische-revision.de>", recipients=["Mareike Oltmann <m.oltmann@hanseatische-revision.de>"], cc=["Dr. Henrik Clausen <h.clausen@hanseatische-revision.de>"], date="Thu, 28 Aug 2025 17:48:00 +0200", subject="Nordwerft – Beschäftigung meines Mannes bei Baltic Rotor", message_id="<ind-25-4471.20250828.174800@hanseatische-revision.de>", body="""Hallo Frau Oltmann,

mein Mann, Jens Brandt, arbeitet seit 01.06. als Leiter Treasury bei Baltic Rotor Services GmbH. Mir war nicht bewusst, dass die Gesellschaft dieses Jahr erstmals in den Nordwerft-Prüfungsumfang aufgenommen wird. Ich habe bisher am Konzernplanungstool und an der Bankbestätigungssteuerung gearbeitet, nicht an Baltic Rotor selbst.

Bitte sagen Sie mir, ob ich aus dem Engagement genommen werde. Mein Zugriff auf den allgemeinen Nordwerft-Datenraum ist noch aktiv. Morgen früh wäre eine Abstimmung hilfreich, weil ich um 10 Uhr den Finanzierungscall moderiere.

Viele Grüße
Lena Brandt
Senior Manager Audit"""
    )
    write_email(d / "08_email_advisory_erweiterungsmodule.eml", sender="René Döring <r.doering@hanseatic-advisory.de>", recipients=["Dr. Martin Veer <martin.veer@nordwerft-energie.de>"], cc=["Dr. Henrik Clausen <h.clausen@hanseatische-revision.de>"], date="Tue, 02 Sep 2025 13:22:00 +0200", subject="OrbisCon – Freigabe Module Intercompany und Cash Forecast", message_id="<orbiscon.20250902.132200@hanseatic-advisory.de>", body="""Sehr geehrter Herr Dr. Veer,

wie im Lenkungskreis besprochen, können wir die beiden Erweiterungsmodule für 155000 EUR je Modul bis Ende November produktiv setzen. Baltic Rotor bleibt für Kontenplan, Freigaberegeln und Abnahme verantwortlich. Unser Team konfiguriert die von Ihnen freigegebenen Parameter und dokumentiert die Tests.

Sie hatten im Call gesagt, Frau Dr. Mahler sei grundsätzlich informiert. Für unsere Akte benötigen wir noch die formale Freigabe des Prüfungsausschusses oder eine Bestätigung, dass der Beschluss vom 22.05. diese Erweiterung umfasst. Bitte senden Sie uns das bis zum Kick-off am 08.09.

Mit freundlichen Grüßen
René Döring
Partner · Hanseatic Advisory GmbH""", attachments=["angebot_orbiscon_module_v4.pdf"])
    write_txt(d / "09_chat_pruefungsakte_nachtraege.txt", """Interner Teams-Export · Kanal Nordwerft-2025 · 07.04.2026

08:12 Torben Lask: APAS-Auswahl ist da. Bitte heute keine Dateien im Engagement verändern, bis Henrik den Export freigibt.
08:18 Dr. Henrik Clausen: Richtig. Nur fehlende Nachtragsdokumentation, keine Änderung von Prüfungsschlüssen.
09:03 Lena Brandt: Mein alter Zugang zeigt noch FIN-4.7 und IND-3.2. Soll ich mich abmelden oder braucht ihr einen Screenshot?
09:05 Mareike Oltmann: Bitte nichts öffnen. IT zieht das Zugriffslog zentral.
10:44 Torben Lask: Log zeigt 17 Aufrufe zwischen 28.08. und 12.09.2025, zwei Kommentare in der Bankbestätigungsliste. Keine Freigabehandlung.
11:02 Dr. Henrik Clausen: Das muss vollständig in den Nachtragsvermerk, einschließlich Datum der heutigen Ermittlung und alter Version.
15:37 Mareike Oltmann: Beim OrbisCon-Memo steht 12.09.2025 im Dokument, Export aus Advisory aber 31.03.2026. Bitte Herkunft nicht als zeitnahen Aktenbestand darstellen.""")
    write_csv(d / "10_honorare_netzwerk_2023_2026.csv", ["Jahr", "Einheit", "Leistung", "Honorar_EUR", "Freigabe", "Rechnung", "Bemerkung"], [[2023, firm, "Abschlussprüfung", "360000", "HV und Aufsichtsrat", "HR-230418", ""], [2023, "Hanseatic Tax GmbH", "Steuer-Compliance", "118000", "PA 2023-04", "HT-231177", ""], [2024, firm, "Abschlussprüfung", "390000", "HV und Aufsichtsrat", "HR-240522", ""], [2024, "Hanseatic Tax GmbH", "Steuer-Compliance", "124000", "PA 2024-06", "HT-241884", ""], [2025, firm, "Abschlussprüfung", "420000", "HV und Aufsichtsrat", "HR-260311", "erteilt 11.03.2026"], [2025, "Hanseatic Tax GmbH", "Steuer-Compliance", "128000", "PA 22.05.2025", "HT-251991", ""], [2025, "Hanseatic Advisory GmbH", "OrbisCon Basismodul", "185000", "PA Information 22.05.2025", "HA-250881", "Managemententscheidungen laut Memo"], [2025, "Hanseatic Advisory GmbH", "OrbisCon Intercompany", "155000", "E-Mail Finanzvorstand", "HA-251104", "kein gesonderter Beschluss im Datenraum"], [2025, "Hanseatic Advisory GmbH", "OrbisCon Cash Forecast", "155000", "E-Mail Finanzvorstand", "HA-251105", "kein gesonderter Beschluss im Datenraum"], [2026, firm, "Prüfungsnahe Fördermittelbestätigung", "85000", "PA 15.01.2026", "HR-260417", "separates Team"]])
    write_csv(d / "11_zugriffs_und_aenderungslog.csv", ["Zeitpunkt", "Nutzer", "Objekt", "Aktion", "Version", "Begründung_im_System"], [["2025-08-29 08:07", "Lena Brandt", "FIN-4.7", "geöffnet", "v12", ""], ["2025-09-03 14:42", "Lena Brandt", "Bankbestätigungen", "Kommentar", "v18", "Rücklauf Elbe Bank fehlt"], ["2025-09-11 09:16", "Lena Brandt", "FIN-4.7", "Kommentar", "v19", "Abgleich Covenant-Spalte"], ["2025-09-12 16:05", "System Admin", "Mandat Nordwerft", "Berechtigung entzogen", "", "Independence Ticket IND-25-4471"], ["2026-03-11 18:22", "Dr. Henrik Clausen", "Prüfungsakte", "Archivmarke", "FINAL", "Bestätigungsvermerk erteilt"], ["2026-03-19 10:08", "Torben Lask", "FIN-4.7", "Querverweis ergänzt", "v22", "Bankbestätigung verknüpft"], ["2026-04-02 12:31", "Torben Lask", "IT-2.3", "Anlage ergänzt", "v17", "Management-Responsibility-Memo"], ["2026-04-07 11:24", "Dr. Henrik Clausen", "IND-3.2", "Sachverhalt ergänzt", "v14", "Zugriffszeitraum Lena Brandt"]])
    write_device_photo(d / "12_foto_datenraum_status.jpg", "Inspection Room APAS-I 4-26/0318", ["Prüfungsakte NWE 2025: 4821 Dateien", "Final Export: 11.03.2026 · 18:22", "Nachträge: 19.03. / 02.04. / 07.04.", "Offene Anfrage: Independence", "Zugriff APAS: bereit ab 24.04.2026"], (74, 86, 99))
    write_readme(d, "Wirtschaftsprüferrecht: Unabhängigkeit und APAS-Inspektion in Hamburg", ["`berufsrecht-wirtschaftspruefer`"], "Eine Wirtschaftsprüfungspraxis prüft den Abschluss eines kapitalmarktorientierten Energieanlagenbauers. Parallel erbringt das Netzwerk Steuer- und Systemleistungen. Eine familiäre Beziehung im Prüfungsteam, verspätet entzogene Zugriffsrechte, unklare Ausschussfreigaben und Änderungen nach dem Bestätigungsvermerk geraten in den Fokus einer APAS-Inspektion.", [("01_pruefungsauftrag_2025.docx", "Auftragsbestätigung mit Honorar, Team und Unabhängigkeitsklausel"), ("02_unabhaengigkeitserklaerung_team.docx", "Konfliktabfrage mit Netzwerkleistungen und familiärer Beziehung"), ("03_pruefungsausschuss_protokoll.docx", "Ausschussprotokoll zu Prüfung und Nichtprüfungsleistungen"), ("04_ankuenigung_apas_inspektion.docx", "Förmliche Inspektionsankündigung und Aktenauswahl"), ("05_zwischenmitteilung_apas.docx", "Vorläufige Feststellungen und Stellungnahmefrist"), ("06_interne_nachtragsdokumentation.docx", "Dokumentation der Änderungen nach Erteilung des Bestätigungsvermerks"), ("07_email_hinweis_pruefungsmanagerin.eml", "Zeitnahe Meldung einer familiären Beziehung"), ("08_email_advisory_erweiterungsmodule.eml", "E-Mail zu Umfang und Freigabe zweier Systemmodule"), ("09_chat_pruefungsakte_nachtraege.txt", "Interner Chat am Tag der Aktenauswahl"), ("10_honorare_netzwerk_2023_2026.csv", "Mehrjährige Honorar- und Leistungsdaten des Netzwerks"), ("11_zugriffs_und_aenderungslog.csv", "Zeitgestempelter Zugriffs- und Versionsverlauf"), ("12_foto_datenraum_status.jpg", "Mobiltelefonfoto der Datenraum-Statusanzeige")], "Der Dokumentensatz ermöglicht eine quellennahe Prüfung von Unabhängigkeit, Nichtprüfungsleistungen, Ausschussfreigabe, Netzwerkzurechnung, Prüfungsdokumentation und Aufsichtsverfahren. Die Feststellungen bleiben zwischen Praxis, Mandant und Aufsicht streitig.")
    write_rubric(d, slug, "berufsrecht-wirtschaftspruefer", "APAS-Inspektion einer PIE-Abschlussprüfung mit Netzwerkleistungen, familiärer Beziehung, Honorarverlauf und Nachtragsdokumentation.", [("zustaendigkeit", "PIE-Eigenschaft, APAS-Zuständigkeit und Inspektionsgegenstand aus den Unterlagen herausarbeiten."), ("unabhaengigkeit", "Paragrafen 319 und 319b HGB, WPO und Verordnung (EU) Nr. 537/2014 auf finanzielle, persönliche und leistungsbezogene Gefährdungen beziehen."), ("honorare", "Prüfungs- und Nichtprüfungshonorare nach Jahr und Netzwerkgesellschaft trennen; eine mögliche Obergrenze nur auf zutreffender Datenbasis berechnen."), ("freigaben", "Ausschussbeschluss, bloße Information und E-Mail des Finanzvorstands nicht gleichsetzen."), ("nachtraege", "Nachträgliche Dokumentation von unzulässiger rückwirkender Aktenänderung anhand des Versionslogs und des Vermerks unterscheiden.")])


def write_readme(d: Path, title: str, plugins: list[str], summary: str, files: list[tuple[str, str]], closing: str) -> None:
    lines = [f"# {title}", "", f"Arbeitsakte zu den Plugins {', '.join(plugins)}.", "", "## Kurzbild", "", summary, "", "## Aktenstücke", "", "| Datei | Inhalt |", "| --- | --- |"]
    lines.extend(f"| `{name}` | {description} |" for name, description in files)
    lines.extend(["", "## Bearbeitung", "", closing, ""])
    (d / "README.md").write_text("\n".join(lines), encoding="utf-8")


def yaml_quote(value: str) -> str:
    return '"' + value.replace("\\", "\\\\").replace('"', '\\"') + '"'


def write_rubric(d: Path, slug: str, plugin: str, description: str, checks: list[tuple[str, str]]) -> None:
    lines = [f"name: {slug}", f"plugin: {yaml_quote(plugin)}", f"description: {yaml_quote(description)}", "checks:", "  - id: dateien", "    check_type: working_file_count", "    description: Mindestens zwölf exportierbare Primärunterlagen vorhanden", "    min: 12"]
    for check_id, text in checks:
        lines.extend([f"  - id: {check_id}", "    check_type: human_review", f"    description: {yaml_quote(text)}"])
    (d / "rubric.yaml").write_text("\n".join(lines) + "\n", encoding="utf-8")


def main() -> int:
    build_apothekenrecht()
    build_krankenhausrecht()
    build_eu_prozessrecht()
    build_treuepflicht()
    build_wirtschaftsprueferrecht()
    print("Fünf priorisierte Schulungsakten mit 61 Primärunterlagen erzeugt.")
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
