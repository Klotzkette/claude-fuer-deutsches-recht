#!/usr/bin/env python3
"""Konvertiert Markdown-Aktenstuecke einer Testakte in lebensechte DOCX-Dateien.

Grundregel (siehe testakten/QUALITAETSSTANDARD.md): Aktenstuecke liegen in
Formaten vor, wie sie im Anwaltsleben anfallen (DOCX, PDF, XLSX, EML, JPG,
TXT-Chatexporte, CSV) - NICHT als Markdown. Markdown bleibt nur fuer
README, rubric und Meta-/Loesungsdateien, die ohnehin nicht in den
Akten-ZIP-Export gelangen (testakte_file_filter).

- Jedes nicht-Meta-.md wird zu einem formatierten DOCX (Times New Roman 11 pt,
  A4, 2.5 cm Raender, echte Word-Ueberschriften und -Tabellen) und geloescht.
- Existiert bereits ein gleichnamiges .docx (Zwilling), wird das .md nur
  entfernt.
- Danach die Gesamt-PDFs neu bauen (build-testakte-gesamt-pdf.py) - der
  Builder liest DOCX nativ.

Aufruf: convert-testakte-aktenstuecke-nativ.py [slug ...]   (ohne Args: alle)
"""
from __future__ import annotations
import re, sys
from pathlib import Path

sys.path.insert(0, str(Path(__file__).resolve().parent))
from testakte_file_filter import is_export_meta_file

from docx import Document
from docx.shared import Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH

REPO = Path(__file__).resolve().parent.parent
ROOT = REPO / "testakten"
SKIP_DIRS = {"megaprompts", "formatvorlagen-paradebeispiele"}

INLINE = re.compile(r"(\*\*[^*]+\*\*|\*[^*\n]+\*|`[^`]+`)")

def add_runs(par, text: str):
    text = re.sub(r"\[([^\]]+)\]\(([^)]+)\)", r"\1", text)  # Links: nur Text
    pos = 0
    for m in INLINE.finditer(text):
        if m.start() > pos:
            par.add_run(text[pos:m.start()])
        tok = m.group(0)
        if tok.startswith("**"):
            r = par.add_run(tok[2:-2]); r.bold = True
        elif tok.startswith("`"):
            par.add_run(tok[1:-1])
        else:
            r = par.add_run(tok[1:-1]); r.italic = True
        pos = m.end()
    if pos < len(text):
        par.add_run(text[pos:])

def setup_doc() -> Document:
    doc = Document()
    sec = doc.sections[0]
    sec.page_width, sec.page_height = Cm(21.0), Cm(29.7)
    for a in ("top_margin", "bottom_margin", "left_margin", "right_margin"):
        setattr(sec, a, Cm(2.5))
    st = doc.styles["Normal"]; st.font.name = "Times New Roman"; st.font.size = Pt(11)
    for name, size in (("Heading 1", 14), ("Heading 2", 12), ("Heading 3", 11)):
        h = doc.styles[name]; h.font.name = "Times New Roman"; h.font.size = Pt(size)
        h.font.bold = True; h.font.color.rgb = None
    return doc

def flush_table(doc, rows):
    rows = [r for r in rows if not re.match(r"^\s*\|?[\s:|-]+\|?\s*$", r)]
    if not rows: return
    cells = [[c.strip() for c in r.strip().strip("|").split("|")] for r in rows]
    ncol = max(len(r) for r in cells)
    t = doc.add_table(rows=len(cells), cols=ncol); t.style = "Table Grid"
    for i, row in enumerate(cells):
        for j in range(ncol):
            txt = row[j] if j < len(row) else ""
            p = t.cell(i, j).paragraphs[0]
            add_runs(p, txt)
            for r in p.runs:
                r.font.name = "Times New Roman"; r.font.size = Pt(10)
                if i == 0: r.bold = True

def md_to_docx(src: Path, dst: Path):
    doc = setup_doc()
    lines = src.read_text(encoding="utf-8", errors="ignore").splitlines()
    table: list[str] = []
    for raw in lines:
        line = raw.rstrip()
        if line.strip().startswith("|"):
            table.append(line); continue
        if table:
            flush_table(doc, table); table = []
        s = line.strip()
        if not s:
            continue
        if s.startswith("<!--"):
            continue
        m = re.match(r"^(#{1,6})\s+(.*)$", s)
        if m:
            lvl = min(len(m.group(1)), 3)
            p = doc.add_heading("", level=lvl); add_runs(p, m.group(2))
            for r in p.runs: r.font.name = "Times New Roman"
            continue
        if re.match(r"^([-*_]\s*){3,}$", s):
            doc.add_paragraph("")  # Trennlinie -> Leerabsatz
            continue
        m = re.match(r"^[-*+]\s+(.*)$", s)
        if m:
            p = doc.add_paragraph(style="List Bullet"); add_runs(p, m.group(1)); continue
        m = re.match(r"^(\d+)[.)]\s+(.*)$", s)
        if m:
            p = doc.add_paragraph(style="List Number"); add_runs(p, m.group(2)); continue
        if s.startswith(">"):
            p = doc.add_paragraph(); p.paragraph_format.left_indent = Cm(1.0)
            add_runs(p, s.lstrip("> ").strip())
            for r in p.runs: r.italic = True
            continue
        p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        add_runs(p, s)
    if table:
        flush_table(doc, table)
    doc.save(dst)

def convert_akte(d: Path) -> tuple[int, int]:
    converted = removed_twin = 0
    for p in sorted(d.rglob("*.md")):
        if "gesamt-pdf" in p.parts: continue
        if is_export_meta_file(p, d): continue
        dst = p.with_suffix(".docx")
        if dst.exists():
            p.unlink(); removed_twin += 1; continue
        md_to_docx(p, dst); p.unlink(); converted += 1
    return converted, removed_twin

def main() -> int:
    slugs = sys.argv[1:]
    dirs = ([ROOT / s for s in slugs] if slugs
            else [d for d in sorted(ROOT.iterdir()) if d.is_dir() and d.name not in SKIP_DIRS])
    tc = tr = 0
    for d in dirs:
        c, r = convert_akte(d)
        tc += c; tr += r
        if c or r:
            print(f"  {d.name}: {c} konvertiert, {r} Zwillings-md entfernt")
    print(f"Fertig: {tc} DOCX erzeugt, {tr} md-Zwillinge entfernt.")
    return 0

if __name__ == "__main__":
    raise SystemExit(main())
