#!/usr/bin/env python3
"""Baut für jede Testakte ein 'gesamt-pdf/<name>_gesamt.pdf', das die
exportfaehigen Aktenstücke (MD/TXT/EML/CSV/XLSX/DOCX/Bilder/PDF) in ein
einziges, sauber gerendertes Dokument mit Dateigrenzen und Seitenzahlen
zusammenfasst.

Aufruf:
  python3 scripts/build-testakte-gesamt-pdf.py                 # alle Testakten
  python3 scripts/build-testakte-gesamt-pdf.py <name1> <name2>  # gezielt
"""

from __future__ import annotations

import io
import re
import sys
import csv
import json
import tempfile
import xml.etree.ElementTree as ET
from copy import deepcopy
from email import policy
from email.parser import BytesParser
from html import unescape as html_unescape
from html.parser import HTMLParser
from pathlib import Path

import yaml

# Drittabhaengigkeiten
from openpyxl import load_workbook
from pypdf import PageObject, PdfReader, PdfWriter, Transformation
from reportlab.lib.pagesizes import A4
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib.units import cm
from reportlab.lib.colors import HexColor, black
from reportlab.platypus import (
    SimpleDocTemplate,
    Image as RLImage,
    Paragraph,
    Spacer,
    PageBreak,
    Table,
    TableStyle,
    HRFlowable,
    KeepTogether,
)
from reportlab.lib.utils import ImageReader
from reportlab.lib.enums import TA_LEFT
from reportlab.pdfgen import canvas
from reportlab.pdfbase import pdfmetrics
from reportlab.pdfbase.ttfonts import TTFont

from testakte_file_filter import include_in_working_dump
from testakte_office_pdf import OFFICE_EXTS, OfficeRenderError, render_office_batch

# DOCX
try:
    from docx import Document
except ImportError:
    Document = None  # type: ignore

REPO_ROOT = Path(__file__).resolve().parent.parent
TESTAKTEN = REPO_ROOT / "testakten"

# Design
TEAL = HexColor("#01696F")
MUTED = HexColor("#7A7974")
BORDER = HexColor("#D4D1CA")
SURFACE = HexColor("#F7F6F2")

# Font: System-Helvetica als Fallback. Inter waere schoener, aber wir verzichten
# auf Netzwerk-Downloads, damit das Skript offline laeuft.
FONT_REG = "Helvetica"
FONT_BOLD = "Helvetica-Bold"
A4_PORTRAIT = (float(A4[0]), float(A4[1]))


def invariant_canvas(*args, **kwargs):
    """Erzeugt PDFs ohne laufzeitabhaengige IDs und Zeitstempel."""
    kwargs["invariant"] = 1
    return canvas.Canvas(*args, **kwargs)

styles = getSampleStyleSheet()
s_cover_label = ParagraphStyle(
    "CoverLabel",
    fontName=FONT_REG, fontSize=14, leading=18,
    textColor=MUTED, spaceAfter=6,
)
s_cover_title = ParagraphStyle(
    "CoverTitle",
    fontName=FONT_BOLD, fontSize=28, leading=34,
    textColor=TEAL, alignment=TA_LEFT, spaceAfter=14,
)
s_cover_sub = ParagraphStyle(
    "CoverSub",
    fontName=FONT_REG, fontSize=12, leading=16,
    textColor=black, spaceAfter=4,
)
s_cover_meta = ParagraphStyle(
    "CoverMeta",
    fontName=FONT_REG, fontSize=9, leading=12,
    textColor=MUTED, spaceAfter=3,
)
s_h1 = ParagraphStyle(
    "H1", parent=styles["Heading1"],
    fontName=FONT_BOLD, fontSize=18, leading=22, textColor=TEAL,
    spaceBefore=18, spaceAfter=8,
)
s_h2 = ParagraphStyle(
    "H2", parent=styles["Heading2"],
    fontName=FONT_BOLD, fontSize=14, leading=18, textColor=black,
    spaceBefore=12, spaceAfter=6,
)
s_h3 = ParagraphStyle(
    "H3", parent=styles["Heading3"],
    fontName=FONT_BOLD, fontSize=11, leading=14, textColor=black,
    spaceBefore=8, spaceAfter=4,
)
s_body = ParagraphStyle(
    "Body", parent=styles["BodyText"],
    fontName=FONT_REG, fontSize=10, leading=14, textColor=black,
    spaceAfter=6,
)
s_meta = ParagraphStyle(
    "Meta", parent=styles["BodyText"],
    fontName=FONT_REG, fontSize=9, leading=12, textColor=MUTED,
    spaceAfter=4,
)
s_partlabel = ParagraphStyle(
    "PartLabel", parent=styles["BodyText"],
    fontName=FONT_BOLD, fontSize=11, leading=14, textColor=MUTED,
    spaceAfter=2,
)
# Kopf jedes Aktenstuecks im Gesamt-PDF: Kategorie-Kicker, Dateititel, Pfad.
s_file_kicker = ParagraphStyle(
    "FileKicker",
    fontName=FONT_BOLD, fontSize=8, leading=10, textColor=MUTED, spaceAfter=1,
)
s_file_title = ParagraphStyle(
    "FileTitle",
    fontName=FONT_BOLD, fontSize=13, leading=16, textColor=TEAL, spaceAfter=1,
)
s_file_path = ParagraphStyle(
    "FilePath",
    fontName=FONT_REG, fontSize=8, leading=10, textColor=MUTED, spaceAfter=4,
)

# Reihenfolge der Datei-Typen im Gesamt-PDF
TYPE_ORDER = [
    "md",
    "txt",
    "eml",
    "csv",
    "xlsx",
    "structured",
    "docx",
    "odt",
    "image",
    "pdf",
]
IMAGE_EXTS = {"jpg", "jpeg", "png"}
STRUCTURED_EXTS = {"json", "yaml", "yml", "xml", "ics", "abc"}
TYPE_LABEL = {
    "md": "Aktenstücke",
    "txt": "Notizen und Textdateien",
    "eml": "E-Mails",
    "csv": "CSV-Tabellen",
    "xlsx": "Excel-Tabellen",
    "structured": "Strukturierte Rohdaten und Kalenderdateien",
    "docx": "Word-Dokumente",
    "odt": "OpenDocument-Textdateien",
    "image": "Bildanlagen und Screenshots",
    "pdf": "PDF-Anhänge (Originaldokumente)",
}


class DocumentRenderError(RuntimeError):
    """Eine Arbeitsunterlage konnte nicht vollständig in PDF überführt werden."""


def escape(s: str) -> str:
    return s.replace("&", "&amp;").replace("<", "&lt;").replace(">", "&gt;")


def md_to_flowables(md_text: str) -> list:
    out: list = []
    lines = md_text.splitlines()
    i = 0
    while i < len(lines):
        line = lines[i].rstrip()
        if not line:
            i += 1
            continue
        if line.startswith("# "):
            out.append(Paragraph(escape(line[2:].strip()), s_h1))
            i += 1
            continue
        if line.startswith("## "):
            out.append(Paragraph(escape(line[3:].strip()), s_h2))
            i += 1
            continue
        if line.startswith("### "):
            out.append(Paragraph(escape(line[4:].strip()), s_h3))
            i += 1
            continue
        if line.startswith("---"):
            out.append(Spacer(1, 6))
            i += 1
            continue
        # Tabelle?
        if (
            line.startswith("|")
            and i + 1 < len(lines)
            and re.match(r"^\|[\s\-:|]+\|$", lines[i + 1])
        ):
            header = [c.strip() for c in line.strip("|").split("|")]
            i += 2
            rows = [header]
            while i < len(lines) and lines[i].startswith("|"):
                cells = [c.strip() for c in lines[i].strip("|").split("|")]
                rows.append(cells)
                i += 1
            col_count = max(1, len(header))
            avail_width = 16 * cm
            col_widths = [avail_width / col_count] * col_count
            if col_count >= 2:
                col_widths[0] = min(4 * cm, avail_width / col_count * 1.5)
                rest = (avail_width - col_widths[0]) / (col_count - 1)
                for k in range(1, col_count):
                    col_widths[k] = rest
            tbl_data = []
            for r in rows:
                tbl_data.append([Paragraph(escape(c), s_body) for c in r])
            tbl = Table(tbl_data, colWidths=col_widths, repeatRows=1)
            tbl.setStyle(
                TableStyle(
                    [
                        ("BACKGROUND", (0, 0), (-1, 0), SURFACE),
                        ("FONTNAME", (0, 0), (-1, 0), FONT_BOLD),
                        ("BOX", (0, 0), (-1, -1), 0.4, BORDER),
                        ("INNERGRID", (0, 0), (-1, -1), 0.3, BORDER),
                        ("VALIGN", (0, 0), (-1, -1), "TOP"),
                        ("LEFTPADDING", (0, 0), (-1, -1), 4),
                        ("RIGHTPADDING", (0, 0), (-1, -1), 4),
                        ("TOPPADDING", (0, 0), (-1, -1), 3),
                        ("BOTTOMPADDING", (0, 0), (-1, -1), 3),
                    ]
                )
            )
            out.append(tbl)
            out.append(Spacer(1, 6))
            continue
        if line.startswith("- ") or line.startswith("* "):
            text = line[2:].strip()
            text = re.sub(r"\*\*([^*]+)\*\*", r"<b>\1</b>", text)
            out.append(Paragraph("• " + _inline_markup(text), s_body))
            i += 1
            continue
        if re.match(r"^\d+\.\s", line):
            text = re.sub(r"\*\*([^*]+)\*\*", r"<b>\1</b>", line)
            out.append(Paragraph(_inline_markup(text), s_body))
            i += 1
            continue
        # Sammle normalen Absatz bis zur naechsten Leerzeile/Sondersyntax
        block = [line]
        j = i + 1
        while (
            j < len(lines)
            and lines[j].strip()
            and not lines[j].startswith(("#", "-", "*", "|", "---"))
            and not re.match(r"^\d+\.\s", lines[j])
        ):
            block.append(lines[j].rstrip())
            j += 1
        text = " ".join(block).strip()
        text = re.sub(r"\*\*([^*]+)\*\*", r"<b>\1</b>", text)
        text = re.sub(r"`([^`]+)`", r"<font face='Courier'>\1</font>", text)
        out.append(Paragraph(_inline_markup(text), s_body))
        i = j
    return out


def _inline_markup(s: str) -> str:
    """Escapt Nutztext und erhält ausschließlich selbst erzeugte Inline-Tags."""
    allowed = re.compile(r"</?(?:b|i|sub|sup)>|<font face='Courier'>|</font>")
    tokens: list[str] = []

    def store(match: re.Match[str]) -> str:
        tokens.append(match.group(0))
        return f"@@INLINE{len(tokens) - 1}@@"

    escaped = escape(allowed.sub(store, s))
    for index, token in enumerate(tokens):
        escaped = escaped.replace(f"@@INLINE{index}@@", token)
    return escaped


def txt_to_flowables(text: str) -> list:
    out = []
    for para in text.split("\n\n"):
        para = para.strip()
        if not para:
            continue
        # Zeilenumbrueche im Absatz erhalten
        out.append(Paragraph(escape(para).replace("\n", "<br/>"), s_body))
    return out


class _MailHTMLTextExtractor(HTMLParser):
    """Reduziert HTML-Mails auf ihren sichtbaren Text für die PDF-Ausgabe."""

    BLOCK_TAGS = {
        "address",
        "blockquote",
        "br",
        "div",
        "h1",
        "h2",
        "h3",
        "h4",
        "h5",
        "h6",
        "hr",
        "li",
        "ol",
        "p",
        "table",
        "td",
        "th",
        "tr",
        "ul",
    }
    SKIP_TAGS = {"head", "script", "style"}

    def __init__(self) -> None:
        super().__init__(convert_charrefs=True)
        self.parts: list[str] = []
        self.skip_depth = 0

    def handle_starttag(self, tag: str, attrs) -> None:
        tag = tag.lower()
        if tag in self.SKIP_TAGS:
            self.skip_depth += 1
        elif not self.skip_depth and tag in self.BLOCK_TAGS:
            self.parts.append("\n")

    def handle_endtag(self, tag: str) -> None:
        tag = tag.lower()
        if tag in self.SKIP_TAGS and self.skip_depth:
            self.skip_depth -= 1
        elif not self.skip_depth and tag in self.BLOCK_TAGS:
            self.parts.append("\n")

    def handle_data(self, data: str) -> None:
        if not self.skip_depth:
            self.parts.append(data)

    def text(self) -> str:
        joined = html_unescape("".join(self.parts)).replace("\xa0", " ")
        lines = [re.sub(r"[ \t]+", " ", line).strip() for line in joined.splitlines()]
        return "\n".join(line for line in lines if line).strip()


def html_mail_to_text(raw_html: str) -> str:
    parser = _MailHTMLTextExtractor()
    parser.feed(raw_html)
    parser.close()
    return parser.text()


def structured_text_to_flowables(path: Path) -> list:
    """Rendert strukturierte Originaldaten lesbar, ohne ihren Inhalt zu deuten."""
    try:
        raw = path.read_text(encoding="utf-8")
    except UnicodeDecodeError:
        raw = path.read_text(encoding="latin-1")
    except Exception as exc:
        raise DocumentRenderError(f"Strukturierte Datei konnte nicht gelesen werden: {exc}") from exc

    ext = path.suffix.lower()
    try:
        if ext == ".json":
            content = json.dumps(json.loads(raw), ensure_ascii=False, indent=2)
        elif ext in {".yaml", ".yml"}:
            content = yaml.safe_dump(
                yaml.safe_load(raw),
                allow_unicode=True,
                sort_keys=False,
                default_flow_style=False,
            )
        elif ext == ".xml":
            root = ET.fromstring(raw)
            ET.indent(root)
            content = ET.tostring(root, encoding="unicode")
        else:
            content = raw
    except Exception as exc:
        raise DocumentRenderError(f"{ext[1:].upper()} konnte nicht geparst werden: {exc}") from exc

    if not content.strip():
        raise DocumentRenderError("strukturierte Datei enthält keinen lesbaren Inhalt")
    return txt_to_flowables(content)


def eml_to_flowables(path: Path) -> list:
    out = []
    try:
        with open(path, "rb") as f:
            msg = BytesParser(policy=policy.default).parse(f)
        headers = [
            ("Von", msg.get("From", "")),
            ("An", msg.get("To", "")),
            ("Kopie", msg.get("Cc", "")),
            ("Datum", msg.get("Date", "")),
            ("Betreff", msg.get("Subject", "")),
            ("Anlagen", msg.get("X-Attachments", "")),
        ]
        headers = [(label, value) for label, value in headers if value]
        body_part = msg.get_body(preferencelist=("plain", "html"))
        body = body_part.get_content() if body_part else ""
        if body_part and body_part.get_content_type() == "text/html":
            body = html_mail_to_text(body)
    except Exception as exc:
        raise DocumentRenderError(f"E-Mail konnte nicht gelesen werden: {exc}") from exc

    rows = [
        [Paragraph(label, s_meta), Paragraph(escape(value), s_meta)]
        for label, value in headers
    ]
    tbl = Table(rows, colWidths=[2.5 * cm, 13.5 * cm])
    tbl.setStyle(
        TableStyle(
            [
                ("BACKGROUND", (0, 0), (0, -1), SURFACE),
                ("BOX", (0, 0), (-1, -1), 0.3, BORDER),
                ("INNERGRID", (0, 0), (-1, -1), 0.2, BORDER),
                ("VALIGN", (0, 0), (-1, -1), "TOP"),
                ("LEFTPADDING", (0, 0), (-1, -1), 4),
                ("RIGHTPADDING", (0, 0), (-1, -1), 4),
                ("TOPPADDING", (0, 0), (-1, -1), 3),
                ("BOTTOMPADDING", (0, 0), (-1, -1), 3),
            ]
        )
    )
    out.append(tbl)
    out.append(Spacer(1, 6))
    out.extend(txt_to_flowables(body))
    return out


def csv_to_flowables(path: Path) -> list:
    out = []
    try:
        text = path.read_text(encoding="utf-8")
    except UnicodeDecodeError:
        text = path.read_text(encoding="latin-1")
    except Exception as exc:
        raise DocumentRenderError(f"CSV konnte nicht gelesen werden: {exc}") from exc

    try:
        dialect = csv.Sniffer().sniff(text[:8192], delimiters=",;\t|")
        rows = list(csv.reader(io.StringIO(text), dialect))
    except csv.Error:
        delimiter = ";" if text.count(";") > text.count(",") else ","
        rows = list(csv.reader(io.StringIO(text), delimiter=delimiter))

    if not rows:
        raise DocumentRenderError("CSV enthält keine lesbare Datenzeile")
    max_cols = max(len(r) for r in rows)
    rows = [r + [""] * (max_cols - len(r)) for r in rows]
    out.extend(_render_table(rows, header=True))
    return out


def xlsx_to_flowables(path: Path) -> list:
    out = []
    try:
        wb = load_workbook(path, data_only=True)
    except Exception as exc:
        raise DocumentRenderError(f"XLSX konnte nicht gelesen werden: {exc}") from exc
    try:
        for sheet_name in wb.sheetnames:
            ws = wb[sheet_name]
            sheet_flow: list = []
            rows = [[_format_cell(c) for c in row] for row in ws.iter_rows(values_only=True)]
            while rows and not any(c.strip() for c in rows[-1]):
                rows.pop()
            if not rows:
                continue
            max_cols = max(len(r) for r in rows)
            while max_cols > 0 and all(
                (len(r) <= max_cols - 1) or (not r[max_cols - 1].strip()) for r in rows
            ):
                max_cols -= 1
            if max_cols == 0:
                continue
            rows = [r[:max_cols] + [""] * (max_cols - len(r[:max_cols])) for r in rows]
            sheet_flow.append(Paragraph(f"Tabellenblatt: {escape(sheet_name)}", s_h3))
            sheet_flow.extend(_render_table(rows, header=True))
            sheet_flow.append(Spacer(1, 6))
            out.extend(sheet_flow)
    finally:
        wb.close()
    if not out:
        raise DocumentRenderError("XLSX enthält kein lesbares Tabellenblatt")
    return out


def _format_cell(c) -> str:
    if c is None:
        return ""
    if isinstance(c, float):
        if c == int(c):
            return str(int(c))
        return f"{c:.4f}".rstrip("0").rstrip(".")
    return str(c)


# Maximalzeichen pro Zelle, ab denen die Tabelle nicht mehr als Table gerendert wird,
# sondern als sequentielle Absatzfolge (verhindert ReportLab-Overflow).
_MAX_CELL_CHARS = 1200


def _split_long_text(text: str, chunk: int = 800) -> list:
    """Schneidet sehr langen Text an Absatz- oder Satzgrenzen in Stuecke."""
    text = text.replace("\r", "")
    if len(text) <= chunk:
        return [text]
    # Erst Absaetze probieren
    paras = [p for p in text.split("\n") if p.strip()]
    if any(len(p) > chunk for p in paras):
        # Weiter an Saetzen schneiden
        out = []
        for p in paras:
            if len(p) <= chunk:
                out.append(p)
                continue
            buf = ""
            for sent in p.replace("; ", "; |").replace(". ", ". |").split("|"):
                if len(buf) + len(sent) > chunk and buf:
                    out.append(buf)
                    buf = sent
                else:
                    buf += sent
            if buf:
                out.append(buf)
        return out
    return paras


def _render_table(rows: list, header: bool = False) -> list:
    """Rendert eine Tabelle. Falls Zellen zu lang werden, faellt es auf eine
    sequentielle Absatzdarstellung zurueck (Reihe fuer Reihe), damit ReportLab
    keine Overflow-Fehler wirft."""
    max_cell_len = max((len(c) for r in rows for c in r), default=0)
    max_cols_in_table = max((len(r) for r in rows), default=0)
    # Bei sehr breiten Tabellen (>12 Spalten) faellt es sequentiell zurueck,
    # weil die Spaltenbreite sonst kleiner ist als die kleinste Wortbreite
    # und ReportLab Cell-Overflow-Fehler wirft.
    record_layout = max_cols_in_table > 5 and max_cell_len > 40
    if max_cell_len > _MAX_CELL_CHARS or max_cols_in_table > 12 or record_layout:
        out = []
        header_row = rows[0] if header else None
        body_rows = rows[1:] if header else rows
        for ri, r in enumerate(body_rows):
            record = []
            if header_row:
                # Reihen-Trennlinie + Spaltenkopf pro Zelle
                for ci, cell in enumerate(r):
                    if not cell.strip():
                        continue
                    label = header_row[ci] if ci < len(header_row) else f"Spalte {ci+1}"
                    if label.strip():
                        record.append(Paragraph(f"<b>{escape(label)}</b>", s_meta))
                    for chunk in _split_long_text(cell):
                        record.append(Paragraph(escape(chunk), s_body))
            else:
                for ci, cell in enumerate(r):
                    for chunk in _split_long_text(cell):
                        record.append(Paragraph(escape(chunk), s_body))
            record.append(Spacer(1, 4))
            record.append(HRFlowable(width="100%", thickness=0.2, color=BORDER))
            record.append(Spacer(1, 4))
            out.append(KeepTogether(record))
        return out

    max_cols = max(len(r) for r in rows)
    avail_width = 16 * cm
    col_widths = [avail_width / max_cols] * max_cols
    data = [
        [Paragraph(escape(c), s_meta) for c in r]
        for r in rows
    ]
    tbl = Table(data, colWidths=col_widths, repeatRows=1 if header else 0, splitByRow=1)
    cmds = [
        ("BOX", (0, 0), (-1, -1), 0.3, BORDER),
        ("INNERGRID", (0, 0), (-1, -1), 0.2, BORDER),
        ("VALIGN", (0, 0), (-1, -1), "TOP"),
        ("LEFTPADDING", (0, 0), (-1, -1), 3),
        ("RIGHTPADDING", (0, 0), (-1, -1), 3),
        ("TOPPADDING", (0, 0), (-1, -1), 2),
        ("BOTTOMPADDING", (0, 0), (-1, -1), 2),
    ]
    if header:
        cmds.insert(0, ("BACKGROUND", (0, 0), (-1, 0), SURFACE))
        cmds.insert(1, ("FONTNAME", (0, 0), (-1, 0), FONT_BOLD))
    tbl.setStyle(TableStyle(cmds))
    return [tbl]


def docx_to_flowables(path: Path) -> list:
    out = []
    if Document is None:
        raise DocumentRenderError("python-docx ist nicht installiert")
    try:
        doc = Document(str(path))
    except Exception as exc:
        raise DocumentRenderError(f"DOCX konnte nicht gelesen werden: {exc}") from exc
    for para in doc.paragraphs:
        text = para.text.strip()
        if not text:
            continue
        style = para.style.name if para.style else ""
        if style.startswith("Heading 1"):
            out.append(Paragraph(escape(text), s_h2))
        elif style.startswith("Heading 2"):
            out.append(Paragraph(escape(text), s_h3))
        elif style.startswith("Heading"):
            out.append(Paragraph(escape(text), s_h3))
        else:
            out.append(Paragraph(escape(text), s_body))
    for table in doc.tables:
        rows = []
        for r in table.rows:
            rows.append([c.text.strip() for c in r.cells])
        if rows:
            out.extend(_render_table(rows, header=True))
            out.append(Spacer(1, 6))
    if not out:
        raise DocumentRenderError("DOCX enthält keinen lesbaren Text und keine Tabelle")
    return out


def odt_to_flowables(path: Path) -> list:
    """Rendert ODT-Inhalte vollständig, wenn native Konvertierung fehlt."""
    try:
        from odf.element import Element
        from odf.opendocument import load as odf_load
    except ImportError as exc:
        raise DocumentRenderError("odfpy ist nicht installiert") from exc
    try:
        doc = odf_load(str(path))
    except Exception as exc:
        raise DocumentRenderError(f"ODT konnte nicht gelesen werden: {exc}") from exc

    out: list = []

    def text_of(node) -> str:
        parts: list[str] = []
        for child in node.childNodes:
            if child.nodeType == child.TEXT_NODE:
                parts.append(child.data)
            elif isinstance(child, Element):
                parts.append(text_of(child))
        return "".join(parts)

    def walk(node) -> None:
        for child in node.childNodes:
            if not isinstance(child, Element):
                continue
            local = child.qname[1]
            if local in ("p", "h"):
                text = text_of(child).strip()
                if text:
                    out.append(Paragraph(escape(text), s_h3 if local == "h" else s_body))
            else:
                walk(child)

    walk(doc.body)
    if not out:
        raise DocumentRenderError("ODT enthält keinen lesbaren Text")
    return out


def image_to_flowables(path: Path) -> list:
    out = []
    try:
        width, height = ImageReader(str(path)).getSize()
        max_width = 16 * cm
        max_height = 22 * cm
        scale = min(max_width / width, max_height / height, 1)
        img = RLImage(str(path), width=width * scale, height=height * scale)
        out.append(img)
        out.append(Spacer(1, 4))
        out.append(Paragraph(f"Bilddatei: {escape(path.name)}", s_meta))
    except Exception as exc:
        raise DocumentRenderError(f"Bild konnte nicht gerendert werden: {exc}") from exc
    return out


def header_footer_factory(testakte_name: str):
    def hf(canv: canvas.Canvas, doc) -> None:
        canv.saveState()
        canv.setFont(FONT_REG, 8)
        canv.setFillColor(MUTED)
        canv.drawString(2 * cm, 1.2 * cm, f"Akte: {testakte_name}")
        canv.drawRightString(19 * cm, 1.2 * cm, f"Seite {doc.page}")
        canv.setStrokeColor(BORDER)
        canv.setLineWidth(0.3)
        canv.line(2 * cm, 1.6 * cm, 19 * cm, 1.6 * cm)
        canv.restoreState()

    return hf


def no_header_footer(canv: canvas.Canvas, doc) -> None:
    return None


def build_cover(_name: str, _readme_summary: str | None, h1: str | None = None) -> list:
    # Historisch gab es hier ein Titelblatt. Das Gesamt-PDF soll jetzt direkt
    # mit dem ersten Aktenstück beginnen; die Funktion bleibt als Kompatibilitaet
    # fuer aeltere Aufrufe erhalten.
    return []


def extract_readme_summary(readme_path: Path) -> tuple[str | None, str | None]:
    """Liest aus der README den H1-Titel und einen kurzen beschreibenden Absatz.

    Der beschreibende Absatz wird absichtlich erst gesucht, NACHDEM Download-Bloecke
    und Aktenstruktur-Blocks uebersprungen wurden, damit nicht der ZIP-Hinweis
    auf dem Cover landet.
    """
    if not readme_path.is_file():
        return None, None
    text = readme_path.read_text(encoding="utf-8")
    # H1
    h1_match = re.search(r"^#\s+(.+?)\s*$", text, flags=re.MULTILINE)
    h1 = h1_match.group(1).strip() if h1_match else None

    # Suche eine Sektion mit beschreibendem Inhalt: 'Kurzbild', 'Worum',
    # 'Sachverhalt', 'Ueberblick', 'Mandat', 'Fall' o.ae.
    section_pattern = re.compile(
        r"^##[^\n]*?(?:kurzbild|worum geht|sachverhalt|ueberblick|\u00fcberblick|mandat|fall|der fall|akte|kontext|ausgangslage|ausgangs|zweck|szenario|idee|einsatz|\u00fcbersicht|uebersicht|verfahrenseckdaten|aktenkern|aktenbestand|mandantenkonstellation|politische vorgabe|enthaltene arbeitsdateien|dateien)[^\n]*\n([\s\S]*?)(?=^## |\Z)",
        re.IGNORECASE | re.MULTILINE,
    )
    m = section_pattern.search(text)
    candidate_text = m.group(1) if m else text
    for para in candidate_text.split("\n\n"):
        para = para.strip()
        if not para:
            continue
        if para.startswith(("#", "-", "*", "|", "<!--", "```")):
            continue
        # Download-/ZIP-Hinweise ueberspringen
        lower = para.lower()
        if any(
            kw in lower
            for kw in (
                "zip-datei",
                "zip datei",
                "direkt-download",
                "als zip",
                "github-release",
                "github release",
                "download",
            )
        ):
            continue
        para = re.sub(r"\*\*([^*]+)\*\*", r"\1", para)
        para = re.sub(r"\s+", " ", para)
        return h1, para[:400]
    return h1, None


def collect_files(testakte_dir: Path) -> dict[str, list[Path]]:
    files_by_type: dict[str, list[Path]] = {t: [] for t in TYPE_ORDER}
    for f in testakte_dir.rglob("*"):
        if not include_in_working_dump(f, testakte_dir):
            continue
        ext = f.suffix.lower().lstrip(".")
        if ext in IMAGE_EXTS:
            files_by_type["image"].append(f)
            continue
        if ext in STRUCTURED_EXTS:
            files_by_type["structured"].append(f)
            continue
        if ext not in TYPE_ORDER:
            continue
        files_by_type[ext].append(f)
    for t in files_by_type:
        files_by_type[t].sort(key=lambda p: str(p.relative_to(testakte_dir)).lower())
    return files_by_type


def file_header_flowables(type_key: str, rel: Path) -> list:
    """Einheitlicher, gut lesbarer Kopf je Aktenstück: Kategorie-Kicker,
    Dateiname als Titel, ggf. Unterordnerpfad, darunter eine Trennlinie."""
    kicker = TYPE_LABEL.get(type_key, "Aktenstück")
    out = [
        Paragraph(escape(kicker.upper()), s_file_kicker),
        Paragraph(escape(rel.name), s_file_title),
    ]
    if str(rel.parent) not in (".", ""):
        out.append(Paragraph(escape(str(rel)), s_file_path))
    out.append(HRFlowable(width="100%", thickness=0.6, color=TEAL,
                          spaceBefore=1, spaceAfter=8))
    return out


def draw_separator_header(c, kicker: str, label: str) -> None:
    """Zeichnet den Kopf einer Trennseite (Anhänge) im gleichen Stil wie
    file_header_flowables: Kategorie-Kicker, Dateititel, ggf. Pfad, Trennlinie."""
    # Aufrufer übergeben ein Label mit vorangestelltem Etikett
    # ("PDF-Anhang: pfad" / "Office-Dokument: pfad"); für die Anzeige wird der
    # reine Pfad verwendet, das Etikett steckt bereits im Kicker.
    rel = label.split(": ", 1)[-1].strip() if ": " in label else label
    name = rel.rsplit("/", 1)[-1]
    c.setFont(FONT_BOLD, 8)
    c.setFillColor(MUTED)
    c.drawString(2 * cm, 26.2 * cm, kicker.upper())
    c.setFont(FONT_BOLD, 13)
    c.setFillColor(TEAL)
    c.drawString(2 * cm, 25.5 * cm, name)
    y_rule = 25.2 * cm
    if "/" in rel:
        c.setFont(FONT_REG, 8)
        c.setFillColor(MUTED)
        c.drawString(2 * cm, 25.0 * cm, rel)
        y_rule = 24.7 * cm
    c.setStrokeColor(TEAL)
    c.setLineWidth(0.6)
    c.line(2 * cm, y_rule, 19 * cm, y_rule)


def build_text_pdf(
    testakte_dir: Path,
    files: dict[str, list[Path]],
    cover: list,
    tmp_path: Path,
) -> tuple[list[tuple[str, bytes]], list[Path], bool]:
    """Baut Textteile und sammelt layoutgetreue Office- sowie Original-PDFs."""
    doc = SimpleDocTemplate(
        str(tmp_path),
        pagesize=A4,
        leftMargin=2 * cm, rightMargin=2 * cm,
        topMargin=2 * cm, bottomMargin=2 * cm,
        title=f"Akte {testakte_dir.name}",
        author="Kanzleiakte",
    )
    flow = list(cover)

    pdf_attachments: list[Path] = []
    office_attachments: list[tuple[str, bytes]] = []
    try:
        office_cache = render_office_batch(files["docx"] + files["odt"])
    except OfficeRenderError as exc:
        raise DocumentRenderError(str(exc)) from exc
    for t in TYPE_ORDER:
        if not files[t]:
            continue
        if t == "pdf":
            # PDFs werden separat angehaengt (Original-Layout bewahren)
            pdf_attachments = files[t]
            continue
        for f in files[t]:
            rel = f.relative_to(testakte_dir)
            if t in OFFICE_EXTS and f in office_cache:
                office_attachments.append((str(rel), office_cache[f]))
                continue
            flow.extend(file_header_flowables(t, rel))
            try:
                if t == "md":
                    rendered = md_to_flowables(f.read_text(encoding="utf-8", errors="strict"))
                elif t == "txt":
                    rendered = txt_to_flowables(f.read_text(encoding="utf-8", errors="strict"))
                elif t == "eml":
                    rendered = eml_to_flowables(f)
                elif t == "csv":
                    rendered = csv_to_flowables(f)
                elif t == "xlsx":
                    rendered = xlsx_to_flowables(f)
                elif t == "structured":
                    rendered = structured_text_to_flowables(f)
                elif t == "docx":
                    rendered = docx_to_flowables(f)
                elif t == "odt":
                    rendered = odt_to_flowables(f)
                elif t == "image":
                    rendered = image_to_flowables(f)
                else:
                    raise DocumentRenderError(f"nicht unterstützter Dokumenttyp: {t}")
            except Exception as exc:
                if isinstance(exc, DocumentRenderError):
                    raise DocumentRenderError(f"{rel}: {exc}") from exc
                raise DocumentRenderError(f"{rel}: {type(exc).__name__}: {exc}") from exc
            if not rendered:
                raise DocumentRenderError(f"{rel}: Konverter lieferte keine PDF-Inhalte")
            flow.extend(rendered)
            # Jedes Aktenstueck beginnt im Gesamt-PDF auf einer neuen Seite
            flow.append(PageBreak())

    # Trailing PageBreak entfernen, damit am Ende keine Leerseite entsteht.
    while flow and isinstance(flow[-1], PageBreak):
        flow.pop()

    # Bei reinen Office-/PDF-Akten beginnt die Ausgabe direkt mit dem ersten
    # Dokumenttrenner. Eine künstliche Füllseite würde nur den Aktenlauf stören.
    if not flow:
        return office_attachments, pdf_attachments, False

    hf = header_footer_factory(testakte_dir.name)
    try:
        doc.build(
            flow,
            onFirstPage=hf,
            onLaterPages=hf,
            canvasmaker=invariant_canvas,
        )
    except Exception as exc:
        raise DocumentRenderError(f"Text-PDF konnte nicht gebaut werden: {exc}") from exc
    return office_attachments, pdf_attachments, True


def a4_normalized_page(source_page):
    """Setzt nicht-A4-Seiten proportional und zentriert auf einen A4-Bogen.

    Das Aktenstück selbst bleibt unverändert. Die Normalisierung betrifft nur
    das zusammengeführte Gesamt-PDF, damit Druck, Durchsicht und Seitennavigation
    nicht zwischen Letter-, A4- und Sonderformaten springen.
    """
    page = deepcopy(source_page)
    if getattr(page, "rotation", 0):
        page.transfer_rotation_to_content()

    width = float(page.mediabox.width)
    height = float(page.mediabox.height)
    portrait = width <= height
    target_width, target_height = A4_PORTRAIT
    if not portrait:
        target_width, target_height = target_height, target_width

    if abs(width - target_width) < 1 and abs(height - target_height) < 1:
        return page

    scale = min(target_width / width, target_height / height)
    x_offset = (target_width - width * scale) / 2
    y_offset = (target_height - height * scale) / 2
    target = PageObject.create_blank_page(width=target_width, height=target_height)
    transform = Transformation().scale(scale).translate(x_offset, y_offset)
    target.merge_transformed_page(page, transform, over=True)
    return target


def append_pdf_with_separator(writer: PdfWriter, label: str, pdf_path: Path, testakte_name: str) -> None:
    try:
        attachment_pages = list(PdfReader(str(pdf_path)).pages)
    except Exception as exc:
        raise DocumentRenderError(f"{pdf_path.name}: PDF konnte nicht gelesen werden: {exc}") from exc
    if not attachment_pages:
        raise DocumentRenderError(f"{pdf_path.name}: PDF enthält keine Seite")

    sep = io.BytesIO()
    c = canvas.Canvas(sep, pagesize=A4, invariant=1)
    c.setTitle(label)
    c.setAuthor("Kanzleiakte")
    draw_separator_header(c, "PDF-Anhang (Originaldokument)", label)
    c.setStrokeColor(BORDER)
    c.setLineWidth(0.3)
    c.line(2 * cm, 1.6 * cm, 19 * cm, 1.6 * cm)
    c.setFont(FONT_REG, 8)
    c.setFillColor(MUTED)
    c.drawString(2 * cm, 1.2 * cm, testakte_name)
    c.showPage()
    c.save()
    sep.seek(0)
    for p in PdfReader(sep).pages:
        writer.add_page(p)
    for page in attachment_pages:
        writer.add_page(a4_normalized_page(page))


def append_pdf_bytes_with_separator(
    writer: PdfWriter,
    label: str,
    data: bytes,
    testakte_name: str,
) -> None:
    try:
        attachment_pages = list(PdfReader(io.BytesIO(data)).pages)
    except Exception as exc:
        raise DocumentRenderError(f"{label}: PDF konnte nicht gelesen werden: {exc}") from exc
    if not attachment_pages:
        raise DocumentRenderError(f"{label}: PDF enthält keine Seite")

    sep = io.BytesIO()
    c = canvas.Canvas(sep, pagesize=A4, invariant=1)
    c.setTitle(label)
    c.setAuthor("Kanzleiakte")
    low = label.lower()
    if low.endswith(".docx"):
        kicker = "Word-Dokument (Original-Layout)"
    elif low.endswith(".odt"):
        kicker = "OpenDocument-Textdatei (Original-Layout)"
    else:
        kicker = "Dokument (Original-Layout)"
    draw_separator_header(c, kicker, label)
    c.setStrokeColor(BORDER)
    c.setLineWidth(0.3)
    c.line(2 * cm, 1.6 * cm, 19 * cm, 1.6 * cm)
    c.setFont(FONT_REG, 8)
    c.setFillColor(MUTED)
    c.drawString(2 * cm, 1.2 * cm, testakte_name)
    c.showPage()
    c.save()
    sep.seek(0)
    for page in PdfReader(sep).pages:
        writer.add_page(page)
    for page in attachment_pages:
        writer.add_page(a4_normalized_page(page))


def build_gesamt_pdf(testakte_dir: Path) -> tuple[str, str]:
    """Gibt (status, info) zurueck. status in {ok, skip, error}."""
    name = testakte_dir.name
    out_dir = testakte_dir / "gesamt-pdf"
    out_dir.mkdir(exist_ok=True)
    out_path = out_dir / f"{name}_gesamt.pdf"

    files = collect_files(testakte_dir)
    total_files = sum(len(v) for v in files.values())
    if total_files == 0:
        return "skip", "keine Quelldateien"

    cover: list = []

    with tempfile.NamedTemporaryFile(prefix=f"gesamt-{name}-", suffix=".pdf", delete=False) as handle:
        tmp_text = Path(handle.name)
    tmp_output = out_path.with_name(f".{out_path.name}.tmp")
    try:
        office_attachments, pdf_attachments, has_text_pdf = build_text_pdf(
            testakte_dir, files, cover, tmp_text
        )
        writer = PdfWriter()
        if has_text_pdf:
            for page in PdfReader(str(tmp_text)).pages:
                writer.add_page(page)
        for label, data in office_attachments:
            append_pdf_bytes_with_separator(writer, f"Office-Dokument: {label}", data, name)
        for pdf in pdf_attachments:
            rel = pdf.relative_to(testakte_dir)
            append_pdf_with_separator(writer, f"PDF-Anhang: {rel}", pdf, name)
        writer.add_metadata(
            {
                "/Title": f"Akte {name}",
                "/Author": "Kanzleiakte",
                "/Subject": "Gesamtakte",
            }
        )
        with tmp_output.open("wb") as handle:
            writer.write(handle)
        if not list(PdfReader(str(tmp_output)).pages):
            raise DocumentRenderError("erzeugtes Gesamt-PDF enthält keine Seite")
        tmp_output.replace(out_path)
    except Exception as exc:
        return "error", str(exc)
    finally:
        tmp_text.unlink(missing_ok=True)
        tmp_output.unlink(missing_ok=True)
    size_kb = out_path.stat().st_size / 1024
    return "ok", f"{out_path.relative_to(REPO_ROOT)} ({size_kb:.0f} KB, {total_files} Quelldateien)"


def main() -> int:
    targets = sys.argv[1:]
    all_dirs = sorted([d for d in TESTAKTEN.iterdir() if d.is_dir()])
    if targets:
        all_dirs = [d for d in all_dirs if d.name in targets]
    print(f"Verarbeite {len(all_dirs)} Testakten")
    print()
    counts = {"ok": 0, "skip": 0, "error": 0}
    for d in all_dirs:
        status, info = build_gesamt_pdf(d)
        counts[status] += 1
        sigil = {"ok": "OK ", "skip": "SK ", "error": "ERR"}[status]
        print(f"  {sigil} {d.name}: {info}")
    print()
    print(f"Fertig: {counts['ok']} OK, {counts['skip']} skip, {counts['error']} Fehler")
    return 1 if counts["error"] else 0


if __name__ == "__main__":
    raise SystemExit(main())
