#!/usr/bin/env python3
"""Setzt Briefkopf, Hausfarbe und Fusszeile in die DOCX-Aktenstuecke.

Die Aktenstuecke lagen bisher ohne Kopfzeile und ohne jede Farbe vor. Dieses
Skript erkennt je Dokument den Absender, waehlt eine stabile Hausfarbe und
setzt Kopf- und Fusszeile. Der Fliesstext bleibt unangetastet: Times New Roman
11 pt, schwarz, A4, dezimale Gliederung.

Aufruf:
  python3 scripts/apply-testakte-briefkopf.py            # alle Testakten
  python3 scripts/apply-testakte-briefkopf.py <slug> ... # gezielt
  python3 scripts/apply-testakte-briefkopf.py --pruefen  # nur berichten
"""

from __future__ import annotations

import sys
from collections import Counter
from pathlib import Path

from docx import Document

sys.path.insert(0, str(Path(__file__).resolve().parent))
from testakte_briefkopf import apply_briefkopf, detect_absender, detect_aktenzeichen

REPO = Path(__file__).resolve().parent.parent
TESTAKTEN = REPO / "testakten"
SKIP_DIRS = {"gesamt-pdf", "megaprompts", "formatvorlagen-paradebeispiele"}


def has_briefkopf(doc) -> bool:
    header = doc.sections[0].header
    if any(p.text.strip() for p in header.paragraphs):
        return True
    return any(
        paragraph.text.strip()
        for table in header.tables
        for row in table.rows
        for cell in row.cells
        for paragraph in cell.paragraphs
    )


def process(path: Path, pruefen: bool) -> str:
    try:
        doc = Document(str(path))
    except Exception as exc:  # defekte oder fremde Datei
        return f"fehler:{exc.__class__.__name__}"
    paragraphs = [p.text for p in doc.paragraphs[:14]]
    absender = detect_absender(paragraphs)
    if absender is None:
        return "ohne-absender"
    if has_briefkopf(doc):
        return "bereits-gesetzt"
    if pruefen:
        return f"wuerde:{absender.typ}"
    apply_briefkopf(doc, absender, detect_aktenzeichen(paragraphs))
    doc.save(str(path))
    return absender.typ


def main() -> int:
    args = [a for a in sys.argv[1:] if not a.startswith("--")]
    pruefen = "--pruefen" in sys.argv
    if args:
        cases = [TESTAKTEN / slug for slug in args]
    else:
        cases = [p for p in sorted(TESTAKTEN.iterdir()) if p.is_dir() and p.name not in SKIP_DIRS]

    stats: Counter[str] = Counter()
    for case in cases:
        if not case.is_dir():
            print(f"unbekannte Testakte: {case.name}", file=sys.stderr)
            return 1
        for docx in sorted(case.rglob("*.docx")):
            if any(part in SKIP_DIRS for part in docx.parts):
                continue
            stats[process(docx, pruefen)] += 1

    gesetzt = sum(v for k, v in stats.items() if k in {
        "kanzlei", "gericht", "behoerde", "kasse", "notariat", "sachverstaendig", "unternehmen",
    })
    print(
        "Briefkopf "
        + ("geprüft" if pruefen else "gesetzt")
        + f": {gesetzt} Dokumente | "
        + ", ".join(f"{k}: {v}" for k, v in sorted(stats.items()))
    )
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
