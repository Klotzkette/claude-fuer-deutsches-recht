#!/usr/bin/env python3
"""Regressionen für die Gliederungsabstände der BWA-Unterlagen."""

import importlib.util
from pathlib import Path
import tempfile
import unittest

from docx import Document

spec = importlib.util.spec_from_file_location(
    "bwa_abstimmung", Path(__file__).with_name("validate-testakte-bwa-abstimmung.py")
)
BWA = importlib.util.module_from_spec(spec)
spec.loader.exec_module(BWA)


class HeadingSpacingTests(unittest.TestCase):
    def check_document(self, *, blank, keep=True):
        document = Document()
        heading = document.add_heading("1 Vertragsgegenstand", level=1)
        heading.paragraph_format.keep_with_next = keep
        if blank:
            empty = document.add_paragraph()
            empty.paragraph_format.keep_with_next = keep
        document.add_paragraph("Der Vertrag betrifft die bezeichnete Lieferung.")
        with tempfile.TemporaryDirectory() as directory:
            path = Path(directory) / "vertrag.docx"
            document.save(path)
            BWA.validate_word_headings(path)

    def test_blank_paragraph_is_required(self):
        with self.assertRaisesRegex(ValueError, "Leerzeile"):
            self.check_document(blank=False)

    def test_heading_cannot_be_stranded(self):
        with self.assertRaisesRegex(ValueError, "folgenden Inhalt"):
            self.check_document(blank=True, keep=False)

    def test_readable_heading_group_passes(self):
        self.check_document(blank=True)


if __name__ == "__main__":
    unittest.main()
