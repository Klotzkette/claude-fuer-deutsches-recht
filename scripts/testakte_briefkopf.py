#!/usr/bin/env python3
"""Briefkopf, Hausfarbe und Fusszeile fuer DOCX-Aktenstuecke.

Aktenstuecke sollen aussehen wie echte Post aus der Akte: eine Kanzlei hat
einen anderen Briefkopf als ein Landgericht, eine Krankenkasse einen anderen
als ein Notariat. Dieses Modul erkennt den Absender aus dem Dokument, waehlt
eine stabile, gedeckte Hausfarbe und setzt Kopf- und Fusszeile.

Grundsaetze:
- Der Fliesstext bleibt Times New Roman 11 pt und schwarz. Farbe traegt nur
  der Briefkopf, die Trennlinie und allenfalls die Dokumentueberschrift.
- Gerichte und Behoerden werden nuechtern gesetzt, ohne Farbe auf dem Namen.
- Die Hausfarbe wird deterministisch aus dem Absendernamen abgeleitet, damit
  dieselbe Kanzlei ueber alle Aktenstuecke hinweg gleich aussieht.
"""

from __future__ import annotations

import hashlib
import re
from dataclasses import dataclass

from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_TAB_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor

# Gedeckte Palette. Keine grellen Toene; alle Werte sind auf weissem Papier
# und im Graustufendruck lesbar.
PALETTE: tuple[tuple[str, str], ...] = (
    ("Dunkelblau", "1F3864"),
    ("Bordeaux", "7B2D3B"),
    ("Petrol", "11555F"),
    ("Waldgruen", "2C5230"),
    ("Anthrazit", "3B3F45"),
    ("Ocker", "8A6A1F"),
)
NEUTRAL = "3B3F45"
GREY = RGBColor(0x6B, 0x6B, 0x66)

KANZLEI = re.compile(
    r"(Rechtsanw[äa]lt\w*|Kanzlei|PartG|PartmbB|mbB|Anwaltskanzlei|Rechtsanwaltsgesellschaft|"
    r"Fachanw[äa]lt\w*\s+f[üu]r|Patentanw[äa]lt\w*|Steuerberat\w+|Wirtschaftspr[üu]f\w+)",
    re.IGNORECASE,
)
GERICHT = re.compile(
    r"((?:Amts|Land|Ober|Verwaltungs|Sozial|Arbeits|Finanz|Verfassungs|Patent)?gericht\w*"
    r"|Bundesgerichtshof|Bundesverfassungsgericht|Oberlandesgericht|Landgericht|Amtsgericht"
    r"|Staatsanwaltschaft|Notariat|Nachlassgericht|Insolvenzgericht|Registergericht)",
    re.IGNORECASE,
)
BEHOERDE = re.compile(
    r"(Beh[öo]rde|Bundesamt|Landesamt|Versorgungsamt|Finanzamt|Jobcenter|Agentur f[üu]r Arbeit"
    r"|Ministerium|Senatsverwaltung|Stadtverwaltung|Landkreis|Gemeinde|Bauamt|Ordnungsamt"
    r"|Veterin[äa]ramt|Aufsichtsbeh[öo]rde|Regierungspr[äa]sidium|Bundesnetzagentur|BaFin"
    r"|Rentenversicherung|Berufsgenossenschaft|Datenschutzbeh[öo]rde)",
    re.IGNORECASE,
)
KASSE = re.compile(
    r"(Krankenkasse|BKK|AOK|Barmer|Pflegekasse|Versicherung\w*|Versicherer|Krankenversicherung"
    r"|Unfallkasse|Zusatzversorgungskasse)",
    re.IGNORECASE,
)
NOTAR = re.compile(r"(Notar(?:in|iat)?\b|Urkundenrolle)", re.IGNORECASE)
SACHVERSTAENDIG = re.compile(
    r"(Sachverst[äa]ndig\w+|Gutachter\w*|Ingenieurb[üu]ro|Pr[üu]finstitut|Dipl\.-Ing\.)",
    re.IGNORECASE,
)
UNTERNEHMEN = re.compile(r"\b(GmbH|AG|KG|OHG|SE|e\.?\s?V\.?|gGmbH|UG)\b")

# Absendertyp -> (Bezeichnung, Farbe erlaubt, Versalien im Namen)
TYP_STIL = {
    "kanzlei": ("Rechtsanwaltskanzlei", True, True),
    "notariat": ("Notariat", True, True),
    "sachverstaendig": ("Sachverständigenbüro", True, False),
    "unternehmen": ("Unternehmen", True, False),
    "kasse": ("Versicherungsträger", True, False),
    "gericht": ("Gericht", False, False),
    "behoerde": ("Behörde", False, False),
    "neutral": ("Aktenstück", False, False),
}


@dataclass(frozen=True)
class Absender:
    name: str
    typ: str
    zusatz: str = ""

    @property
    def farbe(self) -> str:
        _label, farbig, _caps = TYP_STIL[self.typ]
        if not farbig:
            return NEUTRAL
        digest = hashlib.sha256(self.name.lower().encode("utf-8")).hexdigest()
        return PALETTE[int(digest[:8], 16) % len(PALETTE)][1]

    @property
    def kapitaelchen(self) -> bool:
        return TYP_STIL[self.typ][2]


def classify(text: str) -> str:
    if NOTAR.search(text):
        return "notariat"
    if GERICHT.search(text):
        return "gericht"
    if KANZLEI.search(text):
        return "kanzlei"
    if BEHOERDE.search(text):
        return "behoerde"
    if KASSE.search(text):
        return "kasse"
    if SACHVERSTAENDIG.search(text):
        return "sachverstaendig"
    if UNTERNEHMEN.search(text):
        return "unternehmen"
    return "neutral"


def detect_absender(paragraphs: list[str]) -> Absender | None:
    """Liest Absendername und Zusatzzeile aus den ersten Absaetzen.

    Die Aktenstuecke fuehren den Absender regelmaessig in den ersten Zeilen,
    entweder in der Ueberschrift (Gerichtsentscheidungen) oder als eigene
    Anschriftenzeile (Kanzlei-, Behoerden- und Unternehmenspost).
    """
    # Der Briefkopf steht immer oben. Deshalb wird zeilenweise von oben nach
    # unten geprueft und die erste Zeile genommen, die einen Absender traegt.
    # Eine Erwaehnung im Fliesstext ("... vor dem Notar X ...") darf den
    # Absender nicht kapern; lange Fliesstextzeilen bleiben daher aussen vor.
    kandidaten = [p.strip() for p in paragraphs[:8] if p.strip()]
    for zeile in kandidaten:
        if len(zeile) > 160:
            continue
        typ = classify(zeile)
        if typ == "neutral":
            continue
        teile = [t.strip() for t in zeile.split(",") if t.strip()]
        kern = extract_entity(zeile, typ, teile)
        if not kern or not (3 <= len(kern) <= 90):
            continue
        zusatz = ", ".join(t for t in teile[1:] if t.lower() not in kern.lower())[:120]
        return Absender(name=kern, typ=typ, zusatz=zusatz)
    return None


# Firmennamen und Notarnamen stehen oft mitten in einer Titelzeile
# ("Gesellschaftsvertrag der Neuralis MedTech GmbH"). Der Briefkopf soll den
# Absender zeigen, nicht die Dokumentbezeichnung.
FIRMA_NAME = re.compile(
    r"([A-ZÄÖÜ][\w.\-äöüß]*(?:\s+(?:&|und)\s+[A-ZÄÖÜ][\w.\-äöüß]*|\s+[A-ZÄÖÜ0-9][\w.\-äöüß]*){0,4}\s+"
    r"(?:gGmbH|GmbH\s*&\s*Co\.?\s*KG|GmbH|AG|SE|KGaA|KG|OHG|UG|mbH|PartG\s*mbB|PartG|mbB|"
    r"Rechtsanwälte|Rechtsanwaltskanzlei|Steuerberater|Partnerschaft))"
)
NOTAR_NAME = re.compile(
    r"(Notar(?:in)?\s+(?:Dr\.\s+)?[A-ZÄÖÜ][\wäöüß-]+(?:\s+[A-ZÄÖÜ][\wäöüß-]+)?)"
)
# Einzelanwaeltinnen und Einzelanwaelte fuehren keinen Gesellschaftszusatz.
ANWALT_SOLO = re.compile(
    r"((?:Rechtsanwalt|Rechtsanwältin|Fachanwalt|Fachanwältin|Patentanwalt|Patentanwältin"
    r"|Steuerberater|Steuerberaterin)\s+(?:Dr\.\s+|Prof\.\s+Dr\.\s+)?"
    r"[A-ZÄÖÜ][\wäöüß-]+(?:\s+[A-ZÄÖÜ][\wäöüß-]+)?)"
)
# Auch aus Titelzeilen wie "Durchsuchungsbeschluss des Amtsgerichts Leipzig".
GERICHT_NAME = re.compile(
    r"((?:Amts|Land|Ober(?:landes|verwaltungs)?|Verwaltungs|Sozial|Arbeits|Finanz|Bundesarbeits|"
    r"Bundessozial|Bundesfinanz|Bundesverwaltungs)gerichts?\w*\s+[A-ZÄÖÜ][\wäöüß-]+"
    r"|Bundesgerichtshof|Bundesverfassungsgericht"
    r"|Staatsanwaltschaft\s+[A-ZÄÖÜ][\wäöüß-]+)"
)
NUR_KENNUNG = re.compile(
    r"^(?:Urkundenrolle|Aktenzeichen|Az\.|Geschäftszeichen|Tgb\.|Js-Az|Beschluss:|Bruttovergütung)\b",
    re.IGNORECASE,
)


def extract_entity(zeile: str, typ: str, teile: list[str]) -> str:
    """Holt den Absendernamen aus der Zeile, nicht die Dokumentbezeichnung."""
    if NUR_KENNUNG.match(zeile.strip()):
        return ""
    if typ == "notariat":
        m = NOTAR_NAME.search(zeile)
        return m.group(1).strip() if m else ""
    if typ == "gericht":
        m = GERICHT_NAME.search(zeile)
        if m:
            name = m.group(1).strip()
            # "des Amtsgerichts Leipzig" -> "Amtsgericht Leipzig"
            return re.sub(r"gerichts\s", "gericht ", name)
    if typ in {"unternehmen", "kanzlei", "kasse", "sachverstaendig"}:
        m = FIRMA_NAME.search(zeile) or ANWALT_SOLO.search(zeile)
        # Ohne erkennbaren Firmennamen lieber keinen Briefkopf als einen
        # falschen aus der Dokumentueberschrift.
        return m.group(1).strip() if m else ""
    kern = re.sub(r"^#+\s*", "", teile[0] if teile else zeile).strip()
    # Gerichtsentscheidungen: "Landgericht Oldenburg - Urteil vom ..."
    return re.split(r"\s+[—–-]\s+", kern)[0].strip()


def _rule(paragraph, hexcolor: str, size: int = 8) -> None:
    """Zieht eine feine Linie unter den Absatz."""
    pPr = paragraph._p.get_or_add_pPr()
    borders = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), str(size))
    bottom.set(qn("w:space"), "2")
    bottom.set(qn("w:color"), hexcolor)
    borders.append(bottom)
    pPr.append(borders)


def _page_field(paragraph) -> None:
    """Fuegt ein echtes Seitenzahlfeld ein."""
    run = paragraph.add_run()
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = " PAGE "
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    run._r.append(begin)
    run._r.append(instr)
    run._r.append(end)


def apply_briefkopf(doc, absender: Absender, aktenzeichen: str = "") -> None:
    """Setzt Kopf- und Fusszeile im Stil des jeweiligen Absenders."""
    farbe = RGBColor.from_string(absender.farbe)
    section = doc.sections[0]
    section.header_distance = Cm(1.2)
    section.footer_distance = Cm(1.0)

    header = section.header
    # Ohne diese Entkopplung erbt der Abschnitt eine leere Kopfzeile und der
    # Briefkopf bleibt beim Rendern unsichtbar.
    header.is_linked_to_previous = False
    kopf = header.paragraphs[0]
    kopf.text = ""
    kopf.alignment = WD_ALIGN_PARAGRAPH.LEFT
    name = absender.name.upper() if absender.kapitaelchen else absender.name
    run = kopf.add_run(name)
    run.font.name = "Times New Roman"
    run.font.size = Pt(11 if absender.kapitaelchen else 12)
    run.font.bold = True
    run.font.color.rgb = farbe

    zeile2 = absender.zusatz or TYP_STIL[absender.typ][0]
    if aktenzeichen:
        zeile2 = f"{zeile2} · {aktenzeichen}" if zeile2 else aktenzeichen
    unter = header.add_paragraph()
    unter.alignment = WD_ALIGN_PARAGRAPH.LEFT
    r2 = unter.add_run(zeile2)
    r2.font.name = "Times New Roman"
    r2.font.size = Pt(8)
    r2.font.color.rgb = GREY
    _rule(unter, absender.farbe, size=8 if absender.kapitaelchen else 6)

    footer = section.footer
    footer.is_linked_to_previous = False
    fp = footer.paragraphs[0]
    fp.text = ""
    fp.alignment = WD_ALIGN_PARAGRAPH.LEFT
    tabs = fp.paragraph_format.tab_stops
    tabs.add_tab_stop(Cm(16.0), WD_TAB_ALIGNMENT.RIGHT)
    left = fp.add_run(absender.name)
    left.font.name = "Times New Roman"
    left.font.size = Pt(8)
    left.font.color.rgb = GREY
    sep = fp.add_run("\tSeite ")
    sep.font.name = "Times New Roman"
    sep.font.size = Pt(8)
    sep.font.color.rgb = GREY
    _page_field(fp)
    for r in fp.runs:
        r.font.name = "Times New Roman"
        r.font.size = Pt(8)
        r.font.color.rgb = GREY


# Nur kurze Kopfzeilen tragen ein Aktenzeichen. In Fliesstextabsaetzen stehen
# dieselben Woerter mitten im Satz und lieferten sonst abgeschnittene Fetzen.
AKTENZEICHEN = re.compile(
    r"(?:Aktenzeichen|Az\.|Geschäftszeichen|Urkundenrolle(?:\s+Nr\.?)?)[:\s]+"
    r"([A-Za-z0-9][^\n,;()]{2,32})",
    re.IGNORECASE,
)
# Gerichtliche Aktenzeichen wie "5 O 458/25", "8 U 63/26" oder "IV ZR 256/25".
GERICHTS_AZ = re.compile(r"\b((?:[IVX]+ )?\d{0,3}\s?[A-Za-z]{1,4}\s\d{1,5}/\d{2,4})\b")
MAX_KOPFZEILE = 120


def detect_aktenzeichen(paragraphs: list[str]) -> str:
    for p in paragraphs[:10]:
        zeile = p.strip()
        if not zeile or len(zeile) > MAX_KOPFZEILE:
            continue
        m = AKTENZEICHEN.search(zeile)
        if m:
            return re.sub(r"\s+", " ", m.group(0)).strip().rstrip(".")
        m = GERICHTS_AZ.search(zeile)
        if m:
            return re.sub(r"\s+", " ", m.group(1)).strip()
    return ""
